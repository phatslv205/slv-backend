import eventlet
eventlet.monkey_patch()
from app import app, socketio
import os
import sys
import json
import bcrypt
import re
import requests
import imaplib
import email
from email.header import decode_header
import base64
import openai
import fitz  # PyMuPDF
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from flask import Response
from openai_config import call_gpt_viet
from openai import OpenAI
from dotenv import load_dotenv
from flask import url_for
import random
from extensions import db
from models.user import User
from sqlalchemy import text
from random import choices
from extensions import db, migrate, socketio, login_manager
@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, user_id)
from flask_login import login_user
from flask_migrate import Migrate
from sqlalchemy import and_, or_, func
from models.transaction import Transaction
from sqlalchemy import func
from models.friend import Friend
from models.message import Message
from models.leaderboard import NoiTuLeaderboard
from models.saved_chats import SavedChat
from models.leaderboard import update_leaderboard
from models.smartdoc import SmartDoc
from flask import request, send_file, render_template, session
from docx import Document
from models.daily_usage_stats import DailyUsageStats,update_daily_usage
from models.blocked_user_log import BlockedUserLog
from models.dev_chat_history import DevChatHistory
from models.feedback import Feedback
from models.user_memory import UserMemoryItem
from flask_login import login_required
from models.daily_usage_stats import finalize_daily_stats
from models.noi_tu_game_history import NoiTuGameHistory
from apscheduler.schedulers.background import BackgroundScheduler
from datetime import date
from models.image_history import ImageHistory
from flask_login import current_user
import datetime as dt
from datetime import datetime, timedelta, timezone
from app import app, db
from models.chat_history import ChatSession
import os, openai, tempfile
from openai_config import AI_PERSONALITIES
GPT35_KEYS = os.getenv("GPT35_KEYS", "").split(",")
GPT4O_KEYS = os.getenv("GPT4O_KEYS", "").split(",")
ALL_KEYS = [key.strip() for key in GPT35_KEYS + GPT4O_KEYS if key.strip()]
if not ALL_KEYS:
    raise Exception("Không tìm thấy bất kỳ API key nào trong GPT35_KEYS hoặc GPT4O_KEYS")
openai.api_key = random.choice(ALL_KEYS)
from config import Config, OPENAI_KEYS
from models import db  
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy.orm import Session
from models.friend_request import FriendRequest
from flask_socketio import SocketIO, join_room, emit
from flask_socketio import SocketIO, emit, join_room
from flask_socketio import emit 
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
from utils_shared import normalize_package
from pytz import timezone
from zoneinfo import ZoneInfo          
VN_TZ = ZoneInfo("Asia/Ho_Chi_Minh")   
UTC_TZ = ZoneInfo("UTC")
from datetime import timedelta,timezone
from flask import request, jsonify
from werkzeug.utils import secure_filename
from flask import request, session, redirect, render_template
from gpt_vision_ocr import extract_with_gpt_vision
from openai_config import call_gpt_lite


def reset_image_quota_if_needed(user):
    now = datetime.now()
    if not user.image_quota_reset_at or user.image_quota_reset_at.date() < now.date():
        user.image_quota_today = 10
        user.image_quota_reset_at = now
        db.session.commit()

def generate_image_from_prompt(prompt_text):
    try:
        client = create_openai_client()
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt_text,
            size="1024x1024",
            n=1
        )
        image_url = response.data[0].url
        if not image_url:
            return None

        img_data = requests.get(image_url).content
        filename = f"img_{uuid.uuid4().hex[:8]}.png"
        save_path = os.path.join("static", "images", "uploads", filename)

        with open(save_path, "wb") as f:
            f.write(img_data)

        return f"/static/images/uploads/{filename}"

    except Exception:
        return None


def rewrite_prompt_for_image(user_text):
    forbidden_keywords = [
        "gương mặt", "ghép mặt", "gương của tôi", "ảnh của tôi",
        "dựa trên ảnh", "dựa vào ảnh", "ảnh thật", "hình thật",
        "ảnh chụp", "photo of me", "real face", "real photo", "combine with my face"
    ]

    lower_text = user_text.lower()
    if any(keyword in lower_text for keyword in forbidden_keywords):
        return None

    if re.search(r"(phép toán|câu toán|bài toán|đề toán|toán học|toán lớp)", user_text, re.IGNORECASE):
        user_text += (
            "\n👉 Yêu cầu: Tạo ảnh một bảng trắng đẹp có 3 phép toán rõ ràng như:\n"
            "- 8 + 4 = ?\n"
            "- 6 × 7 = ?\n"
            "- 9 − 3 = ?\n"
            "Trình bày như giảng bài, bố cục sạch sẽ, ánh sáng tốt."
        )

    system_instruction = (
        "Bạn là chuyên gia viết prompt cho AI vẽ hình ảnh (DALL·E). "
        "Viết lại mô tả thành prompt tiếng ANH ngắn gọn, bắt đầu bằng danh từ mô tả chủ thể hoặc cảnh vật.\n"
        "- KHÔNG được nói gì ngoài prompt.\n"
        "- KHÔNG giải thích, KHÔNG chào hỏi, KHÔNG giới thiệu bản thân.\n"
        "- KHÔNG được viết: 'I'm here to help' hay 'I can help you'.\n"
        "- Chỉ trả về đúng 1 dòng prompt bằng tiếng Anh mô tả ảnh thật đẹp, cảm xúc, bố cục rõ, ánh sáng tốt."
    )

    try:
        client = create_openai_client(model="gpt-4o")
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_instruction},
                {"role": "user", "content": user_text}
            ],
            temperature=0.7
        )
        prompt = response.choices[0].message.content.strip()
        if not prompt or "i'm here" in prompt.lower() or prompt.lower().startswith("i ") or len(prompt) < 15:
            return None
        return prompt
    except Exception:
        return None


def run_auto_approval():
    while True:
        auto_approve_transactions()
        time.sleep(60)
def cleanup_old_chats():
    folder = "chat_history"
    max_age = 24 * 60 * 60  # 24 giờ
    now = time.time()

    for filename in os.listdir(folder):
        path = os.path.join(folder, filename)
        if os.path.isfile(path):
            if now - os.path.getmtime(path) > max_age:
                os.remove(path)

from datetime import datetime

def save_chat_sql(user_id, session_id, history):
    if len(history) < 2:
        return  # Không có gì mới để lưu

    last_two = history[-2:]

    for msg in last_two:
        role = msg["role"]
        content = msg.get("content")
        image_url = msg.get("image_url")

        # ⚠️ Nếu người dùng gửi ảnh nhưng không có nội dung → gắn nhãn đặc biệt để hiển thị fallback
        if role == "user" and not content and image_url:
            content = "__image_only__"

        # ❌ Nếu không có ảnh và không có nội dung thì bỏ qua luôn (không lưu dòng trống)
        if not content and not image_url:
            continue

        chat = ChatHistory(
            user_id=user_id,
            session_id=session_id,
            role=role,
            content=content,
            timestamp=datetime.utcnow(),
            image_url=image_url,
            reply_to=msg.get("reply_to"),
            is_saved=False 
        )
        db.session.add(chat)

    db.session.commit()


OTP_FILE = "otp_codes.json"


def now_vn():
    return datetime.utcnow() + timedelta(hours=7)



UPLOAD_FOLDER = 'static/images/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

import time
from image_cleaner import clean_old_images
import os
import time
def clean_old_images():
    folder = 'static/images/uploads'
    now = time.time()
    threshold = 7 * 86400  # 7 ngày

    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            if now - os.path.getmtime(file_path) > threshold:
                os.remove(file_path)
from openai_config import create_openai_client
import smtplib
from email.mime.text import MIMEText
def send_otp_email(to_email, otp_code):
    subject = "Mã xác thực OTP"
    body = f"Mã xác thực của bạn là: {otp_code}\nThời hạn hiệu lực: 5 phút."
    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = "SolverViet <your_email@gmail.com>"  # thay bằng email thật
    msg["To"] = to_email

    try:
        server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        server.login("your_email@gmail.com", "your_app_password")
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        return False

def send_user_otp_email(to_email, otp_code):
    subject = "Mã xác thực OTP từ SolverViet"

    html =f"""
<html>
<body style="font-family: Arial, sans-serif; color: #333; background-color: #f4f6f8; padding: 0; margin: 0;">
    <div style="max-width: 500px; margin: auto; background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
        
        <div style="text-align: center; margin-bottom: 20px;">
            <img src="https://i.ibb.co/Z2hCr6B/logo-solverviet.png" alt="SolverViet" style="width: 80px;"/>
        </div>
        
        <h2 style="color: #2e7d32; text-align: center;">Xác thực tài khoản SolverViet</h2>
        <p>Xin chào,</p>
        <p>Chúng tôi đã nhận được yêu cầu xác thực đăng nhập hoặc thao tác quan trọng từ bạn.</p>
        
        <p style="font-size: 16px;"><strong>Mã OTP của bạn là:</strong></p>
        
        <div style="font-size: 26px; font-weight: bold; color: #d32f2f; background-color: #fff3f3; border: 1px dashed #d32f2f; padding: 12px 20px; display: inline-block; border-radius: 6px;">
            {otp_code}
        </div>
        
        <p style="margin-top: 20px;">Mã này có hiệu lực trong <strong>5 phút</strong>.</p>
        
        <p style="color: #e53935; font-weight: bold;">⚠️ Vui lòng không chia sẻ mã này với bất kỳ ai, kể cả người tự xưng là từ SolverViet.</p>
        
        <hr style="margin: 20px 0;">
        
        <p style="font-size: 13px; color: gray;">
            Nếu bạn không yêu cầu mã này, vui lòng bỏ qua email này hoặc liên hệ hỗ trợ.
        </p>
        <p style="margin-top: 20px;">Trân trọng,<br><strong>SolverViet 🇻🇳</strong></p>
    </div>
</body>
</html>
"""


    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = f"SolverViet <{os.getenv('EMAIL_ADDRESS')}>"
    msg["To"] = to_email
    msg.attach(MIMEText(html, "html", "utf-8"))

    try:
        server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        server.login(os.getenv("EMAIL_ADDRESS"), os.getenv("EMAIL_PASSWORD"))
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        return False


import unicodedata
import random
import string

def generate_otp():
    from random import randint
    return str(random.randint(100000, 999999))


TELEGRAM_TOKEN = "7580016404:AAHoWnvRElJD3BXkZyxZQ4A4z34qXB-3s54"
TELEGRAM_CHAT_ID = "6894067779"
def send_telegram_message(message):
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    data = {
        "chat_id": TELEGRAM_CHAT_ID,
        "text": message,
        "parse_mode": "HTML"
    }
    try:
        requests.post(url, data=data)
    except Exception:
        pass
#HÀM GỬI EMAIL
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from ip_blocker import record_ip, is_ip_blocked
from flask import get_flashed_messages
from flask import Flask, render_template, request, redirect, url_for, session, flash, get_flashed_messages
REQUESTS_FILE = "requests.json"
from flask import Flask, render_template, request, session, redirect, url_for
import time
import threading
from flask import flash 
import uuid
# === Thêm thư mục hiện tại vào sys.path để import các module nội bộ ===
sys.path.append(os.path.abspath(os.path.dirname(__file__)))

# ✅ IMPORT các hàm xử lý quản lý người dùng
from admin_utils import is_vip
def generate_unique_referral_code(length=8):
    """Sinh mã giới thiệu duy nhất, không trùng."""
    while True:
        code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))
        if not User.query.filter_by(referral_code=code).first():
            return code
def load_user_messages(user_id):
    inbox_path = f"user_data/{user_id}_inbox.txt"
    if os.path.exists(inbox_path):
        with open(inbox_path, "r", encoding="utf-8") as f:
            return [msg.strip() for msg in f.read().split("---\n\n") if msg.strip()]
    return []
def cap_nhat_trang_thai_vip(user: User):
    now = datetime.utcnow()
    user.vip_gpt_ai = bool(user.vip_until_gpt and now <= user.vip_until_gpt)
    user.vip_ai_lite = bool(user.vip_until_lite and now <= user.vip_until_lite)
    return {
        "vip_gpt_ai": user.vip_gpt_ai,
        "vip_ai_lite": user.vip_ai_lite,
    }
# ===== IMPORT TOÀN BỘ =====
from flask import Flask
from config import Config
from extensions import db
# ====== FLASK APP ======
app = Flask(__name__)
socketio = SocketIO(app, async_mode="threading")  
app.config.from_object(Config)
db.init_app(app)
migrate.init_app(app, db)
app.secret_key = 'b@o_m@t_2025_🔥' 
DATA_FILE = 'friends_data.json'
login_manager.init_app(app)
app.secret_key = Config.SECRET_KEY
app.config['SECRET_KEY'] = Config.SECRET_KEY
app.permanent_session_lifetime = timedelta(days=30)
# Job 15 phút: không cần app context
from flask import render_template, request, redirect, abort, session
def is_admin():
    return session.get("is_admin") is True   
from functools import wraps
def admin_only(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not is_admin():
            return redirect("/admin_login") 
        return f(*args, **kwargs)
    return decorated_function
@app.route("/admin/daily_stats")
@admin_only
def daily_stats():
    page = request.args.get("page", 1, type=int)
    per_page = 20
    pagination = DailyUsageStats.query.order_by(DailyUsageStats.date.desc()).paginate(page=page, per_page=per_page)
    return render_template("admin_daily_stats.html", pagination=pagination, stats=pagination.items)
def delete_old_stats():
    threshold = datetime.utcnow().date() - timedelta(days=180)
    deleted = DailyUsageStats.query.filter(DailyUsageStats.date < threshold).delete()
    db.session.commit()
@app.route("/admin/delete_stat/<stat_id>", methods=["POST"])
@admin_only
def delete_stat(stat_id):
    stat = DailyUsageStats.query.get(stat_id)
    if stat:
        db.session.delete(stat)
        db.session.commit()
    return redirect(url_for("daily_stats"))

from sqlalchemy import or_, func, String
@socketio.on("join_admin")
def join_admin():
    if session.get("is_admin"):
        join_room("admin_dashboard")
def get_time_slot(dt):
    return dt.strftime("%H:00")
@app.route("/admin_users")
@admin_only
def admin_users():
    PER_PAGE = 200
    page   = request.args.get("page", 1, type=int)
    search = request.args.get("search", "").strip().lower()
    base_q = User.query
    if search:
        base_q = base_q.filter(
            func.lower(User.username).like(f"%{search}%") |
            func.lower(User.fullname).like(f"%{search}%") |
            func.lower(User.user_id.cast(String)).like(f"%{search}%")
        )
    total_users = base_q.count()
    active_users = base_q.filter_by(online=True).count()
    gpt_users = base_q.filter(User.vip_gpt.isnot(None)).count()
    blocked_users = base_q.filter_by(over_quota_block=True).count()
    heavy_users = base_q.filter(User.gpt_usage_today > 100).count()
    stats_today = DailyUsageStats.query.filter_by(date=date.today()).first()
    if not stats_today:
        stats_today = DailyUsageStats(slv_count=0, lite_count=0)
    base_q = base_q.order_by(
        User.over_quota_block.desc(),
        (User.gpt_usage_today >= 200).desc(),
        User.online.desc(),
        User.user_id.asc()
    )
    today = now_vn().strftime("%Y-%m-%d")
    now = now_vn()
    expired_users = User.query.filter(
        or_(
            and_(User.vip_gpt.isnot(None), User.vip_until_gpt < now),
            and_(User.vip_ai_lite == True, User.vip_until_lite < now)
        )
    ).all()

    for u in expired_users:
        if u.vip_until_gpt and u.vip_until_gpt < now:
            u.vip_gpt = None
            u.vip_until_gpt = None
        if u.vip_ai_lite and u.vip_until_lite and u.vip_until_lite < now:
            u.vip_ai_lite = False
            u.vip_until_lite = None
    db.session.commit()
    lite_users = User.query.filter(
        User.vip_ai_lite == True,
        User.vip_until_lite > now_vn()
    ).all()
    for u in lite_users:
        if u.lite_date != today:
            u.lite_usage = 0
            u.lite_date = today
    db.session.commit()
    users_paginated = base_q.paginate(page=page, per_page=PER_PAGE)
    users = users_paginated.items
    users_data = []
    for user in users:
        u = {
            "username": user.username,
            "fullname": user.fullname,
            "ai_personality": user.ai_personality or "Không có",
            "bio": user.bio,
            "user_id": user.user_id,
            "referral_code": user.referral_code or "—",
            "online": getattr(user, "online", False),
            "is_blocked": getattr(user, "is_blocked", False),
            "vip_lite_display": user.vip_until_lite.strftime("%Y-%m-%d %H:%M:%S") if user.vip_until_lite else "Không có",
            "vip_gpt_display": user.vip_until_gpt.strftime("%Y-%m-%d %H:%M:%S") if user.vip_until_gpt else "Không có",
            "vip_gpt": user.vip_gpt or "",
            "over_quota_block": getattr(user, "over_quota_block", False),
            "over_quota_block_at": user.over_quota_block_at,
            "quota_alert": (  user.vip_gpt in ["15day", "30day"] and not user.gpt_unlimited and user.gpt_usage_today >= 200),
        }
        gpt_type = user.vip_gpt
        used = getattr(user, "gpt_usage_today", 0) or 0

        if getattr(user, "vip_ai_lite", False) and getattr(user, "vip_until_lite", None):
            lite_used = getattr(user, "lite_usage", 0)
            lite_limit = getattr(user, "vip_lite_daily_limit", 70)
            u["gpt_remaining"] = "None"
            u["lite_remaining"] = f"{lite_limit - lite_used}/{lite_limit}"
        elif gpt_type == "5day":
            u["gpt_remaining"] = f"{100 - used}/100"
            u["lite_remaining"] = "Không có"
        elif gpt_type in ["15day", "30day"]:
            u["gpt_remaining"] = "∞"
            u["lite_remaining"] = "Không có"
        else:
            u["gpt_remaining"] = "Không có"
            u["lite_remaining"] = "Không có"

        if gpt_type == "5day":
            quota_remain = 100 - used
            u["gpt_quota"] = f"{quota_remain}/100"
            u["quota_alert"] = False
        elif gpt_type in ("15day", "30day") and not getattr(user, "gpt_unlimited", False):
            quota_remain = 200 - used
            u["gpt_quota"] = f"{quota_remain}/200"
            if quota_remain <= 0 and not user.over_quota_block:
                user.over_quota_block = True
                user.over_quota_block_at = now_vn()
                db.session.commit()
                socketio.emit("update_quota_alert")
    
                u["over_quota_block"] = True
                u["over_quota_block_at"] = user.over_quota_block_at
                u["quota_alert"] = False
            else:
                u["quota_alert"] = quota_remain <= 0
        else:
            u["gpt_quota"] = "∞" if getattr(user, "gpt_unlimited", False) else "N/A"
            u["quota_alert"] = False

        u["gpt_usage_today"] = user.gpt_usage_today
        u["lite_usage"] = user.lite_usage

        users_data.append(u)
    alert_users = User.query.filter(
        or_(
            User.over_quota_block == True,
            User.gpt_usage_today >= 200
        )
    ).order_by(User.gpt_usage_today.desc()).limit(50).all()

    return render_template(
        "admin_users.html",
        alert_users=alert_users,
        timedelta=timedelta,
        users=users_data,
        user_count=users_paginated.total,
        pagination=users_paginated,
        search=search,
        stats_today=stats_today,
        stats={
            "total": total_users,
            "active": active_users,
            "gpt": gpt_users,
            "blocked": blocked_users,
            "heavy": heavy_users,
        }
    )
import pytz, threading, time
from datetime import timezone
def auto_unblock_loop():
    with app.app_context():
        while True:
            users = User.query.filter_by(over_quota_block=True).all()
            now_utc = datetime.now(pytz.UTC)
            count_unblocked = 0

            for user in users:
                block_time = user.over_quota_block_at
                if not block_time:
                    continue

                # Đảm bảo có timezone (nếu bị thiếu)
                if block_time.tzinfo is None:
                    block_time = pytz.timezone("Asia/Ho_Chi_Minh").localize(block_time).astimezone(pytz.UTC)
                else:
                    block_time = block_time.astimezone(pytz.UTC)

                # Tính thời điểm cần mở khóa
                unblock_time = block_time + timedelta(days=1)

                if now_utc >= unblock_time:
                    user.over_quota_block = False
                    user.over_quota_block_at = None
                    user.gpt_usage_today = 0
                    db.session.commit()
                    count_unblocked += 1

                    socketio.emit("user_usage_update", {
                        "user_id": user.user_id,
                        "username": user.username,
                        "gpt_usage": user.gpt_usage_today,
                        "lite_usage": user.lite_usage,
                        "over_quota_block": user.over_quota_block,
                        "vip_gpt": user.vip_gpt or "",
                        "vip_lite": bool(user.vip_ai_lite),
                        "quota_alert": False
                    }, room="admin_dashboard")

            time.sleep(10)

@app.route("/admin_users/quota_alert")
@admin_only
def quota_alert_partial():
    alert_users = User.query.filter(
        (User.gpt_usage_today >= 200) | (User.over_quota_block == True)
    ).limit(50).all()
    return render_template("admin/_quota_alert_partial.html", alert_users=alert_users, timedelta=timedelta)                                                                                                                                                                                                                                                                                                                       
from sqlalchemy import and_
from sqlalchemy.sql import cast
from sqlalchemy.types import Date
from pytz import timezone
@app.route("/admin_users/block_chat_quota", methods=["POST"])
@admin_only
def block_chat_quota():
    username = request.form.get("username")
    block    = request.form.get("block_chat") == "1"

    user = User.query.filter_by(username=username).first()
    if user:
        user.over_quota_block = block
        user.over_quota_block_at = now_vn() if block else None
        log_added = False  
        if block and user.gpt_usage_today > 0: 
            block_time = user.over_quota_block_at.astimezone(timezone("Asia/Ho_Chi_Minh"))
            socketio.emit("update_quota_alert")
            time_slot = get_time_slot(block_time)
            existing_log = BlockedUserLog.query.filter(
                and_(
                    BlockedUserLog.user_id == user.user_id,
                    BlockedUserLog.time_slot == time_slot,
                    cast(BlockedUserLog.block_time, Date) == block_time.date()
                )
            ).first()
            if not existing_log:
                log = BlockedUserLog(
                    user_id=user.user_id,
                    username=user.username,
                    block_time=block_time,
                    time_slot=time_slot
                )
                db.session.add(log)
                db.session.flush()  
                log_added = True
        if not block:
            user.gpt_usage_today = 0
        db.session.commit()
        socketio.emit("user_usage_update", {
            "user_id": user.user_id,
            "username": user.username,
            "gpt_usage": user.gpt_usage_today,
            "lite_usage": user.lite_usage,
            "vip_gpt": user.vip_gpt or "",
            "vip_lite": bool(user.vip_ai_lite),
            "over_quota_block": user.over_quota_block,
            "over_quota_block_at": user.over_quota_block_at.isoformat() + "+07:00" if user.over_quota_block_at else None,
            "quota_alert": (
                user.vip_gpt in ["15day", "30day"]
                and not user.gpt_unlimited
                and user.gpt_usage_today >= 200
            )
        }, room="admin_dashboard")
    return "", 204

@app.route("/admin_users/auto_unblock_quota")
@admin_only
def auto_unblock_quota():
    expired_time = now_vn() - timedelta(hours=24)
    users = User.query.filter(
        User.over_quota_block == True,
        User.over_quota_block_at != None,
        User.over_quota_block_at < expired_time
    ).all()
    count = 0
    for user in users:
        user.over_quota_block = False
        user.over_quota_block_at = None
        user.gpt_usage_today = 0
        count += 1
        socketio.emit("user_usage_update", {
            "user_id": user.user_id,
            "username": user.username,
            "gpt_usage": user.gpt_usage_today,
            "lite_usage": user.lite_usage,
            "over_quota_block": user.over_quota_block,
            "vip_gpt": user.vip_gpt or "",
            "vip_lite": bool(user.vip_ai_lite),
            "quota_alert": False
        }, room="admin_dashboard")
    db.session.commit()
    return f"✅ Đã gỡ chặn tự động {count} người dùng sau 24h."

def parse_dt(s):
    if not s or s.lower() in ["none", "không có"]:
        return None
    s = s.strip()
    try:
        if len(s) == 10:
            s += " 23:59:59"
        return datetime.strptime(s, "%Y-%m-%d %H:%M:%S")
    except:
        return None
def parse_vip_duration(type_str):
    now = datetime.now()
    if type_str == "5day":
        return now + timedelta(days=4, hours=23, minutes=59, seconds=59)
    elif type_str == "15day":
        return now + timedelta(days=14, hours=23, minutes=59, seconds=59)
    elif type_str == "30day":
        return now + timedelta(days=29, hours=23, minutes=59, seconds=59)
    elif type_str == "7day":  
        return now + timedelta(days=6, hours=23, minutes=59, seconds=59)
    return None

@app.route("/admin_users/update", methods=["POST"])
@admin_only
def admin_users_update():
    old_name = request.form.get("username")
    new_name = request.form.get("new_username", old_name)

    vip_gpt = request.form.get("vip_gpt", "").strip()
    vip_lite = request.form.get("vip_lite", "").strip()

    vip_gpt_type = request.form.get("vip_gpt_type", "").strip()
    vip_lite_type = request.form.get("vip_lite_type", "").strip()

    gpt_unlimited = request.form.get("gpt_unlimited") == "on"

    user = User.query.filter_by(username=old_name).first()
    if user:
        if new_name != old_name:
            user.username = new_name

        # 🎯 Gói GPT
        if vip_gpt_type:
            gpt_time = parse_vip_duration(vip_gpt_type)
        else:
            gpt_time = parse_dt(vip_gpt)

        # 🎯 Gói Lite
        if vip_lite_type == "none":
            lite_time = None
        elif vip_lite_type:
            lite_time = parse_vip_duration(vip_lite_type)
        else:
            lite_time = parse_dt(vip_lite)

        # ❗ Nếu cả GPT và Lite đều được set thì vẫn ưu tiên GPT
        if gpt_time and lite_time:
            lite_time = None

        user.vip_until_gpt = gpt_time
        user.vip_gpt = vip_gpt_type or None
        user.vip_gpt_ai = bool(gpt_time)

        user.vip_until_lite = lite_time
        user.vip_ai_lite = bool(lite_time)
        if lite_time:
            user.lite_usage = 0
        db.session.commit()
 
    return redirect("/admin_users")

@app.route("/admin_users/toggle_block", methods=["POST"])
@admin_only
def admin_toggle_block():
    username = request.form.get("username")
    if not username:
        return "Thiếu tên người dùng", 400

    user = User.query.filter_by(username=username).first()
    if not user:
        return "Không tìm thấy người dùng", 404

    # ✅ Đảo ngược trạng thái block
    user.is_blocked = not getattr(user, "is_blocked", False)
    db.session.commit()

    return redirect("/admin_users")


def get_user_type():
    user_id = session.get("user_id")
    if not user_id:
        return "guest"

    user = User.query.filter_by(user_id=user_id).first()

    if not user:
        return "guest"

    try:
        if user.vip_until and datetime.strptime(user.vip_until, "%Y-%m-%d").date() >= datetime.now().date():
            return "vip"
    except:
        pass

    return "logged_in"


def check_lite_usage(user):
    MAX_FREE = 15
    MID_LIMIT = 5
    DAILY_LIMIT = 30
    now = now_vn()
    today = now.strftime("%Y-%m-%d")

    username = session.get("username")
    if not username:
        usage = session.get("lite_usage", 0)
        if usage >= MAX_FREE:
            return False
        session["lite_usage"] = usage + 1
        return True

    user = User.query.filter_by(username=username).first()
    if not user:
        return False

    verified = user.is_verified or False
    rolls_left = user.ai_personality_rolls_left or 0
    free_uses = user.free_gpt_uses or 0

    # ✅ Gói GPT không giới hạn
    try:
        if user.vip_until_gpt and now <= datetime.strptime(user.vip_until_gpt, "%Y-%m-%d %H:%M:%S"):
            return True
    except:
        pass

    # ✅ Vừa xác thực xong thì cho qua lượt đầu
    if session.pop("just_verified", False):
        return True

    # ✅ Nếu chưa xác thực → chỉ 5 lượt đầu (trừ ai_personality_rolls_left)
    if not verified:
        if rolls_left <= 0:
            return "require_verification"
        user.ai_personality_rolls_left = rolls_left - 1
        user = db.session.merge(user) 
        db.session.commit()

        socketio.emit("user_usage_update", {
            "user_id": user.user_id,
            "username": user.username,
            "gpt_usage": user.gpt_usage_today,
            "free_gpt_uses": user.free_gpt_uses,
            "rolls_left": user.ai_personality_rolls_left,
            "vip_gpt": user.vip_gpt or "",
            "vip_lite": bool(user.vip_ai_lite)
        }, room="admin_dashboard")

        return True

    # ✅ Nếu đã xác thực → 15 lượt miễn phí (trừ free_gpt_uses)
    if verified and free_uses > 0:
        user.free_gpt_uses = free_uses - 1
        user = db.session.merge(user)
        db.session.commit()

        socketio.emit("user_usage_update", {
            "user_id": user.user_id,
            "username": user.username,
            "gpt_usage": user.gpt_usage_today,
            "free_gpt_uses": user.free_gpt_uses,
            "rolls_left": user.ai_personality_rolls_left,
            "vip_gpt": user.vip_gpt or "",
            "vip_lite": bool(user.vip_ai_lite)
        }, room="admin_dashboard")

        return True

    # ✅ Nếu đã mua gói Lite → 50 lượt/ngày
    try:
        if user.vip_until_lite and now <= user.vip_until_lite:
            if user.lite_date != today:
                user.lite_usage = 0
                user.lite_date = today
            usage = user.lite_usage or 0
            if usage < 70:
                user.lite_usage = usage + 1
                user = db.session.merge(user)
                db.session.commit()

                socketio.emit("user_usage_update", {
                    "user_id": user.user_id,
                    "username": user.username,
                    "gpt_usage": user.gpt_usage_today,
                    "lite_usage": user.lite_usage,
                    "free_gpt_uses": user.free_gpt_uses,
                    "rolls_left": user.ai_personality_rolls_left,
                    "vip_gpt": user.vip_gpt or "",
                    "vip_lite": bool(user.vip_ai_lite)
                }, room="admin_dashboard")

                return True
            else:
                return False
    except:
        pass
    if verified:
        has_lite = user.vip_until_lite and now <= user.vip_until_lite
        has_slv = user.vip_until_gpt and now <= user.vip_until_gpt
        if not has_lite and not has_slv:
            if user.daily_fallback_date != today:
                user.daily_fallback_date = today
                user.daily_fallback_uses = 0

            if user.daily_fallback_uses >= DAILY_LIMIT:
                return False

            user.daily_fallback_uses += 1
            user = db.session.merge(user)
            db.session.commit()

            socketio.emit("user_usage_update", {
                "user_id": user.user_id,
                "username": user.username,
                "daily_fallback_uses": user.daily_fallback_uses,
                "free_gpt_uses": user.free_gpt_uses,
                "rolls_left": user.ai_personality_rolls_left,
                "vip_gpt": user.vip_gpt or "",
                "vip_lite": bool(user.vip_ai_lite)
            }, room="admin_dashboard")

            return True

def check_gpt_usage(user):
    today = today_str()  # dùng giờ VN chuẩn
    gpt_until = user.vip_until_gpt
    gpt_type = user.vip_gpt

    if not gpt_until or not gpt_type:
        return False
    if now_vn() > gpt_until:
        return False

    if user.gpt_unlimited:
        return True

    if gpt_type in ["15day", "30day"]:
        # 🔁 Reset nếu qua ngày
        if user.gpt_usage_date != today:
            user.gpt_usage_today = 0
            user.gpt_usage_date = today
        return True

    if gpt_type == "5day":
        if user.gpt_usage_date != today:
            user.gpt_usage_today = 0
            user.gpt_usage_date = today

        if user.gpt_usage_today >= 100:
            return False
        return True

    return False
def get_al_usage():
    return session.get("al_uses", 0)
def increment_al_usage():
    session["al_uses"] = get_al_usage() + 1
# ====== USER MANAGEMENT ======#
def get_username():
    return session.get("username")
def is_admin():
    return session.get("is_admin") is True
def update_solves(username):
    user = User.query.filter_by(username=username).first()
    if user:
        user.solves += 1
        db.session.commit()
def get_today():
    return datetime.now().strftime('%Y-%m-%d')

def get_solve_today(username):
    today = get_today()
    if username == "admin":
        return 0  # Admin không bị giới hạn

    user = User.query.filter_by(username=username).first()
    if not user:
        return 0

    if not user.last_solve_date:
        user.last_solve_date = today
        user.solves = 0
        user.ads_used_today = 0
        db.session.commit()
    elif user.last_solve_date != today:
        if user.solves < 10:
            user.solves = max(0, user.solves - 5)
        else:
            user.solves = 0
        user.last_solve_date = today
        user.ads_used_today = 0
        db.session.commit()

    return user.solves
#==========BẢO TRÌ HỆ THỐNG==========#
def is_maintenance(feature):
    try:
        with open("maintenance_config.json", "r") as f:
            config = json.load(f)
            return config.get("all") or config.get(feature, False)
    except:
        return False
@app.route("/check_maintenance")
def check_maintenance():
    feature = request.args.get("feature", "all")
    return jsonify({"maintenance": is_maintenance(feature)})
@app.route("/maintenance")
def maintenance_page():
    return render_template("maintenance.html")


@app.route("/admin/bao-tri", methods=["GET"])
@admin_only
def bao_tri_router():
    return redirect("/admin/bao-tri-all")
    
@app.route("/admin/bao-tri-all", methods=["GET", "POST"])
@admin_only
def bao_tri_all():
    try:
        with open("maintenance_config.json", "r") as f:
            config = json.load(f)
    except FileNotFoundError:
        config = {
            "all": False,
            "gpt_chat": False,
            "chat_lite": False,
            "chat_ai_lite": False,
            "chat_ai_lite_daily": False,
            "home": False
        }

    if request.method == "POST":
        feature = request.form.get("feature")

        if feature and feature in config:
            # Toggle từng phần cụ thể
            config[feature] = not config[feature]
            flash(f"✅ Đã cập nhật bảo trì cho phần: {feature}", "success")
        else:
            # Toggle toàn hệ thống nếu không có phần cụ thể
            config["all"] = not config.get("all", False)
            flash("✅ Đã cập nhật trạng thái bảo trì toàn bộ hệ thống.", "success")

        with open("maintenance_config.json", "w") as f:
            json.dump(config, f, indent=2)

        return redirect("/admin/bao-tri-all")

    is_on = config.get("all", False)
    return render_template("bao_tri_all.html", is_on=is_on, config=config)

# ====== LOGIN / REGISTER / LOGOUT ======
import secrets
from flask import request, render_template, redirect, url_for, session
from datetime import datetime

@app.route("/login", defaults={"rid": None}, methods=["GET", "POST"])
@app.route("/login/<rid>", methods=["GET", "POST"])
def login(rid):
    if request.method == "GET" and rid is None:
        new_rid = secrets.token_urlsafe(6)  
        return redirect(url_for("login", rid=new_rid, message=request.args.get("message")))

    message = request.args.get("message")

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        user = User.query.filter_by(username=username).first()

        if user:
            if getattr(user, "is_blocked", False):
                error = ("Tài khoản của bạn đã bị khóa. Nếu đây là nhầm lẫn, vui lòng gửi khiếu nại tại "
                         "<a href='/appeal' style='color:#4ea6ff;'>đây</a>.")
                return render_template("login.html", error=error, rid=rid)

            # TODO: nếu dùng hash mật khẩu, thay so sánh này bằng check hash
            if user.password == password:
                if getattr(user, "wants_verification", False) and not getattr(user, "is_verified", False):
                    otp_code = generate_otp()
                    user.otp_code = otp_code
                    send_user_otp_email(user.email, otp_code)
                    session["pending_user"] = username
                    session["otp_sent"] = True
                    db.session.commit()
                    return redirect("/verify-otp")
                user.online = True
                user.last_seen = datetime.utcnow()
                login_user(user)
                session.permanent = True
                session["username"] = user.username
                session["user_id"] = user.user_id          
                session["user_uuid"] = user.user_id         
                session["vip_until_gpt"] = getattr(user, "vip_until_gpt", None)
                session["al_uses"] = 0 
                db.session.commit()
                if getattr(user, "vip_gpt_ai", False):
                    session["just_logged_in"] = True
                return redirect(url_for("home_page"))
        return render_template("login.html", error="Sai tài khoản hoặc mật khẩu.", rid=rid)
    return render_template("login.html", message=message, rid=rid)

from sqlalchemy.exc import IntegrityError
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
ADMIN_ALERT_EMAIL = EMAIL_ADDRESS

def normalize_gmail(email):
    email = email.strip().lower()
    if email.endswith("@gmail.com"):
        local, domain = email.split("@")
        local = local.split("+")[0].replace(".", "")
        return f"{local}@{domain}"
    return email

@app.route("/register", defaults={"rid": None}, methods=["GET", "POST"])
@app.route("/register/<rid>", methods=["GET", "POST"])
def register(rid):
    if request.method == "GET" and rid is None:
        new_rid = secrets.token_urlsafe(6)
        return redirect(url_for("register", rid=new_rid))

    if request.method == "POST":
        ip = request.headers.get('X-Forwarded-For', request.remote_addr)
        if ip and ',' in ip:
            ip = ip.split(',')[0].strip()

        username = request.form.get("username", "").strip().lower()
        email = request.form.get("email", "").strip()
        normalized_email = normalize_gmail(email)
        fullname = request.form.get("fullname", "").strip()
        password = request.form.get("password")
        confirm_password = request.form.get("confirm_password")
        ai_personality = request.form.get("ai_personality", "").strip()
        referral_code_used = request.form.get("referral_code", "").strip()
        want_verification = 'want_verification' in request.form
        existing_user = User.query.filter_by(normalized_email=normalized_email).first()
        if existing_user:
            return render_template("register.html", rid=rid, error="Gmail này đã được sử dụng để tạo tài khoản!")

        final_personality = ai_personality or "SLV"

        new_user = User(
            username=username,
            password=password,
            fullname=fullname,
            email=email,
            normalized_email=normalized_email,
            is_verified=False,
            wants_verification=want_verification,
            referral_code_used=referral_code_used,
            free_gpt_uses=5,
            ai_personality=final_personality,
            referral_code=generate_unique_referral_code()
        )

        db.session.add(new_user)
        try:
            db.session.commit()
        except IntegrityError:
            db.session.rollback()
            return render_template("register.html", rid=rid, error="Email này đã tồn tại trong hệ thống!")

        # Ghi lại IP
        count = record_ip(ip)
        if session.get("username") != "admin" and count >= 3:
            message = f"""
<b>Cảnh báo đăng ký SPAM</b>
🔢 IP: <code>{ip}</code>
🕒 Ngày: {datetime.now().strftime('%Y-%m-%d')}</code>
👤 Username: <code>{username}</code>
📧 Email: <code>{email}</code>
💥 Số lần tạo tài khoản: {count} (hạn mức: 3)
            """
            send_telegram_message(message.strip())
            send_spam_alert_email(ip, username, email, count)

        return redirect(url_for("login", rid=secrets.token_urlsafe(6)))

    return render_template("register.html", rid=rid)

def send_spam_alert_email(ip, username, email, count):
    subject = "⚠️ Cảnh báo SPAM: IP tạo nhiều tài khoản"
    body = f"""
[SPAM PHÁT HIỆN]
🔢 IP: {ip}
🕒 Ngày: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
👤 Username: {username}
📧 Email: {email}
💥 Số lần tạo tài khoản: {count} (hạn mức: 3)
"""

    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = EMAIL_ADDRESS
    msg["To"] = ADMIN_ALERT_EMAIL

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            server.send_message(msg)
    except Exception:
        pass
@app.route("/random_ai_personality", methods=["POST"])
def random_ai_personality():
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"success": False, "message": "Bạn chưa đăng nhập."})

    user = User.query.filter_by(user_id=user_id).first()
    if not user:
        return jsonify({"success": False, "message": "Không tìm thấy người dùng."})

    from openai_config import AI_PERSONALITIES
    chosen = random.choice(AI_PERSONALITIES)

    user.ai_personality = chosen
    db.session.commit()

    return jsonify({"success": True, "new_ai": chosen})


from models.user import User
from extensions import db

from openai_config import AI_PERSONALITIES
@app.route("/save_reward_info", methods=["POST"])
def save_reward_info():
    if "username" not in session:
        return jsonify(success=False, message="Chưa đăng nhập.")

    data = request.get_json()
    fullname = data.get("fullname", "").strip()
    stk = data.get("stk", "").strip()
    bank = data.get("bank", "").strip()

    if not fullname or not stk or not bank:
        return jsonify(success=False, message="Thiếu thông tin.")

    user = User.query.filter_by(username=session["username"]).first()
    if not user:
        return jsonify(success=False, message="Không tìm thấy tài khoản.")

    user.reward_fullname = fullname
    user.reward_stk = stk
    user.reward_bank = bank
    db.session.commit()
    return jsonify(success=True)
def generate_userinfo_slug():
    return str(uuid.uuid4()) 
@app.route("/user-info", methods=["GET"])
def redirect_user_info():
    slug = generate_userinfo_slug()  # uuid
    return redirect(f"/profileSecureGate/{slug}")
# Cho 2 URL dùng chung 1 handler
@app.route("/profileSecureGate/<slug>", methods=["GET", "POST"])
@app.route("/user-info", methods=["GET", "POST"])
def user_info(slug=None):
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")
    if request.method == "GET" and slug is None:
        new_slug = str(uuid.uuid4())
        return redirect(url_for("user_info", slug=new_slug))

    if request.method == "POST":
        # lấy dữ liệu form và cập nhật
        email = request.form.get("email")
        password = request.form.get("new_password")
        fullname = request.form.get("fullname")
        birthyear = request.form.get("birthyear")
        ai_personality = request.form.get("ai_personality")

        if email: user.email = email
        if password: user.password = password
        if fullname: user.fullname = fullname
        if birthyear: user.birthyear = birthyear

        if ai_personality and user.vip_gpt and user.vip_gpt.startswith("SLV"):
            user.ai_personality = ai_personality

        avatar = request.files.get("avatar")
        if avatar and avatar.filename:
            filename = secure_filename(f"{username}_avatar.png")
            path = os.path.join("static/images/avatars", filename)
            avatar.save(path)
            user.avatar_url = f"/static/images/avatars/{filename}"

        db.session.commit()
        flash("Đã cập nhật thông tin thành công!", "success")

        # Sau khi lưu, luôn đưa người dùng về URL có slug mới
        back_slug = slug or str(uuid.uuid4())
        return redirect(url_for("user_info", slug=back_slug))
    
    # request.method == GET với slug đã có
    return render_template(
        "user_info.html",
        user=user,
        username=user.username,
        now=datetime.now(),
        personalities=AI_PERSONALITIES,
        slug=slug,  # nếu template có dùng
    )

@app.route("/user-form-fragment")
def user_form_fragment():
    user = User.query.filter_by(username=session["username"]).first()
    return render_template("user_form.html", user=user, now=datetime.now(), username=user.username)
@app.route("/user-info-popup", methods=["GET"])
def user_info_popup():
    username = session.get("username")
    if not username:
        return "Unauthorized", 403

    user = User.query.filter_by(username=username).first()
    if not user:
        return "User not found", 404

    now = datetime.utcnow()
    current_plan = "Không có"
    days_left = 0

    if user.vip_gpt_ai and user.vip_until_gpt and user.vip_until_gpt > now:
        current_plan = f"SLV ({user.vip_gpt})"
        days_left = (user.vip_until_gpt.date() - now.date()).days
    elif user.vip_ai_lite and user.vip_until_lite and user.vip_until_lite > now:
        current_plan = "Lite"
        days_left = (user.vip_until_lite.date() - now.date()).days

    return render_template(
        "user_info_popup.html",
        user=user,
        personalities=AI_PERSONALITIES,
        current_plan=current_plan,
        days_left=days_left
    )

@app.route("/update_user_info", methods=["POST"])
def update_user_info():
    username = session.get("username")
    if not username:
        return jsonify({"success": False, "error": "Chưa đăng nhập"}), 401

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "error": "Không tìm thấy user"}), 404

    data = request.get_json()
    email = data.get("email")
    password = data.get("password")
    fullname = data.get("fullname")
    ai_personality = data.get("ai_personality")

    if email: user.email = email
    if password: user.password = password
    if fullname: user.fullname = fullname

    # Chỉ cho đổi tính cách AI nếu user đang có gói SLV
    if ai_personality and user.vip_gpt and user.vip_gpt.startswith("SLV"):
        user.ai_personality = ai_personality

    db.session.commit()
    return jsonify({"success": True})

def get_current_user():
    username = session.get("username")
    if not username:
        return None
    return User.query.filter_by(username=username).first()
@app.route("/get_chat_history")
def get_chat_history():
    session_id = request.args.get("session_id")
    user = get_current_user()

    if not user or not session_id:
        return jsonify({"error": "Unauthorized"}), 401

    # Trả toàn bộ lịch sử theo đúng thứ tự
    all_messages = ChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=session_id
    ).order_by(ChatHistory.timestamp.asc()).all()

    result = [
    {
        "role": m.role,
        "content": m.content,
        "image_url": m.image_url,
        "reply_to": m.reply_to  # ✅ THÊM DÒNG NÀY
    }
    for m in all_messages
]


    return jsonify({
        "history": result,
        "too_long": len(all_messages) > 50  # Optional: cảnh báo nếu dài
    })
MAX_CHAT_SESSIONS = 70


@app.route("/save_chat", methods=["POST"])
def save_chat():
    user_id = session.get("user_id") or session.get("anonymous_id")
    session_id = session.get("chat_session_id")
    data = request.get_json()
    custom_title = data.get("title", "").strip()

    if not user_id or not session_id:
        return jsonify({"success": False, "error": "Không tìm thấy phiên chat"})

    # ⚠️ Giới hạn 70 đoạn chat đã lưu
    MAX_SAVED_CHATS = 70
    saved_count = SavedChat.query.filter_by(user_id=user_id).count()
    if saved_count >= MAX_SAVED_CHATS:
        return jsonify({
            "success": False,
            "error": f"Bạn chỉ được lưu tối đa {MAX_SAVED_CHATS} đoạn chat. Vui lòng xóa bớt để lưu mới."
        })

    # 📥 Lấy lịch sử chat hiện tại
    history = ChatHistory.query.filter_by(
        user_id=user_id,
        session_id=session_id
    ).order_by(ChatHistory.timestamp.asc()).all()

    if not history:
        return jsonify({"success": False, "error": "Không có nội dung để lưu"})

    # 🔄 Chuyển thành format JSON
    messages = []
    for m in history:
        raw = m.image_url
        parsed_urls = []

        if raw:
            try:
                parsed_urls = json.loads(raw)
                if not isinstance(parsed_urls, list):
                    parsed_urls = [parsed_urls]
            except:
                parsed_urls = [raw.strip().strip('"').lstrip("/")]

        messages.append({
            "role": m.role,
            "content": m.content,
            "image_url": parsed_urls
        })

    # 💾 Lưu vào bảng SavedChat
    saved = SavedChat(
        id=str(uuid.uuid4()),
        user_id=user_id,
        session_id=session_id,
        title=custom_title or f"Đoạn chat lúc {datetime.now().strftime('%H:%M %d/%m/%Y')}",
        messages=messages
    )
    db.session.add(saved)

    # ✅ Đánh dấu đoạn chat này là đã lưu
    for m in history:
        m.is_saved = True

    db.session.commit()

    return jsonify({"success": True})

@app.route("/saved_chats", methods=["GET"])
def get_saved_chats():
    user_id = session.get("user_id") or session.get("anonymous_id")
    if not user_id:
        return jsonify({"success": False, "data": []})

    # ❌ Bỏ giới hạn limit(10) để trả về toàn bộ
    saved = SavedChat.query.filter_by(user_id=user_id)\
        .order_by(SavedChat.created_at.desc())\
        .all()

    return jsonify({
        "success": True,
        "data": [
            {
                "id": chat.id,
                "title": chat.title,
                "created_at": chat.created_at.strftime("%d/%m/%Y"),
            }
            for chat in saved
        ]
    })

def strip_dangerous_tags(html):
    if not html:
        return ""
    # Escape các thẻ nguy hiểm
    tags = ["button", "input", "svg", "form", "style", "script", "iframe", "object", "embed", "link"]
    for tag in tags:
        html = re.sub(f"<{tag}([^>]*)>", r"&lt;" + tag + r"\1&gt;", html, flags=re.IGNORECASE)
        html = re.sub(f"</{tag}>", r"&lt;/" + tag + r"&gt;", html, flags=re.IGNORECASE)
    return html
def remove_ai_image_note(html):
    if not html:
        return ""
    # Xoá toàn bộ <svg>...</svg>
    html = re.sub(r"<svg.*?</svg>", "", html, flags=re.DOTALL|re.IGNORECASE)
    # Xoá luôn dòng 'AI vừa tạo ảnh này'
    html = re.sub(r"AI vừa tạo ảnh này", "", html, flags=re.IGNORECASE)
    return html.strip()
@app.route("/shared_chat/<chat_id>")
def shared_chat(chat_id):
    chat = db.session.get(SavedChat, chat_id)
    if not chat:
        return abort(404)

    def clean_image_url(urls):
        if not urls:
            return []
        if isinstance(urls, str):
            try:
                urls = json.loads(urls)
                if not isinstance(urls, list):
                    urls = [urls]
            except:
                urls = [urls]
        result = []
        for url in urls:
            if isinstance(url, str) and url.strip() != "":
                url = url.strip().strip('"').lstrip("/")
                result.append(f"/{url}")
        return result

    for msg in chat.messages:
        msg["image_url"] = clean_image_url(msg.get("image_url"))
        msg["content"] = remove_ai_image_note(msg.get("content", ""))
        msg["content"] = strip_dangerous_tags(msg.get("content", ""))


    return render_template("shared_chat.html", chat=chat, user_id=chat.user_id)


@app.route("/saved_chat/<chat_id>")
def view_saved_chat(chat_id):
    user_id = session.get("user_id") or session.get("anonymous_id")
    saved = SavedChat.query.filter_by(id=chat_id, user_id=user_id).first()
    if not saved:
        return jsonify({"success": False, "error": "Không tìm thấy đoạn chat."})

    def clean_image_url(urls):
        if not urls:
            return []
        if isinstance(urls, str):
            try:
                urls = json.loads(urls)
                if not isinstance(urls, list):
                    urls = [urls]
            except:
                urls = [urls]
        result = []
        for url in urls:
            if isinstance(url, str) and url.strip() != "":
                url = url.strip().strip('"').lstrip("/")
                result.append(f"/{url}")
        return result

    cleaned_messages = []
    for msg in saved.messages:
        msg["image_url"] = clean_image_url(msg.get("image_url"))
        cleaned_messages.append(msg)

    return jsonify({
        "success": True,
        "title": saved.title,
        "created_at": saved.created_at.strftime("%Y-%m-%d"),
        "messages": cleaned_messages
    })

@app.route("/delete_saved_chat/<chat_id>", methods=["DELETE"])
def delete_saved_chat(chat_id):
    try:
        user_id = session.get("user_id")
        if not user_id:
            return jsonify({"success": False, "error": "Không có user_id trong session"})

        # 🔍 Tìm đoạn chat đã lưu
        saved = SavedChat.query.filter_by(id=chat_id, user_id=user_id).first()
        if not saved:
            return jsonify({"success": False, "error": "Không tìm thấy đoạn chat đã lưu"})

        # 🧹 B1: Xóa hẳn các tin nhắn trong ChatHistory thuộc session này
        ChatHistory.query.filter_by(user_id=user_id, session_id=saved.session_id).delete()

        # 🧹 B2: Xóa các ảnh thuộc đoạn chat này
        upload_folder = os.path.abspath(os.path.join("static", "images", "uploads", user_id))
        print("[DEBUG] Thư mục uploads:", upload_folder)

        for msg in saved.messages:
            raw_urls = msg.get("image_url", [])

            if isinstance(raw_urls, str):
                try:
                    raw_urls = json.loads(raw_urls)
                except:
                    raw_urls = [raw_urls]
            elif not isinstance(raw_urls, list):
                raw_urls = [raw_urls]

            for url in raw_urls:
                try:
                    relative_path = url.replace("/static/", "", 1)
                    abs_path = os.path.abspath(os.path.join("static", relative_path))

                    if abs_path.startswith(upload_folder) and os.path.isfile(abs_path):
                        print("[🧹] Đang xóa ảnh:", abs_path)
                        os.remove(abs_path)
                    else:
                        print(f"❌ Không xóa ảnh: {abs_path} (không đúng thư mục hoặc không tồn tại)")
                except Exception as e:
                    print(f"❌ Không thể xóa ảnh '{url}':", e)

        # 🧹 B3: Xóa khỏi bảng SavedChat
        db.session.delete(saved)
        db.session.commit()

        return jsonify({"success": True})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)})



#XÓA AVTR
@app.route("/user-info/delete-avatar", methods=["POST"])
def delete_avatar():
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    if user.avatar_url and user.avatar_url != "/static/logos/logo.png":
        try:
            os.remove(user.avatar_url.replace("/", os.sep)[1:])
        except:
            pass
    user.avatar_url = "/static/logos/logo.png"

    db.session.commit()

    flash("Ảnh đại diện đã được đặt lại về mặc định!", "info")
    return redirect("/user-info")

#========ROUTE XÁC THỰC TÀI KHOẢN==========#
@app.route("/verify-otp", methods=["GET", "POST"])
def verify_otp():
    username = session.get("pending_user") or session.get("username")
    if not username:
        return redirect(url_for("login"))

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect(url_for("login"))

    if user.is_verified:
        return redirect(url_for("home_page"))

    # ✅ GET: chỉ hiển thị form, không gửi lại OTP nữa
    if request.method == "GET":
        return render_template(
            "verify_otp.html",
            username=username,
            method="email",
            error="",
            user=user
        )

    # ✅ POST: xử lý OTP
    if request.method == "POST":
        otp_input = request.form.get("otp")
        if otp_input == user.otp_code:
            user.is_verified = True
            user.otp_code = None
            user.wants_verification = True

            if user.email_temp and user.email_temp != user.email:
                user.email = user.email_temp
                user.email_temp = None

            if (user.free_gpt_uses or 0) <= 5:
                user.free_gpt_uses = (user.free_gpt_uses or 0) + 10

            db.session.commit()

            session["username"] = username
            session["user_id"] = user.user_id
            session["is_verified"] = True
            session.pop("pending_user", None)
            session.pop("otp_sent", None)

            for key in ["chat_history", "chat_ai_lite", "chat_ai_lite_history"]:
                session.pop(key, None)

            session["just_verified"] = True
            session.modified = True

            return redirect(url_for("home_page"))

        return render_template(
            "verify_otp.html",
            username=username,
            method="email",
            error="❌ Sai mã OTP. Vui lòng thử lại.",
            user=user
        )
    return render_template(
        "verify_otp.html",
        username=username,
        method="email",
        error="",
        user=user
    )

@app.route("/resend-otp", methods=["POST"])
def resend_otp():
    username = session.get("pending_user") or session.get("username")
    if not username:
        return Response(json.dumps({
            "status": "error",
            "message": "❌ Không tìm thấy tài khoản."
        }, ensure_ascii=False), content_type="application/json")

    user = User.query.filter_by(username=username).first()
    if not user:
        return Response(json.dumps({
            "status": "error",
            "message": "❌ Không tìm thấy tài khoản."
        }, ensure_ascii=False), content_type="application/json")

    email = user.email
    otp = str(random.randint(100000, 999999))
    user.otp_code = otp
    db.session.commit()
    success = send_user_otp_email(email, otp)

    if success:
        return Response(json.dumps({
            "status": "ok",
            "message": "✅ Đã gửi lại mã xác thực qua email."
        }, ensure_ascii=False), content_type="application/json")
    else:
        return Response(json.dumps({
            "status": "error",
            "message": "❌ Gửi email thất bại. Vui lòng thử lại sau."
        }, ensure_ascii=False), content_type="application/json")

@app.route("/logout")
def logout():
    username = session.get("username")

    if username:
        user = User.query.filter_by(username=username).first()
        if user:
            user.online = False
            db.session.commit()

    if session.get("admin"):
        session.pop("admin", None)
        return redirect("/admin_login")

    session.pop("username", None)
    session.pop("user_id", None)
    return redirect("/login")

#KHÔI PHỤC TÀI KHOẢN 
def generate_forgot_pass_slug():
    return str(uuid.uuid4())
@app.route("/forgot-password", methods=["GET"])
def redirect_forgot_password():
    slug = generate_forgot_pass_slug()
    return redirect(f"/forgotPassBridge/{slug}")
@app.route("/forgotPassBridge/<slug>", methods=["GET", "POST"])
def forgot_password(slug):
    if request.method == "POST":
        input_text = request.form["username"].strip().lower()

        # ✅ Tìm user theo username hoặc email
        user = User.query.filter(
            (User.username.ilike(input_text)) | (User.email.ilike(input_text))
        ).first()

        if not user:
            return render_template("forgot_password.html", error="Tài khoản không tồn tại.")

        if not user.is_verified:
            return render_template("forgot_password.html", error="Chức năng này chỉ hỗ trợ tài khoản đã xác thực.")

        # ✅ Tạo và gửi mã OTP
        otp_code = generate_otp()
        send_user_otp_email(user.email, otp_code)

        # ✅ Ghi trực tiếp vào user.otp_code
        user.otp_code = otp_code
        db.session.commit()

        # ✅ Lưu session
        session["reset_user"] = user.username
        return redirect("/reset-password")

    return render_template("forgot_password.html")

@app.route("/change-email", methods=["GET", "POST"])
def change_email():
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    avatar_filename = user.avatar_filename or "default.png"
    avatar_url = url_for("static", filename=f"avatars/{avatar_filename}")

    if not user.is_verified and not user.can_change_email_once:
        flash("Tài khoản của bạn chưa xác thực và đã đổi email 1 lần. Vui lòng liên hệ hỗ trợ.", "danger")
        return redirect("/user-info")

    if request.method == "POST":
        new_email = request.form.get("new_email", "").strip()

        if not new_email or "@" not in new_email:
            flash("Email không hợp lệ!", "danger")
            return redirect("/change-email")

        user.email_temp = new_email
        user.is_verified = False
        user.otp_code = generate_otp()

        db.session.commit()

        send_otp_email(new_email, user.otp_code)

        flash("Đã gửi mã xác nhận đến email mới! Vui lòng kiểm tra và xác thực lại.", "info")
        return redirect("/verify-otp")

    return render_template("change_email.html", avatar_url=avatar_url)

@app.route("/reset-password", methods=["GET", "POST"])
def reset_password():
    session_username = session.get("reset_user")
    if not session_username:
        return redirect("/forgot-password")

    session_username = session_username.strip().lower()

    if request.method == "POST":
        input_username = request.form["username"].strip().lower()
        otp = request.form["otp"].strip()
        new_password = request.form["new_password"]
        confirm = request.form["confirm_password"]

        user = User.query.filter_by(username=session_username).first()
        avatar_url = user.avatar_url if user and user.avatar_url else "/static/logos/logo.png"

        if input_username != session_username:
            return render_template("reset_password.html", error="⚠️ Tên tài khoản không khớp.", avatar_url=avatar_url)

        if not user or user.otp_code != otp:
            return render_template("reset_password.html", error="⚠️ Mã OTP không đúng.", avatar_url=avatar_url)

        if new_password != confirm:
            return render_template("reset_password.html", error="⚠️ Mật khẩu xác nhận không khớp.", avatar_url=avatar_url)

        if len(new_password) < 8:
            return render_template("reset_password.html", error="⚠️ Mật khẩu phải từ 8 ký tự trở lên.", avatar_url=avatar_url)

        # ✅ Cập nhật mật khẩu và xoá OTP
        user.password = new_password
        user.otp_code = None
        user.online = True
        db.session.commit()

        session.pop("reset_user", None)

        # ✅ Tự đăng nhập lại
        session["username"] = user.username
        session["user_id"] = user.user_id
        session["vip_until_gpt"] = user.vip_until_gpt
        session["vip_ai_lite"] = user.vip_ai_lite
        session["vip_until_lite"] = user.vip_until_lite or ""
        session["al_uses"] = 0

        return redirect("/")

    # ✅ Trường hợp GET: chỉ tìm user để lấy avatar
    user = User.query.filter_by(username=session_username).first()
    avatar_url = user.avatar_url if user and user.avatar_url else "/static/logos/logo.png"
    return render_template("reset_password.html", avatar_url=avatar_url)

def generate_change_pass_slug():
    return str(uuid.uuid4())
@app.route("/change-password", methods=["GET"])
def redirect_change_password():
    slug = generate_change_pass_slug()
    return redirect(f"/changePassX7vault/{slug}")
@app.route("/changePassX7vault/<slug>", methods=["GET", "POST"])
def change_password(slug):
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")
    avatar_url = user.avatar_url or "/static/logos/logo.png"
    if request.method == "POST":
     
        current_password = request.form.get("current_password")
        new_password = request.form.get("new_password")
        confirm_password = request.form.get("confirm_password")

        if user.password != current_password:
            return render_template("change_password.html", error="⚠️ Mật khẩu hiện tại không đúng.")

        if new_password != confirm_password:
            return render_template("change_password.html", error="⚠️ Mật khẩu xác nhận không khớp.")

        if len(new_password) < 8:
            return render_template("change_password.html", error="⚠️ Mật khẩu phải từ 8 ký tự trở lên.")

        user.password = new_password
        db.session.commit()

        return redirect("/")

    return render_template("change_password.html", avatar_url=avatar_url)



# TRANG CHỦ
from sqlalchemy import distinct
from models import Message, User  # đảm bảo đã import model Message

@app.route("/")
def home_page():
    if is_maintenance("home"):
        return render_template("maintenance.html")

    username = session.get("username")

    # Nếu chưa login, vẫn cho hiển thị trang chủ (home), nhưng dưới dạng "khách"
    if not username:
        return render_template("app.html", user=None)

    user = User.query.filter_by(username=username).first()
    if not user:
        if not session.get("reloading_after_restore"):
            session["reloading_after_restore"] = True
            time.sleep(0.3)
            return redirect("/")
        else:
            session.pop("reloading_after_restore", None)
            return render_template("login.html", error="⚠️ Dữ liệu tài khoản không tồn tại. Vui lòng đăng nhập lại.")

    if user.is_blocked:
        session.clear()
        return render_template("login.html", error="🚫 Tài khoản của bạn đã bị khóa. Nếu đây là nhầm lẫn, vui lòng gửi khiếu nại tại <a href='/appeal' style='color:#4ea6ff;'>đây</a>.")

    session.pop("reloading_after_restore", None)
    session.pop("just_logged_in", None)

    return render_template("app.html", user=user)



#TỔNG QUÁT
from flask import render_template, request, session, redirect, abort
login_attempts = {}

@app.route("/solverviet_control_x2025")
@admin_only
def admin_panel():
    feedback_count = count_feedback_entries()
    return render_template("admin.html",feedback_count=feedback_count)
@app.route("/admin/check_tu", methods=["POST"])
@admin_only
def admin_check_tu_trong_tu_dien():
    data = request.get_json()
    tu = data.get("tu", "").strip()

    def normalize(text):
        return unicodedata.normalize("NFC", text.lower().strip())

    try:
        with open("viet_dictionary.json", "r", encoding="utf-8") as f:
            dic = json.load(f)

        normalized_keys = {normalize(k): k for k in dic.keys()}
        tu_norm = normalize(tu)
        exists = tu_norm in normalized_keys

        return jsonify({"exists": exists})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def count_feedback_entries():
    try:
        with open("feedback_log.txt", "r", encoding="utf-8") as f:
            blocks = f.read().split("---\n\n")
            return sum(1 for block in blocks if block.strip())
    except FileNotFoundError:
        return 0
import pyotp
from flask import Flask, request, session, redirect, render_template

# ---- Route 1: admin_login - kiểm tra 4 trường ----
@app.route("/admin_login", methods=["GET", "POST"])
def admin_login():
    ip = request.remote_addr
    login_attempts.setdefault(ip, 0)

    if login_attempts[ip] >= 5:
        return "⚠️ You have entered incorrectly too many times. Please try again later.", 429

    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        code1 = request.form.get("backdoor_code", "")
        code2 = request.form.get("backdoor_code2", "")

        ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "")
        ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")
        BACKDOOR_CODE = os.getenv("BACKDOOR_CODE", "")
        BACKDOOR_CODE2 = os.getenv("BACKDOOR_CODE2", "")

        if (username == ADMIN_USERNAME and password == ADMIN_PASSWORD and
            code1 == BACKDOOR_CODE and code2 == BACKDOOR_CODE2):
            session["admin_otp_ready"] = True
            session["admin_last_active"] = time.time()
            return "OK"

        else:
            login_attempts[ip] += 1
            return "❌ Invalid credentials.", 403

    return render_template("admin_login.html")

@app.route("/admin_prepare_secret")
def admin_prepare_secret():
    hint_map = {
        "ADMIN_SECRET_CODE_1": "🔺",
        "ADMIN_SECRET_CODE_2": "🌙",
        "ADMIN_SECRET_CODE_3": "💡",
        "ADMIN_SECRET_CODE_4": "🔒",
        "ADMIN_SECRET_CODE_5": "🧠",
    }

    pool = []
    for key, hint in hint_map.items():
        val = os.getenv(key)
        if val:
            pool.append({"code": val, "hint": hint})

    if not pool:
        return jsonify({"error": "No codes found"}), 500

    selected = random.choice(pool)
    session["expected_secret_code"] = selected["code"]
    return jsonify(hint=selected["hint"])

@app.route("/admin_secret_entry", methods=["POST"])
def admin_secret_entry():
    data = request.get_json()
    code_input = data.get("code", "").strip()
    expected = session.get("expected_secret_code", "")

    if code_input == expected:
        session["is_admin"] = True
        session["admin_last_active"] = time.time()
        return jsonify(success=True)
    return jsonify(success=False)

@app.route("/admin_2fa", methods=["POST"])
def admin_2fa():
    if not session.get("admin_otp_ready"):
        return "Phiên xác thực không hợp lệ.", 403

    code = request.form.get("otp", "").strip()
    secret = os.getenv("ADMIN_2FA_SECRET")

    if not secret:
        return "❌ SECRET không tồn tại", 500

    totp = pyotp.TOTP(secret)

    if totp.verify(code, valid_window=1):
        session.pop("admin_otp_ready", None)
        session["is_admin"] = True
        return redirect("/solverviet_control_x2025")
    else:
        return "❌ Mã xác thực sai", 403

@app.route("/verify_user_info", methods=["POST"])
def verify_user_info():
    data = request.get_json()
    username = data.get("username", "").strip()
    user_id = data.get("user_id", "").strip()
    user = User.query.filter_by(username=username, user_id=user_id).first()
    return jsonify({"valid": bool(user)})

def send_upgrade_email(to_email, username, package, amount, method, note, created_at, txn_id):
    from flask import url_for
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    email_user = os.getenv("EMAIL_ADDRESS")
    email_pass = os.getenv("EMAIL_PASSWORD")
    BASE_URL = os.getenv("BASE_URL", "http://127.0.0.1:5000")

    approve_path = url_for('admin_review') + f"?txn_id={txn_id}"
    approve_link = "https://solverviet.com/admin/review"
    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = to_email
    msg['Subject'] = " Yêu cầu nâng cấp gói mới từ người dùng"

    body = f"""
    <h3> Yêu cầu nâng cấp mới</h3>
    <p><b> User:</b> {username}</p>
    <p><b> Gói:</b> {package} ({amount})</p>
    <p><b> Phương thức:</b> {method}</p>
    <p><b> Ghi chú:</b> {note}</p>
    <p><b> Thời gian:</b> {created_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
    <br>
    <p><b>Link duyệt:</b> <a href="{approve_link}" target="_blank">{approve_link}</a></p>
    <p>Hoặc copy thủ công: {approve_link}</p>
    """

    msg.attach(MIMEText(body, 'html'))

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(email_user, email_pass)
            server.send_message(msg)
    except Exception:
        pass

@app.route("/upgradeWaguri_9d7s2x4kP1tY0mVn6cQ8hR5eB3aZxLw_fJ7k", methods=["GET", "POST"])
def upgrade():
    username = session.get("username")
    if not username:
        return redirect(url_for("login"))

    user = User.query.filter_by(username=username).first()
    if not user:
        return render_template("upgrade.html", error="❌ Không tìm thấy tài khoản.")

    if request.method == "POST":
        package = request.form["package"]
        note = request.form["note"]
        referral_code_input = request.form.get("referral_code", "").strip().upper()

        # ✅ Ưu tiên mã đã lưu, nếu chưa có thì dùng mã vừa nhập
        referral_code = user.referral_code_used or referral_code_input

        # ✅ Nếu user chưa từng nhập mã → lưu lại vĩnh viễn
        if referral_code and not user.referral_code_used:
            user.referral_code_used = referral_code

        # ✅ Gắn mã vào ghi chú để admin dễ nhìn
        if referral_code:
            note += f" | Mã giới thiệu: {referral_code}"

        method = request.form.get("method", "Không rõ")

        if not note.strip():
            return render_template("upgrade.html", error="⚠️ Vui lòng nhập mã giao dịch hoặc ghi chú.")

        now = now_vn()
        has_gpt = user.vip_until_gpt and user.vip_until_gpt > now
        has_lite = user.vip_until_lite and user.vip_until_lite > now

        if package.startswith("vip_gpt") and has_lite:
            return render_template("upgrade.html", error="🚫 Bạn đang dùng gói AI Lite. Không thể mua thêm gói GPT cùng lúc.")
        if package == "vip_ai_lite" and has_gpt:
            return render_template("upgrade.html", error="🚫 Bạn đang dùng gói GPT. Không thể mua thêm gói AI Lite cùng lúc.")

        # ✅ Tạo ID giao dịch
        created_at = now
        txn_id = f"txn_{int(now.timestamp())}_{random.randint(1000,9999)}"
        session["last_txn_id"] = txn_id

        package_display = {
            "vip_gpt_5d": "49K",
            "vip_gpt_15d": "109K",
            "vip_gpt_30d": "149K",
            "vip_ai_lite": "25K"
        }
        amount = package_display.get(package, "Không rõ")

        # ✅ Lưu giao dịch
        txn = Transaction(
            txn_id=txn_id,
            username=username,
            amount=amount,
            package=package,
            method=method,
            note=note,
            referral_code=referral_code,
            status="pending",
            created_at=created_at
        )
        db.session.add(txn)
        db.session.commit()  # ⚠️ commit cả giao dịch và cập nhật mã giới thiệu

        # ✅ Gửi email cho admin
        try:
            send_upgrade_email(
                to_email=os.getenv("ADMIN_EMAIL"),
                username=username,
                package=package,
                amount=amount,
                method=method,
                note=note,
                created_at=created_at,
                txn_id=txn_id 
            )
        except Exception as e:
            pass
        try:
            send_telegram_message(
                f"🔔 <b>Yêu cầu nâng cấp mới</b>\n"
                f"👤 User: <code>{username}</code>\n"
                f"💳 Gói: {package} ({amount})\n"
                f"🏦 Phương thức: {method}\n"
                f"📝 Ghi chú: <code>{note}</code>\n"
                f"🕒 Thời gian: {created_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"➡️ Vào trang admin để duyệt: {url_for('admin_review', _external=True)}"
            )
        except:
            pass

        return render_template(
            "upgrade.html",
            success=True,
            txn_id=txn_id,
            flash_message="✅ Giao dịch đã được ghi nhận! Bạn sẽ được duyệt trong vòng 5 phút."
        )
    display_name = user.fullname if user and user.fullname else username
    now = now_vn()
    vip_gpt_active = user.vip_until_gpt and user.vip_until_gpt > now
    vip_lite_active = user.vip_until_lite and user.vip_until_lite > now

    return render_template(
        "upgrade.html",
        fullname=display_name,
        username=username,
        user_id=user.user_id if user else "",
        vip_gpt_active=vip_gpt_active,
        vip_lite_active=vip_lite_active,
        user=user
    )
from flask import jsonify
import logging
log = logging.getLogger('werkzeug')

@app.route("/check_status")
def check_status():
    log.setLevel(logging.ERROR) 
    txn_id = session.get("last_txn_id")
    if not txn_id:
        return jsonify({"status": "none"})
    from models.transaction import Transaction
    txn = Transaction.query.filter_by(txn_id=txn_id).first()
    if txn:
        return jsonify({
            "status": txn.status,
            "package": txn.package,
            "method": txn.method,
            "created_at": txn.created_at.strftime("%Y-%m-%d %H:%M:%S")
        })
    else:
        return jsonify({"status": "not_found"})

def grant_vip(username, package, method=""):
    user = User.query.filter_by(username=username).first()
    if not user:
        return " Người dùng không tồn tại."

    now = datetime.utcnow()
    message = " Gói đã được cấp thành công!"

    gpt_vip_until = user.vip_until_gpt or now
    lite_vip_until = user.vip_until_lite or now

    if package in ["vip_gpt_5d", "vip_gpt_15d", "vip_gpt_30d"]:
        if lite_vip_until > now:
            return "Bạn đang có gói Lite. Không thể mua gói SLV. Vui lòng đợi hết hạn hoặc liên hệ nhà phát triển."
        days_map = {
            "vip_gpt_5d": (5, "5day"),
            "vip_gpt_15d": (15, "15day"),
            "vip_gpt_30d": (30, "30day")
        }

        duration, gpt_type = days_map[package]

        user.vip_gpt_ai = True
        user.vip_gpt = gpt_type
        new_vip = (now + timedelta(days=duration)).replace(hour=23, minute=59, second=59, microsecond=0)
        user.vip_until_gpt = new_vip

        if gpt_type == "5day":
            user.gpt_usage_today = 0
            user.gpt_usage_date = now.strftime("%Y-%m-%d")
            user.gpt_unlimited = False

        user.vip_until_lite = None
        user.vip_ai_lite = None
        user.vip_lite_daily_limit = None
        user.lite_usage = None
        user.lite_date = None
    elif package == "vip_ai_lite":
        if gpt_vip_until > now:
            return "Bạn đang có gói SLV. Không thể mua gói Lite. Vui lòng đợi hết hạn hoặc liên hệ nhà phát triển."

        # ✅ Luôn cấp gói Lite mới nếu hợp lệ
        new_vip = (now + timedelta(days=7)).replace(hour=23, minute=59, second=59, microsecond=0)
        user.vip_ai_lite = True
        user.vip_until_lite = new_vip
        user.vip_lite_daily_limit = 50
        user.lite_usage = 0
        user.lite_date = now.strftime("%Y-%m-%d")
    else:
        return "Gói không hợp lệ."

    try:
        db.session.commit()
        return message
    except Exception as e:
        db.session.rollback()
        return " Lỗi khi lưu dữ liệu."


def review_authenticated_only(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('review_authenticated'):
            return redirect('/admin/review/auth')
        return f(*args, **kwargs)
    return decorated_function
import os, smtplib, time
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import pyotp
from flask import session, redirect, request, render_template, abort, url_for
EMAIL_USER = os.getenv("EMAIL_ADDRESS")      
EMAIL_PASS = os.getenv("EMAIL_PASSWORD")   
def send_payment_success_email(to_email: str, username: str, user_id: str, plan_name: str,
                               amount: str, txn_id: str,
                               effective_at: datetime, expire_at: datetime):
    if not (EMAIL_USER and EMAIL_PASS):
        raise RuntimeError("Missing EMAIL_ADDRESS/EMAIL_PASSWORD")

    subject = f"Thanh toán thành công • {plan_name} • SolverViet"
    body = f"""
    <div style="font-family:Segoe UI,Arial,sans-serif;color:#111;line-height:1.6">
      <h2 style="margin:0 0 8px;color:#0ea5e9;">SolverViet</h2>
      <p><b>Xin chào {username}</b> (ID: {user_id}),</p>
      <p>Giao dịch của bạn đã được duyệt thành công.</p>

      <div style="border:1px solid #eee;border-radius:10px;padding:14px;margin:12px 0;background:#fafafa">
        <p style="margin:6px 0;"><b>Gói:</b> {plan_name} ({amount})</p>
        <p style="margin:6px 0;"><b>Mã giao dịch:</b> {txn_id}</p>
        <p style="margin:6px 0;"><b>Ngày hiệu lực:</b> {effective_at.strftime("%d/%m/%Y")}</p>
        <p style="margin:6px 0;"><b>Ngày hết hạn:</b> {expire_at.strftime("%d/%m/%Y")}</p>
      </div>

      <p>Cảm ơn bạn đã sử dụng dịch vụ của SolverViet! 💙</p>
      <hr style="margin:20px 0;">

      <h3 style="color:#0ea5e9;">English Version</h3>
      <p><b>Hello {username}</b> (ID: {user_id}),</p>
      <p>Your transaction has been successfully approved.</p>

      <div style="border:1px solid #eee;border-radius:10px;padding:14px;margin:12px 0;background:#fafafa">
        <p style="margin:6px 0;"><b>Plan:</b> {plan_name} ({amount})</p>
        <p style="margin:6px 0;"><b>Transaction ID:</b> {txn_id}</p>
        <p style="margin:6px 0;"><b>Effective Date:</b> {effective_at.strftime("%d/%m/%Y")}</p>
        <p style="margin:6px 0;"><b>Expiry Date:</b> {expire_at.strftime("%d/%m/%Y")}</p>
      </div>

      <p>Thank you for choosing SolverViet! 💙</p>
    </div>
    """

    msg = MIMEMultipart()
    msg["From"] = EMAIL_USER
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "html"))

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(EMAIL_USER, EMAIL_PASS)
        server.send_message(msg)
import os
import pyotp
import time
from flask import session, redirect, request, render_template
@app.route("/admin/review", methods=["GET", "POST"])
def admin_review():
    if request.method == "POST":
        code = request.form.get("code")
        secret = os.getenv("ADMIN_2FA_SECRET")
        if not secret:
            return render_template("login.html", error="System validation failed. Contact administrator for access.")
        totp = pyotp.TOTP(secret)
        if totp.verify(code):
            session['review_authenticated'] = True
            session['review_last_active'] = time.time()
        else:
            abort(403)

    if (not session.get("review_authenticated")
        or time.time() - session.get("review_last_active", 0) > 300):
        session.pop("review_authenticated", None)
        return render_template("admin_review.html", txs=[])

    session['review_last_active'] = time.time()
    txs = Transaction.query.filter_by(status="pending").order_by(Transaction.created_at.desc()).all()
    return render_template("admin_review.html", txs=txs)

@app.route("/admin/review/requests")
@review_authenticated_only
def get_review_requests():
    txs = Transaction.query.filter_by(status="pending").order_by(Transaction.created_at.desc()).all()
    return render_template("admin/_requests_table.html", txs=txs)

@app.route("/admin/approve/<txn_id>")
@review_authenticated_only
def approve_transaction(txn_id):
    tx = Transaction.query.filter_by(txn_id=txn_id, status="pending").first()
    if not tx:
        return redirect(url_for("admin_review"))

    tx.status = "approved"
    db.session.commit()

    package = normalize_package(tx.package)
    method = tx.method
    username = tx.username
    result = grant_vip(username, package, method)
    if isinstance(result, str) and result.startswith("❌"):
        return redirect(url_for("admin_review"))

    user = User.query.filter_by(username=username).first()
    if user and user.email:
        now = now_vn()
        if package == "vip_gpt_5d":
            days, plan_name = 5, "SLV 5 ngày"
        elif package == "vip_gpt_15d":
            days, plan_name = 15, "SLV 15 ngày"
        elif package == "vip_gpt_30d":
            days, plan_name = 30, "SLV 30 ngày"
        elif package == "vip_ai_lite":
            days, plan_name = 7, "AI Lite 7 ngày"
        else:
            days, plan_name = 0, package

        effective_at = now
        expire_at = (now + timedelta(days=days)).replace(hour=23, minute=59, second=59, microsecond=0)
        try:
            send_payment_success_email(
                to_email=user.email,
                username=username,
                user_id=user.user_id,  
                plan_name=plan_name,
                amount=tx.amount,
                txn_id=tx.txn_id,
                effective_at=effective_at,
                expire_at=expire_at
            )
        except Exception:
            pass

    return redirect(url_for("admin_review"))

@app.route("/admin/reject/<txn_id>")
@review_authenticated_only
def reject_transaction(txn_id):
    tx = Transaction.query.filter_by(txn_id=txn_id, status="pending").first()
    if not tx:
        return redirect(url_for("admin_review"))

    tx.status = "rejected"
    db.session.commit()
    return redirect(url_for("admin_review"))

def is_follow_up(msg):
    msg = msg.lower().strip()
    follow_keywords = ["bài", "phần", "tiếp", "tiếp theo", "tiếp tục", "rồi sao", "câu", "nữa", "b nhé", "c thì sao"]
    return any(kw in msg for kw in follow_keywords) and len(msg.split()) <= 6

#=====BẠN BÈ=====#
@app.route("/unblock_user", methods=["POST"])
def unblock_user():
    username = session.get("username")
    data = request.get_json()
    target = data.get("target")

    if not username or not target:
        return jsonify({"status": "error", "message": "Thiếu dữ liệu"})

    user = User.query.filter_by(username=username).first()
    target_user = User.query.filter_by(username=target).first()

    if not user or not target_user:
        return jsonify({"status": "error", "message": "Không tìm thấy user"})

    if target_user in user.blocked_users:
        user.blocked_users.remove(target_user)
        db.session.commit()
        return jsonify({"status": "success"})

    return jsonify({"status": "error", "message": "User chưa bị chặn"})

@app.route("/blocked_users")
def get_blocked_users():
    username = session.get("username")
    if not username:
        return jsonify([])

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify([])

    blocked_list = [u.username for u in user.blocked_users]
    return jsonify(blocked_list)


@app.route("/block_user", methods=["POST"])
def block_user():
    username = session.get("username")
    data = request.get_json()
    target = data.get("target")

    if not username or not target:
        return jsonify({"status": "error", "message": "Thiếu dữ liệu"})

    user = User.query.filter_by(username=username).first()
    target_user = User.query.filter_by(username=target).first()

    if not user or not target_user:
        return jsonify({"status": "error", "message": "User không tồn tại"})

    if target_user not in user.blocked_users:
        user.blocked_users.append(target_user)
        db.session.commit()

    return jsonify({"status": "success"})


@app.route("/chat/delete/<friend>", methods=["POST"])
def delete_chat(friend):
    username = session.get("username")
    if not username:
        return jsonify({"success": False, "error": "Chưa đăng nhập"})

    from models.message import Message
    data = request.get_json()
    mode = data.get("mode", "one_side")

    if mode == "one_side":
        msgs = Message.query.filter(
            db.or_(
                db.and_(Message.sender == username, Message.receiver == friend),
                db.and_(Message.sender == friend, Message.receiver == username)
            )
        ).all()
        count_marked = 0
        for msg in msgs:
            deleted = msg.deleted_for or []
            if username not in deleted:
                deleted.append(username)
                msg.deleted_for = deleted
                count_marked += 1
        db.session.commit()
        return jsonify({"success": True})

    # ------------------- BOTH SIDE -----------------------
    msgs = Message.query.filter(
        db.or_(
            db.and_(Message.sender == username, Message.receiver == friend),
            db.and_(Message.sender == friend, Message.receiver == username)
        )
    ).all()
    if not msgs:
        return jsonify({"success": False, "error": "Không tìm thấy đoạn chat."})

    deleted_images = 0
    deleted_voices = 0

    for msg in msgs:
        if msg.image_urls:
            for url in msg.image_urls:
                try:
                    image_path = os.path.join(app.root_path, url.strip("/"))
                    if os.path.exists(image_path):
                        os.remove(image_path)
                        deleted_images += 1
                except Exception:
                    pass

        if msg.voice_url:
            try:
                voice_path = os.path.join(app.root_path, msg.voice_url.strip("/"))
                if os.path.exists(voice_path):
                    os.remove(voice_path)
                    deleted_voices += 1
            except Exception:
                pass

        db.session.delete(msg)

    db.session.commit()

    socketio.emit("chat_deleted_notice", {
        "from": username,
        "to": friend
    }, room=f"user_{friend}")

    return jsonify({"success": True})

@app.route("/update_privacy", methods=["POST"])
def update_privacy():
    username = session.get("username")
    if not username:
        return jsonify({"success": False, "error": "Chưa đăng nhập."})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "error": "Không tìm thấy tài khoản."})

    data = request.get_json()
    user.privacy = {
        "hide_online": data.get("hide_online", False),
        "hide_avatar": data.get("hide_avatar", False),
        "hide_bio":    data.get("hide_bio", False),     
        "hide_all":    data.get("hide_all", False),
        "blocked_users": data.get("blocked_users", []),
    }

    db.session.commit()
    return jsonify({"success": True})


@app.route("/get_profile")
def get_profile():
    username = session.get("username")
    if not username:
        return jsonify({})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({})

    return jsonify({
        "fullname": user.fullname or "",
        "bio": user.bio or "",
        "privacy": user.privacy or {}
    })

@app.route("/update_profile", methods=["POST"])
def update_profile():
    data = request.json
    username = session.get("username")
    if not username:
        return jsonify({"success": False, "message": "Chưa đăng nhập"})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "message": "Không tìm thấy user"})

    user.fullname = data.get("fullname", "")
    user.bio = data.get("bio", "")
    user.bio_updated_at = datetime.utcnow()

    db.session.commit()

    return jsonify({"success": True})

def is_friend(user_id_1, user_id_2):
    return Friend.query.filter(
        or_(
            and_(Friend.user_id == user_id_1, Friend.friend_id == user_id_2),
            and_(Friend.user_id == user_id_2, Friend.friend_id == user_id_1)
        )
    ).first() is not None

@app.route('/users/search')
def users_search():
    query = request.args.get('q', '').strip().lower()
    current_username = session.get("username", "").lower()

    if not query:
        return jsonify([])

    current_user = User.query.filter_by(username=current_username).first()
    if not current_user:
        return jsonify([])

    all_users = (
        User.query
        .filter(func.lower(User.username) != current_username)
        .filter(
            or_(
                func.lower(User.username).like(f"%{query}%"),
                func.lower(User.fullname).like(f"%{query}%"),
                func.lower(User.user_id).like(f"%{query}%"),
            )
        )
        .all()
    )

    visible_users = []
    for u in all_users:
        hide_all = u.privacy.get("hide_all") if u.privacy else False

        if hide_all and not is_friend(current_user.user_id, u.user_id):
            continue  # 🔥 Bỏ qua user đang ẩn danh nếu chưa là bạn

        visible_users.append(u)

    return jsonify([
    {
        'username': u.username,
        'name': u.fullname or u.username,
        'fullname': u.fullname,
        'user_id': u.user_id,
        'online': u.online,
        'is_friend': is_friend(current_user.user_id, u.user_id),
        'last_seen': u.last_seen.strftime("%Y-%m-%d %H:%M:%S") if u.last_seen else "",
        'avatar_url': u.avatar_url or '/static/logos/logo.png'
    }
            for u in visible_users
        ])

def auto_offline():
    timeout = datetime.utcnow() - timedelta(minutes=5)
    users_to_update = User.query.filter(User.last_seen < timeout, User.online == True).all()

    for user in users_to_update:
        user.online = False

    db.session.commit()

@app.before_request
def update_last_seen_and_protect_admin():
    ua = request.headers.get("User-Agent", "")
    path = request.path

    # ✅ Nếu là các route duyệt (review/approve/reject)
    if path.startswith("/admin/review") or path.startswith("/admin/approve") or path.startswith("/admin/reject"):
        if session.get("review_authenticated"):
            return  # Cho phép nếu đã xác thực
        else:
            return  # Cho phép truy cập giao diện để nhập mã xác thực

    # 🚫 Các route admin khác yêu cầu Electron hoặc admin login hoặc key app thật
    if path.startswith("/admin") or path.startswith("/solverviet_control"):
        admin_key = request.headers.get("X-Admin-App-Key", "")
        is_valid_electron = "Electron" in ua
        is_valid_key = admin_key == "SLV_ADMIN_2025"
        is_admin_login = session.get("is_admin")

        if not (is_valid_electron or is_valid_key or is_admin_login):
            return "⛔ Access denied. Internal use only.", 403

        # ⏱ Kiểm tra timeout session admin
        if is_admin_login:
            last_active = session.get("admin_last_active", 0)
            now = time.time()
            if now - last_active > 1800:
                session.clear()
                return redirect("/admin_login")
            else:
                session["admin_last_active"] = now

    # 👤 Cập nhật trạng thái người dùng thường
    username = session.get("username")
    if username:
        try:
            user = User.query.filter_by(username=username).first()
            if user:
                user.online = True
                user.last_seen = datetime.utcnow()
                db.session.commit()
        except Exception as e:
            pass

@app.route("/friends")
def friends_page():
    current_username = session.get("username")
    if not current_username:
        return redirect("/login")
    
    current_user = User.query.filter_by(username=current_username).first()
    if not current_user:
        return redirect("/login")

    auto_offline()
 
    friend_links = Friend.query.filter(
        (Friend.user_id == current_user.user_id) | (Friend.friend_id == current_user.user_id)
    ).all()

    friend_list = []
    seen = set()  

    for link in friend_links:
        friend_id = link.friend_id if link.user_id == current_user.user_id else link.user_id
        friend = User.query.filter_by(user_id=friend_id).first()

        if not friend or friend.username in seen:
            continue
        seen.add(friend.username)

        display_name = friend.fullname or friend.username
        show_bio = not (friend.privacy and friend.privacy.get("hide_bio"))
        online, status_text = get_status(friend.last_seen)
        friend_list.append({
            "username": friend.username,
            "name": friend.fullname,
            "display_name": display_name,
            "fullname": friend.fullname,
            "avatar_url": friend.avatar_url or "/static/logos/logo.png",
            "bio": friend.bio if show_bio else "",
            "last_message": "...",
            "online": online,
            "friend_status": status_text
        })
    
    # Thông tin người dùng hiện tại
    display_name = current_user.fullname or current_user.username
    user_avatar = current_user.avatar_url or url_for('static', filename='logos/logo.png')
    is_online = current_user.online
    user_id = current_user.user_id
    bio = current_user.bio or ""

    show_bio = True 

    pending_invites = FriendRequest.query.filter_by(to_user_id=current_user.user_id).count()

    return render_template("friends.html",
        friends=friend_list,
        username=current_username,
        fullname=display_name,
        user_avatar=user_avatar,
        is_online=is_online,
        bio=bio if show_bio else "",
        current_user=current_user,
        user_id=user_id,
        pending_invites=pending_invites  
    )


def load_messages():
    messages_by_chat = {}

    all_msgs = Message.query.order_by(Message.timestamp).all()
    for msg in all_msgs:
        if msg.chat_key not in messages_by_chat:
            messages_by_chat[msg.chat_key] = []

        # Lọc ảnh mất
        valid_images = [
            img for img in msg.image_urls or []
            if os.path.exists("." + img)
        ]

        messages_by_chat[msg.chat_key].append({
            "sender": msg.sender,
            "content": msg.content,
            "image_urls": valid_images,
            "timestamp": msg.timestamp.strftime("%Y-%m-%d %H:%M:%S")
        })

    return messages_by_chat

def save_messages(data):
    Message.query.delete()

    for chat_key, msgs in data.items():
        for m in msgs:
            sender = m.get("sender", "")
            u1, u2  = chat_key.split("__")
            receiver = u2 if u1 == sender else u1

            vn_time   = datetime.strptime(m["timestamp"], "%Y-%m-%d %H:%M:%S")
            utc_time  = vn_time.replace(tzinfo=VN_TZ).astimezone(UTC_TZ).replace(tzinfo=None)

            new_msg = Message(
                chat_key = chat_key,
                sender   = sender,
                receiver = receiver,
                content  = m.get("content", ""),
                image_urls = m.get("image_urls", []),
                voice_url  = m.get("voice_url"),           
                timestamp  = utc_time
            )
            db.session.add(new_msg)
    db.session.commit()

def get_status(last_seen):
    if not last_seen:
        return False, "offline"

    now = datetime.utcnow()
    delta = now - last_seen

    if delta < timedelta(minutes=5):
        return True, "Đang hoạt động"
    elif delta < timedelta(hours=1):
        minutes = int(delta.total_seconds() // 60)
        return False, f"Đã offline {minutes} phút trước"
    elif delta < timedelta(hours=24):
        hours = int(delta.total_seconds() // 3600)
        return False, f"Đã offline {hours} giờ trước"
    else:
        return False, "Đã offline"

@app.route("/chat/send/<username>", methods=["POST"])
def send_message(username):
    current_username = session.get("username")
    if not current_username:
        return jsonify({"error": "Not logged in"}), 403

    sender = User.query.filter_by(username=current_username).first()
    receiver = User.query.filter_by(username=username).first()

    if not sender or not receiver:
        return jsonify({"error": "User not found"}), 404

    # Kiểm tra bị chặn
    if sender in receiver.blocked_users:
        return jsonify({"success": False, "error": "BẠN ĐÃ BỊ CHẶN"})

    # Cập nhật trạng thái hoạt động
    sender.last_seen = datetime.utcnow()
    db.session.commit()

    # Nhận dữ liệu gửi
    text = request.form.get("text", "").strip()
    images = request.files.getlist("images")
    voice_file = request.files.get("voice")
    reply_raw = request.form.get("reply_to")
    reply_to = None

    if reply_raw:
        try:
            raw = json.loads(reply_raw)
            reply_id = raw.get("id")
            try:
                reply_msg = Message.query.get(int(reply_id)) if reply_id is not None else None
                reply_sender = User.query.filter_by(username=reply_msg.sender).first() if reply_msg else None
            except Exception:
                reply_msg = None

            if reply_msg:
                if reply_msg.unsent:
                    reply_text = "Tin nhắn đã được thu hồi."
                elif reply_msg.content:
                    reply_text = reply_msg.content
                elif reply_msg.image_urls:
                    reply_text = "[ảnh]"
                elif reply_msg.voice_url:
                    reply_text = "[voice]"
                else:
                    reply_text = "Tin nhắn không tồn tại."

                reply_to = {
                    "id": reply_msg.id,
                    "sender": reply_msg.sender,
                    "text": reply_text,
                    "unsent": reply_msg.unsent,
                    "fullname": reply_sender.fullname if reply_sender and reply_sender.fullname else reply_msg.sender
                }
                reply_user = User.query.filter_by(username=reply_msg.sender).first()
                reply_to["fullname"] = reply_user.fullname if reply_user and reply_user.fullname else reply_msg.sender
            else:
                reply_to = {
                    "id": int(reply_id),
                    "sender": raw.get("sender"),
                    "text": "Tin nhắn không tồn tại.",
                    "unsent": False
                }
                fallback_user = User.query.filter_by(username=raw.get("sender")).first()
                reply_to["fullname"] = fallback_user.fullname if fallback_user and fallback_user.fullname else raw.get("sender")

        except Exception:
            reply_to = None

    image_urls = []
    voice_url = None

    # Xử lý ảnh
    for image in images:
        if image and image.filename:
            ext = image.filename.rsplit(".", 1)[-1].lower()
            if ext in ["png", "jpg", "jpeg", "gif"]:
                filename = f"{uuid.uuid4().hex}.{ext}"
                save_path = os.path.join("static", "images", "uploads", filename)
                os.makedirs(os.path.dirname(save_path), exist_ok=True)
                try:
                    image.save(save_path)
                    start_time = time.time()
                    while not os.path.exists(save_path) or os.path.getsize(save_path) == 0:
                        if time.time() - start_time > 2.0:
                            raise Exception("File save timeout")
                        time.sleep(0.05)
                    image_urls.append(f"/static/images/uploads/{filename}")
                except Exception:
                    pass

    # Xử lý voice
    if voice_file and voice_file.filename:
        ext = voice_file.filename.rsplit(".", 1)[-1].lower()
        if ext in ["mp3", "wav", "m4a", "ogg", "webm"]:
            filename = f"{uuid.uuid4().hex}.{ext}"
            save_path = os.path.join("static", "voices", filename)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            try:
                voice_file.save(save_path)
                start_time = time.time()
                while not os.path.exists(save_path) or os.path.getsize(save_path) == 0:
                    if time.time() - start_time > 2.0:
                        raise Exception("Voice save timeout")
                    time.sleep(0.05)
                voice_url = f"/static/voices/{filename}"
            except Exception:
                pass

    # Xử lý video
    video_file = request.files.get("video")
    video_url = None
    if video_file and video_file.filename:
        ext = video_file.filename.rsplit(".", 1)[-1].lower()
        if ext in ["mp4", "webm", "mov", "avi", "mkv"]:
            filename = f"{uuid.uuid4().hex}.{ext}"
            save_path = os.path.join("static", "videos", filename)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            try:
                video_file.save(save_path)
                start_time = time.time()
                while not os.path.exists(save_path) or os.path.getsize(save_path) == 0:
                    if time.time() - start_time > 2.0:
                        raise Exception("Video save timeout")
                    time.sleep(0.05)
                video_url = f"/static/videos/{filename}"
            except Exception:
                pass

    if not text and not image_urls and not voice_url and not video_url:
        return jsonify({"error": "No content"}), 400

    # Tạo chat_key
    chat_key_1 = f"{current_username}__{username}"
    chat_key_2 = f"{username}__{current_username}"
    chat_key = chat_key_1 if current_username < username else chat_key_2

    # Lưu vào PostgreSQL
    msg = Message(
        chat_key=chat_key,
        sender=current_username,
        receiver=username,
        content=text if text else None,
        image_urls=image_urls,
        voice_url=voice_url,
        video_url=video_url,
        timestamp=datetime.utcnow() + timedelta(hours=7),
        read=False,
        reply_to=reply_to
    )

    db.session.add(msg)
    db.session.commit()

    # Emit đến người nhận
    socketio.emit("private_message", {
        "from": current_username,
        "to": username,
        "text": msg.content,
        "image_urls": msg.image_urls,
        "voice_url": msg.voice_url,
        "video_url": msg.video_url,
        "time": msg.timestamp.strftime("%H:%M"),
        "reply_to": reply_to,
    }, room=f"user_{username}")

    socketio.emit("new_unread_message", {
        "from": current_username
    }, room=f"user_{username}")

    return jsonify({
        "success": True,
        "message": {
            "sender": msg.sender,
            "text": msg.content,
            "image_urls": msg.image_urls,
            "voice_url": msg.voice_url,
            "video_url": msg.video_url,
            "time": msg.timestamp.strftime("%H:%M"),
            "time_full": msg.timestamp.strftime("%Y-%m-%d %H:%M:%S"),
            "reply_to": reply_to
        }
    })


@app.route("/chat/<username>", methods=["GET"])
def chat(username):
    from models.message import Message
    current_username = session.get("username")
    if not current_username:
        return redirect("/login")
    current_user = User.query.filter_by(username=current_username).first()
    friend_user = User.query.filter_by(username=username).first()

    if not friend_user:
        return "User not found", 404

    current_user.last_seen = datetime.utcnow()
    current_user.online = True
    db.session.commit()

    chat_key_1 = f"{current_username}__{username}"
    chat_key_2 = f"{username}__{current_username}"
    chat_key = chat_key_1 if current_username < username else chat_key_2

    messages = Message.query.filter_by(chat_key=chat_key).order_by(Message.timestamp).all()
    now = datetime.now(VN_TZ) 
    messages_list = []
    for msg in messages:
        if msg.deleted_for and current_username in msg.deleted_for:
            continue

        vn_time = msg.timestamp.replace(tzinfo=VN_TZ)

        days_since = (now - vn_time).days
        days_left  = 30 - days_since if days_since >= 25 else None
        reply_data = None
        if isinstance(msg.reply_to, dict):
            reply_data = msg.reply_to.copy()
            reply_msg_id = reply_data.get("id")

            if reply_msg_id:
                try:
                    real_id = int(reply_msg_id)
                    reply_msg = Message.query.get(real_id)
                    if not reply_msg:
                        reply_data["text"] = "Tin nhắn không tồn tại."
                    else:
                        reply_sender = User.query.filter_by(username=reply_msg.sender).first()
                        reply_data["fullname"] = reply_sender.fullname if reply_sender and reply_sender.fullname else reply_msg.sender

                        if reply_msg.unsent:
                            reply_data["text"] = "Tin nhắn đã được thu hồi."
                        else:
                            reply_data["text"] = reply_msg.content or "[Ảnh/voice]"
                            if not reply_msg.content and not reply_msg.image_urls and not reply_msg.voice_url:
                                reply_data["text"] = "Tin nhắn không tồn tại."
                except Exception:
                    reply_data["text"] = "Tin nhắn không tồn tại."
            else:
                reply_data["text"] = "Tin nhắn không tồn tại."
        messages_list.append({
            "id":         msg.id,
            "sender":     msg.sender,
            "text":       "(đã thu hồi)" if msg.unsent else msg.content,
            "image_urls": [] if msg.unsent else (msg.image_urls or []),
            "voice_url": None if msg.unsent else msg.voice_url,
            "video_url":  None if msg.unsent else msg.video_url, 
            "unsent":     msg.unsent,
            "time":       vn_time.strftime("%H:%M"),
            "time_full":  vn_time.strftime("%Y-%m-%d %H:%M:%S"),
            "days_left":  days_left,
            "reply_to": reply_data
        })

    friend_name = friend_user.fullname or friend_user.username
    hide_online = friend_user.privacy.get("hide_online", False) if friend_user.privacy else False

    if hide_online:
        online = False
        status_text = ""
    else:
        online, status_text = get_status(friend_user.last_seen)

    return render_template("chat.html",
        username=current_username,
        friend_username=username,
        current_user=current_user,
        friend=friend_user,
        friend_name=friend_name,
        friend_status=status_text,
        messages=messages_list,
        online=online,
        hide_online=hide_online
    )

@app.route("/upload_video", methods=["POST"])
def upload_video():
    video_file = request.files.get("video")
    if not video_file or not video_file.filename:
        return jsonify({"error": "No video uploaded"}), 400

    ext = video_file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ["mp4", "webm", "mov", "avi", "mkv"]:
        return jsonify({"error": "Invalid video format"}), 400

    filename = f"{uuid.uuid4().hex}.{ext}"
    save_path = os.path.join("static", "videos", filename)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)

    try:
        video_file.save(save_path)
        start_time = time.time()
        while not os.path.exists(save_path) or os.path.getsize(save_path) == 0:
            if time.time() - start_time > 2.0:
                raise Exception("Video save timeout")
            time.sleep(0.05)

        return jsonify({"success": True, "url": f"/static/videos/{filename}"})

    except Exception as e:
        return jsonify({"error": "Failed to save video"}), 500

@app.route("/chat/unsend/<msg_id>", methods=["POST"])
def unsend_message(msg_id):
    current_username = session.get("username")
    if not current_username:
        return jsonify({"success": False, "error": "Not logged in"}), 403

    msg = Message.query.get(msg_id)
    if not msg or msg.sender != current_username:
        return jsonify({"success": False, "error": "Not allowed"}), 403

    # Xoá ảnh/voice nếu có
    if msg.image_urls:
        for url in msg.image_urls:
            try:
                path = os.path.join(app.root_path, url.strip("/"))
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass

    if msg.voice_url:
        try:
            path = os.path.join(app.root_path, msg.voice_url.strip("/"))
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass

    msg.content = None
    msg.image_urls = []
    msg.voice_url = None
    msg.unsent = True

    db.session.commit()

    socketio.emit("message_unsent", {
        "msg_id": msg_id,
        "from": current_username,
        "to": msg.receiver
    }, room=f"user_{msg.receiver}")

    socketio.emit("message_unsent", {
        "msg_id": msg_id,
        "from": current_username,
        "to": msg.receiver
    }, room=f"user_{msg.sender}")

    return jsonify({"success": True})


@app.route("/api/album-images")
def api_album_images():
    current_username = session.get("username")
    friend_username = request.args.get("friend")

    if not current_username or not friend_username:
        return jsonify([])

    messages_data = load_messages()
    chat_key_1 = f"{current_username}__{friend_username}"
    chat_key_2 = f"{friend_username}__{current_username}"
    chat_key = chat_key_1 if chat_key_1 in messages_data else chat_key_2

    chat_history = messages_data.get(chat_key, [])
    image_list = []

    for msg in chat_history:
        for img in msg.get("image_urls", []):
            full_path = os.path.join(".", img.lstrip("/"))
            if os.path.exists(full_path):
                image_list.append({"url": img})

    return jsonify(image_list[::-1])
@app.route("/admin/cleanup_chats_30_days")
def cleanup_old_chats():
    env_admin = os.getenv("ADMIN_USERNAME")
    if not (session.get("is_admin") or session.get("username") == env_admin):
        return "Unauthorized", 403

    messages = load_messages()
    removed_chats = []
    removed_images = []

    now = datetime.now()
    cutoff = now - timedelta(days=30)

    for chat_key in list(messages.keys()):
        chat_history = messages[chat_key]
        if not chat_history:
            continue

        # Kiểm tra nếu tất cả tin nhắn đều cũ hơn 30 ngày
        all_old = True
        for msg in chat_history:
            try:
                msg_time = datetime.strptime(msg["time_full"], "%Y-%m-%d %H:%M:%S")
                if msg_time > cutoff:
                    all_old = False
                    break
            except:
                all_old = False  # Nếu không có time_full hợp lệ thì bỏ qua xoá

        if all_old:
            # Xoá ảnh liên quan
            for msg in chat_history:
                for img_url in msg.get("image_urls", []):
                    img_path = os.path.join(".", img_url.lstrip("/"))
                    if os.path.exists(img_path):
                        os.remove(img_path)
                        removed_images.append(img_url)

            # Xoá đoạn chat
            removed_chats.append(chat_key)  
            del messages[chat_key]

    save_messages(messages)

    return f"✅ Đã xoá {len(removed_chats)} đoạn chat & {len(removed_images)} ảnh sau 30 ngày."
@app.route("/delete_old_chats")
def delete_old_chats():
    messages_data = load_messages()
    now = datetime.now()
    deleted = 0

    for key in list(messages_data.keys()):
        new_messages = []
        for msg in messages_data[key]:
            time_full = msg.get("time_full")
            if time_full:
                msg_time = datetime.strptime(time_full, "%Y-%m-%d %H:%M:%S")
                if now - msg_time < timedelta(days=30):
                    new_messages.append(msg)
                else:
                    # Nếu có ảnh thì xóa ảnh
                    for img in msg.get("image_urls", []):
                        img_path = img.lstrip("/")
                        if os.path.exists(img_path):
                            os.remove(img_path)
                    deleted += 1
            else:
                new_messages.append(msg)  # giữ lại tin cũ không có time_full
        messages_data[key] = new_messages

    save_messages(messages_data)
    return f"Đã xoá {deleted} tin nhắn cũ hơn 30 ngày."
def load_data():
    if not os.path.exists(DATA_FILE):
        return {}
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_data(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_username():
    # Lấy username hiện tại, ví dụ từ session
    return session.get('username')

@app.route('/friends/request', methods=['POST'])
def send_friend_request():
    current_username = session.get("username")
    to_username = request.json.get('to_user')

    if not current_username or not to_username:
        return jsonify({'success': False, 'message': 'Thiếu thông tin'}), 400
    if current_username == to_username:
        return jsonify({'success': False, 'message': 'Không thể gửi lời mời cho chính mình'}), 400

    from_user = User.query.filter_by(username=current_username).first()
    to_user = User.query.filter_by(username=to_username).first()

    if not from_user or not to_user:
        return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

    # Đã là bạn bè
    existing = Friend.query.filter_by(user_id=from_user.user_id, friend_id=to_user.user_id).first()
    if existing:
        return jsonify({'success': False, 'message': 'Đã là bạn bè'}), 400

    # Đã gửi lời mời
    already_sent = FriendRequest.query.filter_by(from_user_id=from_user.user_id, to_user_id=to_user.user_id).first()
    if already_sent:
        return jsonify({'success': False, 'message': 'Đã gửi lời mời trước đó'}), 400

    fr = FriendRequest(from_user_id=from_user.user_id, to_user_id=to_user.user_id)
    db.session.add(fr)
    db.session.commit()
    return jsonify({'success': True, 'message': 'Đã gửi lời mời kết bạn'})


@app.route('/friends/requests', methods=['GET'])
def get_friend_requests():
    current_username = session.get("username")
    current_user = User.query.filter_by(username=current_username).first()

    if not current_user:
        return jsonify([])

    requests = FriendRequest.query.filter_by(to_user_id=current_user.user_id).all()
    results = []
    for r in requests:
        from_user = User.query.filter_by(user_id=r.from_user_id).first()
        if from_user:
            results.append({
                'from_username': from_user.username,
                'from_name': from_user.fullname
            })

    return jsonify(results)

@app.route('/friends/remove', methods=['POST'])
def remove_friend():
    current_username = session.get("username")
    target_username = request.json.get('target_user')

    if not current_username or not target_username:
        return jsonify({'success': False, 'message': 'Thiếu thông tin'}), 400

    current_user = User.query.filter_by(username=current_username).first()
    target_user = User.query.filter_by(username=target_username).first()

    if not current_user or not target_user:
        return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

    # ❌ Xóa cả 2 chiều
    Friend.query.filter_by(user_id=current_user.user_id, friend_id=target_user.user_id).delete()
    Friend.query.filter_by(user_id=target_user.user_id, friend_id=current_user.user_id).delete()

    db.session.commit()

    return jsonify({'success': True, 'message': f'Đã xoá bạn với {target_username}'})

@app.route('/friends/requests/accept', methods=['POST'])
def accept_friend_request():
    current_username = session.get("username")
    from_username = request.json.get('from_user')

    if not current_username or not from_username:
        return jsonify({'success': False, 'message': 'Thiếu thông tin'}), 400

    current_user = User.query.filter_by(username=current_username).first()
    from_user = User.query.filter_by(username=from_username).first()

    if not current_user or not from_user:
        return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

    already_friends = Friend.query.filter_by(user_id=current_user.user_id, friend_id=from_user.user_id).first()
    if already_friends:
        return jsonify({'success': False, 'message': 'Đã là bạn bè'}), 400

    # ✅ Thêm 2 chiều
    db.session.add(Friend(user_id=current_user.user_id, friend_id=from_user.user_id))
    db.session.add(Friend(user_id=from_user.user_id, friend_id=current_user.user_id))

    # ✅ Xoá lời mời kết bạn
    FriendRequest.query.filter_by(from_user_id=from_user.user_id, to_user_id=current_user.user_id).delete()

    db.session.commit()

    return jsonify({'success': True, 'message': 'Đã chấp nhận lời mời', 'from_username': from_username, 'from_name': from_user.fullname})

@app.route('/friends/requests/reject', methods=['POST'])
def reject_friend_request():
    current_username = session.get("username")
    from_username = request.json.get('from_user')

    if not current_username or not from_username:
        return jsonify({'success': False, 'message': 'Thiếu thông tin'}), 400

    current_user = User.query.filter_by(username=current_username).first()
    from_user = User.query.filter_by(username=from_username).first()

    if not current_user or not from_user:
        return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

    req = FriendRequest.query.filter_by(from_user_id=from_user.user_id, to_user_id=current_user.user_id).first()
    if not req:
        return jsonify({'success': False, 'message': 'Không có lời mời từ người này'}), 400

    db.session.delete(req)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Đã từ chối lời mời'})


@app.route("/friends/list", methods=["GET"])
def get_friends_list():
    current_username = session.get("username")
    if not current_username:
        return jsonify([])

    current_user = User.query.filter_by(username=current_username).first()
    if not current_user:
        return jsonify([])

    now = datetime.utcnow()
    result = []

    friend_links = Friend.query.filter(
        (Friend.user_id == current_user.user_id) | (Friend.friend_id == current_user.user_id)
    ).all()

    seen = set()

    for link in friend_links:
        friend_id = link.friend_id if link.user_id == current_user.user_id else link.user_id
        friend = User.query.filter_by(user_id=friend_id).first()
        if not friend or friend.username in seen:
            continue
        seen.add(friend.username)

        # ------ trạng thái online ------
        last_active = friend.last_seen
        is_online = (now - last_active) < timedelta(seconds=60) if last_active else False

        # ------ xác định chat_key ------
        chat_key_1 = f"{current_username}__{friend.username}"
        chat_key_2 = f"{friend.username}__{current_username}"
        chat_key = chat_key_1 if current_username < friend.username else chat_key_2

        # ------ Đếm tin nhắn CHƯA đọc gửi tới current user ------
        unread = Message.query.filter_by(
            chat_key=chat_key,
            receiver=current_username,
            read=False
        ).count()

        # ------ Kiểm tra tin chưa đọc mới nhất có phải chỉ 1 ảnh ------
        image_only = False
        if unread:
            newest = Message.query.filter_by(
                        chat_key=chat_key,
                        receiver=current_username,
                        read=False
                     ).order_by(Message.timestamp.desc()).first()
            if newest and newest.image_urls and not newest.content:
                image_only = len(newest.image_urls) == 1

        # ------ Gộp vào kết quả ------
        show_bio = not (friend.privacy and friend.privacy.get("hide_bio"))

        result.append({
            "username": friend.username,
            "display_name": friend.fullname or friend.username,
            "name": friend.fullname or friend.username,
            "bio": friend.bio if show_bio else "",
            "unread": unread,
            "image_only": image_only,
            "is_online": is_online,
            "avatar_url": friend.avatar_url or "/static/logos/logo.png",
            "hide_avatar": friend.privacy.get("hide_avatar") if friend.privacy else False,
            "chat_key": chat_key
        })

    return jsonify(result)
@app.route("/chat/mark_read/<username>", methods=["POST"])
def mark_messages_as_read(username):
    current_username = session.get("username")
    if not current_username:
        return jsonify({"error": "Not logged in"}), 403

    chat_key_1 = f"{current_username}__{username}"
    chat_key_2 = f"{username}__{current_username}"
    chat_key = chat_key_1 if current_username < username else chat_key_2

    # ✅ Cập nhật tất cả tin gửi tới current_user và chưa đọc
    Message.query.filter_by(chat_key=chat_key, receiver=current_username, read=False).update({"read": True})
    db.session.commit()

    return jsonify({"success": True})

@app.route("/chat/unread_status")
def unread_status():
    username = session.get("username")
    if not username:
        return jsonify({})

    unread = Message.query.filter_by(receiver=username, read=False).all()
    return jsonify({m.sender: True for m in unread})


@app.route("/rename", methods=["POST"])
def rename_friend():
    current_username = session.get("username")
    if not current_username:
        return "Unauthorized", 403
    
    data = request.get_json()
    target = data.get("target_username")
    nickname = data.get("nickname", "").strip()

    user_data = load_data()
    user_data.setdefault(current_username, {}).setdefault("nicknames", {})

    if nickname:
        user_data[current_username]["nicknames"][target] = nickname
    else:
        user_data[current_username]["nicknames"].pop(target, None)

    save_data(user_data)
    return "OK"





@app.route("/upload_image", methods=["POST"])
def upload_image():
      # ✅ Nếu đã hết lượt dùng thử AI Toán → chặn lại
    if session.get("username") != "admin":
        user_type = get_user_type()
        if user_type != "vip":
            if not check_lite_usage():
                return jsonify({
                    "reply": "⚠️ Bạn đã dùng hết 10 lượt AI miễn phí. Vui lòng nâng cấp VIP để tiếp tục sử dụng."
                })

    image = request.files.get("image")
    if image:
        filename = secure_filename(image.filename)
        save_path = os.path.join("static", "images", "uploads", filename)
        image.save(save_path)

        try:
            solution_raw = extract_with_gpt_vision(save_path)
            session["last_ai_math_answer"] = solution_raw
            return jsonify({"reply": str (solution_raw)})
        except Exception as e:
            return jsonify({"reply": "⚠️ Ảnh có vài chi tiết hơi mờ,Anh/Chị có thể chụp lại rõ hơn rồi gửi lại giúp em nhé."})
    else:
        return jsonify({"reply": "❌ Không nhận được ảnh nào."})




@app.route("/upload_pdf", methods=["POST"])
def upload_pdf():
    file = request.files.get("pdf")
    if not file:
        return jsonify({"reply": "❌ Không nhận được file PDF."})

    pdf_text = ""
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                pdf_text += page.get_text()

        reply = call_gpt_viet(pdf_text)

        reply = "✏️ Rồi nha, em bắt đầu giải từng câu cho anh/chị nè!\n\n" + reply
        reply += "\n\n💬 Nếu cần em giải tiếp bài khác thì cứ gửi thêm ảnh hoặc gõ tiếp nhé!"
        return jsonify({"reply": reply})
    except Exception as e:
        return jsonify({"reply": "⚠️ Không thể xử lý file PDF. Vui lòng thử lại với file khác."})
@app.route("/update_ai_personality", methods=["POST"])
def update_ai_personality():
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"success": False})
    
    user = User.query.filter_by(user_id=user_id).first()
    if not user or not user.vip_gpt or not user.vip_gpt.startswith("SLV"):
        return jsonify({"success": False, "message": "Bạn không có quyền đổi tính cách AI."})

    new_personality = request.form.get("personality", "").strip()
    user.ai_personality = new_personality
    db.session.commit()
    return jsonify({"success": True})

#==========GPT==========#
def generate_obfuscated_slug():
    part1 = "0x" + secrets.token_hex(4)      
    part2 = secrets.token_hex(3)            
    part3 = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(8))  # obaq1267
    return f"{part1}-{part2}-?{part3}"
@app.route("/solvervietAIwatch", methods=["GET"])
def redirect_to_fancy_slug():
    # Nếu đã có slug và session_id → giữ nguyên, không reset
    slug_data = session.get("slug_to_session")
    if slug_data and "slug" in slug_data and "session_id" in slug_data:
        return redirect(f"/solvervietAIwatch/{slug_data['slug']}")
    
    # Nếu chưa có → tạo mới
    chat_session_id = str(uuid.uuid4())
    fancy_slug = generate_obfuscated_slug()
    session["chat_session_id"] = chat_session_id
    session["slug_to_session"] = {
        "slug": fancy_slug,
        "session_id": chat_session_id
    }

    return redirect(f"/solvervietAIwatch/{fancy_slug}")
@app.route("/solvervietAIwatch/<chat_uuid>", methods=["GET"])
def gpt_chat(chat_uuid):
    print("== SESSION SLUG ==", session.get("slug_to_session"))
    print("== CHAT SESSION ID ==", session.get("chat_session_id"))
    if is_maintenance("gpt_chat"):
        return render_template("gpt_chat.html",
            user_id=None,
            username=None,
            user_vip_gpt=False,
            user_vip_al=False,
            user_lite_exhausted=False,
            is_vip_chat=False,
            is_verified=False,
            chat_history=[],
            chat_id=None,
            chat_title="",
            is_maintenance=True,
            user=None  
        )

    cleanup_old_chats()

    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    if user.is_blocked:
        session.clear()
        return render_template("login.html", error= "🚫 Tài khoản của bạn đã bị khóa. Bạn sẽ được đăng xuất. Nếu đây là nhầm lẫn, vui lòng <a href='/appeal' style='color:#4ea6ff;'>liên hệ hỗ trợ</a>.")

    now = datetime.utcnow()

    # ✅ Ưu tiên gói GPT nếu còn hạn
    is_vip = False
    if username == "admin":
        is_vip = True
    elif user.vip_until_gpt and user.vip_until_gpt >= now:
        is_vip = True

    # ❗ Nếu không có gói GPT → kiểm tra AI Lite
    if not is_vip and user.vip_until_lite and user.vip_until_lite >= now:
        return redirect("/chat_lite")

    # ✅ Chặn nếu là gói 5 ngày và đã hết 100 lượt (trừ khi gpt_unlimited)
    if is_vip and user.vip_gpt == "5day" and not user.gpt_unlimited:
        usage_today = user.gpt_usage_today or 0
        usage_date = user.gpt_usage_date or ""
        today = now.strftime("%Y-%m-%d")
        if usage_date != today:
            usage_today = 0  # reset lại
        if usage_today >= 100:
            return render_template("gpt_chat.html",
                user_id=user.user_id, 
                username=user.username,
                user_vip_gpt=True,
                user_vip_al=True,
                user_lite_exhausted=False,
                is_vip_chat=False,
                is_verified=user.is_verified,
                chat_history=[],
                chat_id=None,
                chat_title=""
            )

    # ✅ Kiểm tra xác thực và giới hạn Lite
    lite_used = user.lite_usage or 0
    is_verified = user.is_verified
    lite_exhausted = not is_verified and lite_used >= 5

    # 🖼️ Lấy toàn bộ ảnh trước để dùng map
    image_logs = ImageHistory.query.filter_by(user_id=user.user_id).order_by(ImageHistory.created_at).all()
    image_map = {log.message_index: log.image_url for log in image_logs}

    # 🗂 Chat history
    chat_history = []
    slug_data = session.get("slug_to_session")
    if slug_data and slug_data.get("slug") == chat_uuid:
        session_id = slug_data["session_id"]
    else:
        session_id = str(uuid.uuid4())
        session["slug_to_session"] = {
            "slug": chat_uuid,
            "session_id": session_id
        }

    session["chat_session_id"] = session_id

    history_records = ChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=session_id
    ).order_by(ChatHistory.timestamp).all()

    for idx, record in enumerate(history_records):
        if idx in image_map:
            chat_history.append({
                "role": "user",
                "type": "image_url",
                "image_url": {"url": image_map[idx]}
            })
        else:
            entry = {
                "role": record.role,
                "content": record.content
            }

            # ✅ Nếu là câu phản hồi sau khi thả cảm xúc
            if record.role == "assistant" and record.content.startswith("[Thả cảm xúc "):
                entry["source"] = "emoji"
                # Tách emoji từ đoạn đầu
                start = record.content.find("❤️")
                if start == -1:
                    for emo in ["😂", "😢", "🤔", "😡"]:
                        if emo in record.content:
                            entry["emoji_used"] = emo
                            break
                else:
                    entry["emoji_used"] = "❤️"

            chat_history.append(entry)

    return render_template("gpt_chat.html",
        user_id=user.user_id, 
        username=user.username,
        user_vip_gpt=is_vip,
        user_vip_al=is_vip,
        user_lite_exhausted=lite_exhausted,
        is_vip_chat=is_vip,
        is_verified=is_verified,
        message_from_home=None,
        chat_history=chat_history,  
        chat_id=session["chat_session_id"],
        chat_title="",
        is_maintenance=False,
        user=user,
        chat_session_id=session["chat_session_id"],
        chat_uuid=chat_uuid
    )
def reset_image_quota_if_new_day(user):
    today = datetime.now().date()
    if not user.image_quota_reset_at or user.image_quota_reset_at.date() != today:
        user.image_quota_today = 10
        user.image_quota_reset_at = datetime.now()

def should_generate_image(message: str, has_uploaded_image: bool) -> bool:
    """
    Trả về True nếu hệ thống nên tạo ảnh từ mô tả (prompt).
    Trả về False nếu có ảnh upload hoặc không rõ ràng, hoặc là các yêu cầu không liên quan đến hình ảnh.
    """

    if has_uploaded_image:
        return False

    lower_msg = message.lower().strip()

    # ===== TẦNG 1: LỌC NHANH BẰNG TỪ KHÓA =====
    non_image_keywords = [
        "ý nghĩa", "bài hát", "bài thơ", "bài tập", "giải thích", "nội dung", "lời bài", 
        "gpt vision", "upload ảnh", "đây là ảnh", "trong ảnh", "phân tích ảnh",
        "câu hỏi", "kiến thức", "ai có biết", "là gì", "viết lại", "dịch", "tóm tắt", 
        "so sánh", "phân biệt", "hỏi", "nghĩa là", "mang ý nghĩa gì", "phát âm"
    ]

    if any(kw in lower_msg for kw in non_image_keywords):
        return False

    # Các câu hỏi về khả năng vẽ (chung chung) → cũng loại trừ
    if (
        ("có thể" in lower_msg or "có khả năng" in lower_msg or "vẽ được" in lower_msg)
        and ("tạo ảnh" in lower_msg or "vẽ" in lower_msg or "minh họa" in lower_msg)
        and any(kw in lower_msg for kw in ["không", "chứ", "à", "?"])
    ):
        return False

    # ===== TẦNG 2: GỌI GPT (có thể thay model tùy mức chính xác) =====
    try:
        client = create_openai_client("gpt-4o-mini")  # hoặc "gpt-4o" nếu cần chính xác hơn
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Bạn là bộ lọc yêu cầu tạo ảnh từ văn bản. Trả lời duy nhất dưới dạng JSON:\n"
                        '{"is_image_request": true} hoặc {"is_image_request": false}.\n\n'
                        "Trả lời `true` chỉ khi người dùng **thật sự yêu cầu tạo hình ảnh** từ mô tả cụ thể, "
                        "ví dụ như: 'Hãy vẽ...', 'Tạo ảnh...', 'Cho tôi một bức tranh...', 'Hãy minh họa...'.\n\n"
                        "Các trường hợp sau đây PHẢI trả lời `false`:\n"
                        "- Câu hỏi về ý nghĩa, lời bài hát, thơ, triết lý, tôn giáo, kiến thức, nội dung văn học\n"
                        "- Các yêu cầu như: 'ý nghĩa bài Cupid', 'Cupid là ai', 'Cupid tượng trưng cho điều gì'\n"
                        "- Các câu hỏi về ảnh đã upload, bài tập, GPT Vision, ảnh gửi kèm\n"
                        "- Những câu hỏi chung chung như: 'Bạn có vẽ được không?', 'GPT có tạo được ảnh không?'\n\n"
                        "Không cần giải thích gì thêm, chỉ trả về đúng JSON như yêu cầu."
                    )
                },
                {
                    "role": "user",
                    "content": f'Câu sau có phải là yêu cầu tạo ảnh không?\n"{message}"'
                }
            ]
        )

        content = response.choices[0].message.content.strip()

        # Trích JSON trong câu trả lời
        if "{" in content and "}" in content:
            content = content[content.find("{"):content.rfind("}") + 1]

        result = json.loads(content)
        return result.get("is_image_request", False)

    except Exception as e:
        print(f"[Image Intent Check Error] {e}")
        return False

@app.route("/core_inference", methods=["POST"])
def core_inference():
    
    if "username" not in session:
        return jsonify({"success": False, "reply": " Bạn chưa đăng nhập."})

    username = session["username"]
    user_id = session.get("user_id", "guest")
    session_id = request.form.get("session_id")  
    if not session_id:
        session_id = session.get("chat_session_id")  
    if not session_id:
        session_id = str(uuid.uuid4())
        session["chat_session_id"] = session_id

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "reply": "❌ Không tìm thấy người dùng."})
    # --- NGĂN SPAM GỬI LIÊN TIẾP ---
    reset_pending_if_needed(user)
    if user.pending_messages is None:
        user.pending_messages = 0
    if user.pending_messages >= 4:
        return jsonify({
            "success": False,
            "reply": "⚠️ Bạn đang gửi quá nhanh. Vui lòng đợi AI phản hồi trước khi gửi tiếp."
        })
    now = datetime.utcnow()
    if user.last_sent_at and (now - user.last_sent_at).total_seconds() < 3:
        return jsonify({
            "success": False,
            "reply": "⏳ Bạn thao tác quá nhanh. Vui lòng chờ 3 giây giữa các lần gửi."
        })
    user.pending_messages += 1
    user.last_sent_at = now
    db.session.commit()

    if user.over_quota_block:
        return jsonify({
            "success": False,
            "reply": (
                    "Tài khoản của bạn đã đạt giới hạn sử dụng hôm nay. Hãy quay lại sau 24 giờ nhé.br>"
                    " You’ve reached today’s usage limit. Please try again in 24 hours."
            )
        })
    cap_nhat_trang_thai_vip(user)
    if username != "admin" and not user.vip_gpt_ai:
        return jsonify({"success": False, "reply": " Bạn chưa mở khóa gói Solver Chat."})
    reset_usage_if_new_day(user)
    reset_image_quota_if_new_day(user)
    can_chat = check_and_update_quota(user)
    if not can_chat:
        return jsonify({
           "reply": (
                    "Tài khoản của bạn đã đạt giới hạn sử dụng hôm nay. Hãy quay lại sau 24 giờ nhé.<br>"
                    "  You’ve reached today’s usage limit. Please try again in 24 hours."
                )
        })
    message = request.form.get("message", "").strip()
    reply_to = request.form.get("reply_to", "").strip()
    img_url = None
    try:
        history = json.loads(request.form.get("history", "[]"))
        has_uploaded_image = bool(request.files)
        if should_generate_image(message, has_uploaded_image):
            image_prompt = rewrite_prompt_for_image(message)
            if not image_prompt:
                return jsonify({"success": False, "reply": "The request has been denied, please try again later."})
            if user.image_quota_today is not None and user.image_quota_today <= 0:
                return jsonify({
                    "success": False,
                    "reply": """
                    <span style="color:#ffcc66; font-style:italic;">
                        ⚠️ The image generation system seems temporarily overloaded. Please try again later.
                    </span>
                    """
                })
            img_url = generate_image_from_prompt(image_prompt)
            if not img_url:
                return jsonify({"success": False, "reply": "Không tạo được ảnh, thử lại sau nha!"})
            save_generated_image_log_backend(user.user_id, img_url, image_prompt, source="gpt_chat")
            user.image_quota_today = (user.image_quota_today or 0) - 1
            user.image_quota_reset_at = datetime.now()

            reply_text = """
                    <div style='text-align:center; margin-bottom:4px; font-size:14px; color:#aaa; animation:fadeIn 0.8s ease;'>
                    <div style='display:inline-flex; align-items:center; gap:6px; justify-content:center;'>
                        <svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" fill="none" stroke="#0ea5e9"
                        stroke-width="2" stroke-linecap="round" stroke-linejoin="round" class="lucide lucide-palette"
                        style="animation: pulseGlow 1.5s infinite ease-in-out;">
                        <circle cx="13.5" cy="6.5" r=".5"/>
                        <circle cx="17.5" cy="10.5" r=".5"/>
                        <circle cx="8.5" cy="7.5" r=".5"/>
                        <circle cx="6.5" cy="12.5" r=".5"/>
                        <path d="M12 22a10 10 0 1 1 8.5-4.7c-.8 1.4-2.6 1.2-3.5.6-.8-.5-1.6-1.4-2.5-1.4-1.3 0-2 1.3-2 2.5 0 1.3 1.2 2.5 2 3z"/>
                        </svg>
                        <span style="font-style:italic;">AI vừa tạo ảnh này</span>
                    </div>
                    </div>
                    """ + f"<br><img src='{img_url}' style='max-width:60%; border-radius:12px; margin:10px auto; display:block; box-shadow:0 2px 8px rgba(0,0,0,0.2);'>" 

            save_chat_sql(
                user_id=user.user_id,
                session_id=session_id,
                history=history + [
                    {"role": "user", "content": message},
                    {"role": "assistant", "content": reply_text, "image_url": img_url}
                ]
            )
            increase_gpt_usage(user)
            user.pending_messages = 0
            db.session.commit()
            return jsonify({"success": True, "reply": reply_text, "img_url": img_url})
    except Exception as e:
        history = []
    try:
        vision_texts = []
        vision_image_urls = []  
        vision_image_url = None
        for key in request.files:
            file = request.files[key]
            if file.filename:
                ext = os.path.splitext(file.filename)[-1]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = secure_filename(f"{timestamp}_{username}_{key}{ext}")
                user_folder = os.path.join('static', 'images', 'uploads', str(user_id))
                os.makedirs(user_folder, exist_ok=True)
                save_path = os.path.join(user_folder, filename)
                file.save(save_path)
                vision_output = extract_with_gpt_vision(save_path, user_request=message)
                vision_texts.append(f"[Ảnh: {filename}]\n{vision_output}")
                vision_image_url = f"/static/images/uploads/{user_id}/{filename}"
                vision_image_urls.append(vision_image_url)
        user_original_message = message  
        if vision_texts:
            vision_combined = "\n\n".join(vision_texts)
            if message:
                message += f"\n\n🖼 Dưới đây là nội dung AI trích từ ảnh:\n{vision_combined}"
            else:
                message = f"🖼 AI đã trích nội dung từ ảnh như sau:\n{vision_combined}"
        if reply_to:
            message = f'💬 Bạn đang trả lời lại đoạn AI trước đây:\n"{reply_to}"\n\n👉 Tin nhắn mới của bạn:\n{message}'
    
        reply = call_gpt_viet(vision_combined if vision_texts else message, history)
        increase_gpt_usage(user)
        save_chat_sql(
            user_id=user.user_id,
            session_id=session_id,
            history=history + [
                {
                    "role": "user",
                    "content": user_original_message,
                    "reply_to": reply_to,
                    "image_url": json.dumps(vision_image_urls)

                },
                {
                    "role": "assistant",
                    "content": reply,
                    "reply_to": reply_to,  
                    "image_url": None
                }
            ]
        )
        if user.vip_gpt_ai:  
            update_daily_usage(is_slv=True)
        else:
            update_daily_usage(is_slv=False)
        user = db.session.merge(user)

        try:
            db.session.commit()
        except Exception as e:
            pass
            from sqlalchemy.orm import Session

            # 🧠 Cập nhật ChatHistory như bình thường
            db.session.add(ChatHistory(
                user_id=user.user_id,
                session_id=session_id,
                role="user",
                content=message
            ))
            db.session.add(ChatHistory(
                user_id=user.user_id,
                session_id=session_id,
                role="assistant",
                content=reply
            ))
            today = now_vn().strftime("%Y-%m-%d")
            if user.gpt_usage_date != today:
                user.gpt_usage_today = 0
                user.gpt_usage_date = today
            user.gpt_usage_today += 1
            user = db.session.merge(user)  
            try:
                db.session.commit()
            except Exception as e:
                pass
        socketio.emit("user_usage_update", {
                    "user_id": user.user_id,
                    "username": user.username,
                    "gpt_usage": user.gpt_usage_today,
                    "lite_usage": user.lite_usage,
                    "over_quota_block": user.over_quota_block,
                    "vip_gpt": user.vip_gpt or "",
                    "vip_lite": bool(user.vip_ai_lite),
                    "quota_alert": (
                        (user.vip_gpt in ["15day", "30day"] and not user.gpt_unlimited and user.gpt_usage_today >= 200)
                        if user.vip_gpt else False
                    )
                }, room="admin_dashboard")
        try:
            if session.get("dev_mode"):
                full_history = []
                for msg in history:
                    role = msg.get("role", "")
                    content = msg.get("content", "")
                    full_history.append({"role": role, "content": content})
                full_history.append({"role": "user", "content": message})
                full_history.append({"role": "assistant", "content": reply})

                db.session.add(DevChatHistory(
                    user_id=user.user_id,
                    session_id=session_id,
                    history=full_history,
                    created_at=datetime.utcnow()
                ))
                db.session.commit()
        except Exception as e:
            pass
        user.pending_messages = 0
        db.session.commit()
        return jsonify({
            "success": True,
            "reply": reply,
            "img_url": vision_image_url or img_url
        })
    except Exception as e:
        pass
        return jsonify({
            "reply": " Hệ thống quá tải, bạn hãy thử lại trong ít phút nữa hoặc tạo đoạn chat mới."
        })
    
@app.route("/submit_rating", methods=["POST"])
def submit_rating():
    if "user_id" not in session:
        return jsonify({"error": "Not logged in"}), 401

    data = request.get_json()
    stars = data.get("score")
    comment = data.get("comment", "")
    user_id = session["user_id"]
    session_id = session.get("chat_session_id")  
    user = User.query.filter_by(user_id=user_id).first()
    ai_personality = getattr(session, "ai_personality", None) 
    package_type = (
        "GPT"
        if user and user.vip_gpt_ai
        else "Lite"
    )

    fb = Feedback(
        user_id=user_id,
        session_id=session_id,
        stars=stars,
        comment=comment.strip(),
        ai_personality=ai_personality,
        package_type=package_type,
    )
    db.session.add(fb)
    db.session.commit()

    return jsonify({"message": "Đã gửi đánh giá thành công!"})    
def increase_gpt_usage(user):
    today = now_vn().strftime("%Y-%m-%d")

    if user.gpt_usage_date != today:
        user.gpt_usage_today = 0
        user.gpt_usage_date = today

    user.gpt_usage_today += 1
#==========CHẾ ĐỘ DEV CODE===========#
@app.route("/solvervietCode", methods=["GET"])
def redirect_to_code_slug():
    # Nếu đã có rồi thì redirect lại slug cũ
    if session.get("code_slug") and session.get("code_session_id"):
        return redirect(f"/solvervietCode/{session['code_slug']}")

    # Ngược lại thì tạo mới
    code_session_id = str(uuid.uuid4())
    code_slug = "code-" + secrets.token_hex(6)

    session["code_session_id"] = code_session_id
    session["code_slug"] = code_slug

    return redirect(f"/solvervietCode/{code_slug}")

@app.route("/solvervietCode/<slug>", methods=["GET"])
def code_chat(slug):
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    if session.get("code_slug") != slug:
        session["code_slug"] = slug

        existing = DevChatHistory.query.filter_by(user_id=user.user_id).filter(
            DevChatHistory.session_id.like(f"%{slug}")
        ).order_by(DevChatHistory.created_at.desc()).first()

        if existing:
            session["code_session_id"] = existing.session_id
        else:
            # Nếu chưa có thì tạo mới
            session["code_session_id"] = f"{str(uuid.uuid4())[:8]}-{slug}"

    code_session_id = session.get("code_session_id")

    # ✅ Lấy toàn bộ lịch sử của session và gộp lại
    all_history = DevChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=code_session_id
    ).order_by(DevChatHistory.created_at.asc()).all()

    chat_history = []
    for item in all_history:
        if isinstance(item.history, list):
            for msg in item.history:
                if msg.get("role") == "user" and msg.get("type") == "file":
                    filename = msg.get("filename")
                    file_content = msg.get("content", "")
                    if filename:
                        msg = {
                            "role": "user",
                            "type": "text",
                            "content": "",
                            "attachedFiles": [{
                                "name": filename,
                                "content": file_content  # ✅ gắn luôn nội dung nếu có
                            }]
                        }
                chat_history.append(msg)
    return render_template("devcode.html",
        user=user,
        chat_session_id=code_session_id,
        slug=slug,
        chat_history_json=json.dumps(chat_history)
    )

def split_reply_to_blocks(content, language="plaintext"):
    pattern = re.compile(r"```(\w+)?\n([\s\S]*?)```")
    result = []
    last_end = 0

    for match in pattern.finditer(content):
        # đoạn text trước code
        if match.start() > last_end:
            text = content[last_end:match.start()].strip()
            if text:
                result.append({
                    "role": "assistant",
                    "type": "text",
                    "content": text
                })
        
        lang = match.group(1) or language
        code = match.group(2).strip()
        result.append({
            "role": "assistant",
            "type": "code",
            "language": lang.strip().lower(),
            "content": code
        })

        last_end = match.end()

    # phần text sau cùng (nếu có)
    if last_end < len(content):
        text = content[last_end:].strip()
        if text:
            result.append({
                "role": "assistant",
                "type": "text",
                "content": text
            })

    return result

from openai_config import (
    call_gpt_python, call_gpt_js, call_gpt_html,
    call_gpt_flutter, call_gpt_sql, call_gpt_java, call_gpt_cpp
)
@app.route("/code_infer", methods=["POST"])
def code_infer():
    if "username" not in session:
        return jsonify({"success": False, "reply": "Bạn chưa đăng nhập."})

    user = User.query.filter_by(username=session["username"]).first()
    if not user:
        return jsonify({"success": False, "reply": "❌ Không tìm thấy người dùng."})

    # ✅ Khóa spam gửi nếu AI chưa trả lời xong
    if session.get("code_infer_locked"):
        return jsonify({"success": False, "reply": "Vui lòng đợi AI trả lời xong trước khi gửi tiếp."})
    
    session["code_infer_locked"] = True  # khóa lại ngay

    try:
        message = request.form.get("message", "").strip()
        files_json = request.form.get("files", "[]")
        file_contents = []

        try:
            attached_files = json.loads(files_json)
            if not isinstance(attached_files, list): attached_files = []
        except Exception:
            attached_files = []

        for f in attached_files:
            fname = f.get("name", "file.txt")
            content = f.get("content", "")
            if len(content.strip()) > 0:
                file_contents.append(f"📄 Nội dung file **{fname}**:\n```text\n{content.strip()}\n```")
        if file_contents:
            message += "\n\n" + "\n\n".join(file_contents)
        session_id = request.form.get("session_id") or session.get("code_session_id")
        history_json = request.form.get("history", "[]")
        language = request.form.get("language", "python").lower()
        images_json = request.form.get("images", "[]")

        try:
            images = json.loads(images_json)
            if not isinstance(images, list): images = []
        except Exception: images = []

        try:
            history = json.loads(history_json)
            if not isinstance(history, list): history = []
        except Exception: history = []

        def clean_base64(img):
            return img.split(",")[-1]

        # ✅ Gọi GPT theo ngôn ngữ
        if language == "python":
            reply = call_gpt_python(message, history, images=images)
        elif language in ["js", "javascript"]:
            reply = call_gpt_js(message, history, images=images)
        elif language == "html":
            reply = call_gpt_html(message, history, images=images)
        elif language in ["flutter", "dart"]:
            reply = call_gpt_flutter(message, history, images=images)
        elif language == "sql":
            reply = call_gpt_sql(message, history, images=images)
        elif language == "java":
            reply = call_gpt_java(message, history, images=images)
        elif language in ["cpp", "c++"]:
            reply = call_gpt_cpp(message, history, images=images)
        else:
            reply = "❌ Ngôn ngữ không được hỗ trợ."

        assistant_blocks = split_reply_to_blocks(reply.strip(), language)
        user_blocks = []
        if message:
            user_blocks.append({
                "role": "user",
                "type": "text",
                 "content": request.form.get("message", "").strip()  
            })
        for f in attached_files:
            fname = f.get("name", "file.txt")
            fcontent = f.get("content", "")
            user_blocks.append({
                "role": "user",
                "type": "file",
                "filename": fname,
                "content": fcontent  
            }) 
        for img_b64 in images:
            user_blocks.append({
                "role": "user",
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{clean_base64(img_b64)}",
                    "detail": "auto"
                }
            })

        combined_history = history + user_blocks + assistant_blocks

        db.session.add(DevChatHistory(
            user_id=user.user_id,
            session_id=session_id,
            history=combined_history,
            created_at=datetime.utcnow()
        ))
        db.session.commit()

        return jsonify({
            "success": True,
            "reply": reply,
            "blocks": assistant_blocks
        })

    except Exception as e:
        print("❌ Lỗi xử lý GPT:", str(e))
        return jsonify({"success": False, "reply": "❌ Hệ thống gặp lỗi, thử lại sau."})

    finally:
        session["code_infer_locked"] = False

@app.route("/get_code_history/<slug>")
def get_code_history(slug):
    username = session.get("username")
    if not username:
        return jsonify({"success": False, "messages": []})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "messages": []})

    session_id = session.get("code_session_id")
    if not session_id:
        return jsonify({"success": False, "messages": []})

    # ✅ Lấy tất cả các bản ghi theo session_id và gộp lại
    all_history = DevChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=session_id
    ).order_by(DevChatHistory.created_at.asc()).all()

    combined = []
    for item in all_history:
        if isinstance(item.history, list):
            for block in item.history:
                # 🐞 THÊM DÒNG SAU ĐỂ DEBUG
                if block.get("role") == "user":
                    print("🧩 USER BLOCK:", block)
            combined.extend(item.history)

    return jsonify({
        "success": True,
        "messages": combined
    })

@app.route("/new_code_chat", methods=["POST"])
def new_code_chat():
    if "username" not in session:
        return jsonify({"success": False, "error": "Bạn chưa đăng nhập."})

    user = User.query.filter_by(username=session["username"]).first()
    if not user:
        return jsonify({"success": False, "error": "Không tìm thấy người dùng."})

    old_session_id = session.get("code_session_id")
    print("🔍 Session cũ:", old_session_id)

    # ✅ Xoá lịch sử nếu có session cũ
    if old_session_id:
        try:
            deleted = DevChatHistory.query.filter_by(user_id=user.user_id, session_id=old_session_id).delete()
            db.session.commit()
            print(f"✅ Đã xoá {deleted} bản ghi của session cũ: {old_session_id}")
        except Exception as e:
            print("❌ Lỗi khi xóa session cũ:", str(e))
            db.session.rollback()

    # 🔄 Tạo session mới
    code_session_id = f"{str(uuid.uuid4())[:8]}-code-{secrets.token_hex(6)}"
    code_slug = code_session_id.split("-code-")[-1]
    session["code_session_id"] = code_session_id
    session["code_slug"] = code_slug

    print(f"🚀 Tạo session mới: {code_session_id}")
    return jsonify({"success": True, "redirect": f"/solvervietCode/{code_slug}"})



def today_str():
    return now_vn().strftime("%Y-%m-%d")

def check_and_update_quota(user):
    now = datetime.utcnow()

    # ---------- 1. Reset khi sang ngày mới ----------
    if user.gpt_usage_date != today_str():
        user.gpt_usage_today = 0
        user.gpt_usage_date  = today_str()

    # ---------- 2. Tự động gỡ chặn nếu đã quá 24h ----------
    if user.over_quota_block and user.over_quota_block_at:
        if now - user.over_quota_block_at > timedelta(hours=24):
            user.over_quota_block = False
            user.over_quota_block_at = None
            user.gpt_usage_today = 0  # reset lượt để tránh chặn lại liền

    # ---------- 3. Xác định hạn mức ----------
    if user.gpt_unlimited:
        return True  # Gói không giới hạn

    limit = None
    if user.vip_gpt == "5day":
        limit = 100
    elif user.vip_gpt in ("15day", "30day"):
        limit = 200

    # Nếu không có gói tính lượt → cho phép
    if limit is None:
        return True

    # ---------- 4. Kiểm tra quota hiện tại ----------
    if user.gpt_usage_today >= limit:
        if not user.over_quota_block:
            user.over_quota_block = True
            user.over_quota_block_at = now
        return False  # Vượt quota → chặn
    else:
        return True  # Chưa vượt quota → cho phép

def reset_usage_if_new_day(user):
    today = now_vn().strftime("%Y-%m-%d")
    if user.gpt_usage_date != today:
        user.gpt_usage_today = 0
        user.gpt_usage_date = today
import datetime as dt

def to_dt(val):
    if not val:
        return None
    if isinstance(val, dt.datetime):
        return val
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            return dt.datetime.strptime(str(val), fmt)
        except:
            pass
    return None
def reset_pending_if_needed(user):
    now = datetime.utcnow()

    if user.last_sent_at and (now - user.last_sent_at) >= timedelta(minutes=17):
        user.pending_messages = 0
        db.session.commit()


@app.route("/chat_ai_lite", methods=["POST"])
def chat_ai_lite():
    if is_maintenance("chat_ai_lite_daily") or is_maintenance("all"):
        return jsonify({"reply": "🚧 Hệ thống đang bảo trì. Vui lòng quay lại sau!"}), 503

    username = session.get("username")
    user = User.query.filter_by(username=username).first() if username else None
    if not user:
        session.clear()
        return jsonify({"reply": "❌ Phiên đăng nhập của bạn không hợp lệ. Vui lòng đăng nhập lại."}), 403
    reset_pending_if_needed(user)
    if user.pending_messages is None:
        user.pending_messages = 0

    if user.pending_messages >= 4:
        return jsonify({
            "reply": "⚠️ Bạn đang gửi quá nhanh. Vui lòng quay lại sau 17 phút để tiếp tục trò chuyện."
        })

    now = datetime.utcnow()
    if user.last_sent_at and (now - user.last_sent_at).total_seconds() < 3:
        return jsonify({
            "reply": "⏳ Bạn thao tác quá nhanh. Vui lòng chờ 3 giây giữa các lần gửi."
        })

    user.pending_messages += 1
    user.last_sent_at = datetime.utcnow()
    user.last_sent_at = now
    db.session.commit()
    if user.is_blocked:
        session.clear()
        return jsonify({
            "reply": "🚫 Tài khoản của bạn đã bị khóa. Bạn sẽ được đăng xuất. Nếu đây là nhầm lẫn, vui lòng <a href='/appeal' style='color:#4ea6ff;'>liên hệ hỗ trợ</a>."
        }), 403

    now_local = now_vn()  
    vip_dt = to_dt(user.vip_until_lite)
    if vip_dt and now_local <= vip_dt:
        if is_maintenance("chat_lite"):
            return jsonify({"reply": "🚧 AI Lite đang bảo trì. Vui lòng quay lại sau!"}), 503
    elif user.is_verified and is_maintenance("chat_ai_lite"):
        return jsonify({"reply": "🚧 AI Free 15 lượt đang bảo trì. Vui lòng quay lại sau!"}), 503

    usage_check = check_lite_usage(user)
    user = db.session.merge(user)
    if usage_check == "require_verification":
        return jsonify({"reply": "Bạn cần <a href='/verify-otp' style='color:#00e676;font-weight:bold;'>xác minh tài khoản</a> để sử dụng tiếp."})
    elif usage_check is False:
        now = now_vn()
        if user.is_blocked:
            return jsonify({
                "reply": "Tài khoản của bạn đã bị khóa. Nếu đây là nhầm lẫn, vui lòng gửi khiếu nại tại <a href='/appeal' style='color:#4ea6ff;'>đây</a>."
            })
        try:
            vip_dt = to_dt(user.vip_until_lite)
            if vip_dt and now <= vip_dt:
                return jsonify({
                    "reply": (
                            "<strong> Gói AI Lite:</strong> Đã dùng hết 70 lượt hôm nay. "
                            "Vui lòng quay lại vào ngày mai.<br>"
                            "<em>AI Lite:</em> 70 turns used today. Please try again tomorrow!"
                        )
                })
        except:
            pass
        if user.is_verified:
            return jsonify({
                "reply": "You have reached today's limit, please come back tomorrow"
            })

        # ✅ Nếu chưa xác thực → gợi ý xác thực
        return jsonify({
            "reply": " Bạn đã dùng hết 5 lượt miễn phí. Vui lòng xác thực tài khoản để nhận thêm 10 lượt nữa!"
        })
    # --- xử lý message ---
    message = request.form.get("message", "").strip()
    if not message:
        return jsonify({"reply": "Bạn chưa nhập nội dung câu hỏi."})

    history_str = request.form.get("history", "[]")
    try:
        history = json.loads(history_str)
        if isinstance(history, list):
            history = history[-15:]
    except Exception as e:
        history = []
    if re.search(r"\b(vẽ ảnh|tạo ảnh|minh hoạ|minh họa|vẽ tranh|vẽ|tạo tranh)\b", message, re.IGNORECASE):
        reply = "Rất tiếc, phiên bản AI này không hỗ trợ tạo ảnh hoặc vẽ. Vui lòng thử lại sau.\nSorry, this AI version does not support image generation or drawing. Please try again later."
        history.append({"role": "user", "content": message})
        history.append({"role": "assistant", "content": reply})

        user_id = session.get("user_id")
        if user_id:
            if "chat_session_id" not in session:
                session["chat_session_id"] = str(uuid.uuid4())
            session_id = session["chat_session_id"]
            save_chat_sql(user_id, session_id, history)

       
        return jsonify({"reply": reply})
    # --- xử lý ảnh ---
    images = []
    vision_texts = []
    image_tags = []
    for key in request.files:
        file = request.files[key]
        if file.filename:
            ext = os.path.splitext(file.filename)[-1]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = secure_filename(f"{timestamp}_guest_{key}{ext}")
            save_path = os.path.join('static/images/uploads', filename)
            file.save(save_path)
            image_url = f"/static/images/uploads/{filename}"
            image_tags.append(f"<img src='{image_url}' style='max-width:200px; border-radius:8px; margin-top:8px;'>")
            vision_output = extract_with_gpt_vision(save_path, user_request=message)
            vision_texts.append(f"[Ảnh: {filename}]\n{vision_output}")

    if vision_texts:
        if image_tags:
            message += "<br>" + "<br>".join(image_tags)
        vision_context = "\n\n".join(vision_texts)
    else:
        vision_context = ""

    try:
        if vision_context:
            message_for_gpt = message + "\n\n# Gợi ý từ ảnh:\n" + vision_context
        else:
            message_for_gpt = message

        reply = call_gpt_lite(message_for_gpt, history)
    except Exception as e:
        return jsonify({"reply": "⚠️ Hệ thống quá tải, bạn hãy thử lại trong ít phút nữa hoặc tạo đoạn chat mới."})

    if vision_texts:
        reply += "\n\nCó thể một vài chi tiết trong ảnh hơi mờ nhạt hoặc sai. Mong bạn kiểm tra lại giúp mình nha."

    history.append({"role": "user", "content": message})
    history.append({"role": "assistant", "content": reply})
    user_id = session.get("user_id")
    if "chat_session_id" not in session:
        session["chat_session_id"] = str(uuid.uuid4())
    session_id = session["chat_session_id"]
    save_chat_sql(user_id, session_id, history)
    update_daily_usage(is_slv=False)
    # --- trừ lượt nếu chưa xác thực ---
    if not user.is_verified:
        vip_dt = to_dt(user.vip_until_lite)
        if not vip_dt or now_vn() > vip_dt:
            user.free_gpt_uses = max(0, (user.free_gpt_uses or 0) - 1)
            db.session.commit()
    user.pending_messages = 0
    db.session.commit()
    return jsonify({"reply": reply})

#AI LITE
def generate_chat_lite_slug():
    return str(uuid.uuid4())
@app.route("/chat_lite", methods=["GET"])
def redirect_chat_lite():
    slug = generate_chat_lite_slug()
    return redirect(f"/chatLiteG9yZ/{slug}")    
@app.route("/chatLiteG9yZ/<slug>", methods=["GET"])
def chat_lite(slug):
    if is_maintenance("chat_lite"):
        return render_template("maintenance.html")

    username = session.get("username")
    user_id = session.get("user_id")

    if not username or not user_id:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    # 🛡️ Nếu chưa xác thực và đã hết lượt miễn phí → bắt xác thực
    if not user.is_verified and (user.free_gpt_uses or 0) <= 0:
        return redirect("/verify-otp")

    # ❌ Nếu có gói GPT → redirect
    now = now_vn()
    def valid(s):
        try:
            return datetime.strptime(s, "%Y-%m-%d %H:%M:%S") > now
        except:
            return False

    if user.vip_gpt_ai and valid(user.vip_until_gpt or ""):
        return redirect("/chat_redirect")

    # ✅ Nếu không có chat_session_id thì tạo session mới
    if "chat_session_id" not in session:
        session["chat_session_id"] = str(uuid.uuid4())
    session_id = session["chat_session_id"]

    # ✅ Kiểm tra xem session_id có trong DB chưa, nếu chưa thì tạo mới
    existing_session = ChatSession.query.filter_by(id=session_id).first()
    if not existing_session:
        new_session = ChatSession(
            id=session_id,
            user_id=user.user_id,
            created_at = dt.datetime.now(dt.timezone.utc),
            title="Đoạn chat mới"
        )
        db.session.add(new_session)
        db.session.commit()

    # ✅ Lấy lịch sử chat từ DB
    chat_history_objs = ChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=session_id
    ).order_by(ChatHistory.id.asc()).all()

    chat_history = [
        {"role": msg.role, "content": msg.content}
        for msg in chat_history_objs
    ]
    if chat_history and chat_history[-1]["role"] == "assistant":
        chat_history = chat_history[:-1]

    return render_template(
        "chat_ai_lite.html",
        user_id=user_id,
        username=username,
        user_lite_exhausted=False,
        is_verified=user.is_verified,
        chat_history=chat_history,
        chat_id=session_id,
        chat_title="Đoạn chat mới",
        is_maintenance=is_maintenance("chat_lite")
    )

#RESET LUỢT MỖI NGÀY
@app.route("/chat_ai_lite", methods=["GET"])
def chat_ai_lite_page():
    if is_maintenance("chat_ai_lite_daily"):
        return render_template("maintenance.html")

    username = session.get("username")
    user_id = session.get("user_id")

    if not username or not user_id:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    if not user.is_verified and (user.free_gpt_uses or 0) <= 0:
        return redirect("/verify-otp")

    now = now_vn()
    def valid(date_str):
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
            try:
                return datetime.strptime(date_str, fmt) > now
            except:
                continue
        return False
    if (user.vip_gpt_ai and valid(user.vip_until_gpt or "")) or \
       (user.vip_ai_lite and valid(user.vip_until_lite or "")):
        return redirect("/chat_redirect")

    # 🗂 Lịch sử đoạn chat
    chat_history = []
    chat_file = os.path.join("chat_history", f"{user_id}.json")
    if os.path.exists(chat_file):
        try:
            with open(chat_file, "r", encoding="utf-8") as f:
                chat_history = json.load(f)
        except Exception as e:
            print(f"[💥] Không đọc được chat_history: {e}")

    return render_template(
        "chat_ai_lite_daily.html",
        user_id=user_id,
        username=username,
        user_lite_exhausted=False,
        is_verified=user.       verified,
        chat_history=chat_history,
        chat_id=None,
        chat_title="",
        is_maintenance=is_maintenance("chat_ai_lite")
    )


@app.route("/chat_redirect", methods=["GET"])
def chat_redirect():
    if is_maintenance("chat_ai_lite_daily") or \
       is_maintenance("chat_ai_lite") or \
       is_maintenance("chat_lite") or \
       is_maintenance("gpt_chat"):
        return render_template("maintenance.html")

    username = session.get("username")
    if not username:
        return redirect("/login")
    
    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")
    if getattr(user, "is_blocked", False):
        session.clear()
        error = ("🚫 Tài khoản của bạn đã bị khoá. "
                 "Nếu đây là nhầm lẫn, vui lòng gửi khiếu nại "
                 "tại <a href='/appeal' style='color:#4ea6ff;'>đây</a>.")
        return render_template("login.html", error=error)
    now = now_vn()

    def valid(dt):
        try:
            if isinstance(dt, str):
                dt = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
            return dt > now
        except:
            return False

    # Gói GPT còn hạn
    if user.vip_gpt and valid(user.vip_until_gpt):
        return redirect("/solvervietAIwatch")  # gpt_chat.html

    # Gói AI Lite còn hạn
    if user.vip_ai_lite and valid(user.vip_until_lite):
        slug = generate_chat_lite_slug()
        return redirect(f"/chatLiteG9yZ/{slug}")

    # Tài khoản mới tạo (15 lượt free)
    if (user.free_gpt_uses or 0) > 0:
        return render_template("chat_ai_15.html", chat_history=[])

    # Đã xác thực → Free 5 lượt mỗi ngày
    if user.is_verified:
        return render_template("chat_ai_lite_daily.html", chat_history=[])

    # Chưa xác thực, hết free → Bắt xác thực
    flash("📩 Bạn cần xác thực email để tiếp tục sử dụng AI.")
    return redirect("/verify-otp")


#GỬI TIN TỪ HOME QUA AI VÀ TỰ KIỂM TRA THÔNG MINH
@app.route("/get-user-package")
def get_user_package():
    username = session.get("username")
    if not username:
        return jsonify({"status": "not_logged_in"})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"status": "not_logged_in"})

    now = now_vn()

    # Kiểm tra VIP GPT
    is_vip_gpt = False
    if user.vip_gpt:
        try:
            vip_until = datetime.strptime(user.vip_until_gpt, "%Y-%m-%d %H:%M:%S")
            is_vip_gpt = now <= vip_until
        except:
            pass

    # Kiểm tra gói AI Lite
    is_lite = False
    if user.vip_ai_lite:
        try:
            lite_until = datetime.strptime(user.vip_until_lite, "%Y-%m-%d %H:%M:%S")
            is_lite = now <= lite_until
        except:
            pass

    return jsonify({
        "status": "ok",
        "is_vip_gpt": is_vip_gpt,
        "is_lite": is_lite,
        "verified": user.is_verified
    })


#AI VẼ HÌNH
def generate_ai_prompt(user_text):
    return (
        f"Tạo hình ảnh theo mô tả: {user_text}. "
        "Ảnh nên rõ ràng, trình bày đẹp, không thêm chi tiết dư thừa, phù hợp sách vở hoặc minh hoạ học thuật."
    )
def translate_to_english(text):
    try:
        client = create_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional translator. Translate Vietnamese prompts into natural English for AI image generation. Only return the English text."
                },
                {"role": "user", "content": text}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return text

def enhance_prompt_style(eng_prompt):
    styles = [
        "highly detailed",
        "ultra realistic",
        "cinematic lighting",
        "4K resolution",
        "soft glow",
        "beautiful composition",
        "emotional atmosphere"
    ]
    return f"{eng_prompt}, {', '.join(styles)}"
def get_final_prompt(prompt_vn):
    rewritten = rewrite_prompt_for_image(prompt_vn)
    # 🧠 Không cần translate nữa vì GPT đã viết sẵn prompt tiếng Anh
    enhanced = enhance_prompt_style(rewritten)
    return enhanced

def generate_image_from_user_input(prompt_vn):
    try:
        final_prompt = get_final_prompt(prompt_vn)
        return generate_image_from_prompt(final_prompt)
    except Exception as e:
        return None
@app.route("/generate_image_from_text", methods=["POST"])
def draw_math_figure():
    data = request.json
    user_input = data.get("text", "").strip()

    if not user_input:
        return jsonify({"error": "Bạn chưa nhập nội dung cần vẽ."})

    # ✅ Kiểm tra quota nếu đã đăng nhập
    if current_user.is_authenticated:
        if not check_and_update_quota(current_user):
            return jsonify({
                "error": " Hệ thống đang tạm ngưng xử lý do lưu lượng truy cập cao từ tài khoản của bạn. Vui lòng quay lại sau 24 giờ.<br>"
                    " System temporarily paused your access due to high activity. Please try again later."
            })

    try:
        img_url = generate_image_from_user_input(user_input)

        # ✅ Lưu lịch sử nếu user đã đăng nhập
        if current_user.is_authenticated:
            log = ImageHistory(
                user_id=current_user.user_id,
                prompt=user_input,
                image_url=img_url
            )
            db.session.add(log)
            db.session.commit()

        return jsonify({"img_url": img_url, "source": "ai"})

    except Exception as e:
        return jsonify({"error": "Không tạo được hình từ AI. Vui lòng thử lại."})

@app.route("/generate_image", methods=["POST"])
def generate_image():
    username = session.get("username")
    if not username:
        return jsonify({"error": "🔒 Bạn chưa đăng nhập."})

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"error": "🔒 Không tìm thấy tài khoản."})

    cap_nhat_trang_thai_vip(user)
    db.session.commit()

    # ✅ Chặn nếu không phải gói GPT AI
    if username != "admin":
        if not user.vip_gpt_ai:
            return jsonify({"error": "⚠️ Gói của bạn không hỗ trợ tạo ảnh. Hãy nâng cấp lên gói SLV."})
        
        # ✅ GỌI HÀM KIỂM TRA QUOTA GPT
        if not check_and_update_quota(user):
            return jsonify({
                "error": " Hệ thống đang tạm ngưng xử lý do lưu lượng truy cập cao từ tài khoản của bạn. Vui lòng quay lại sau 24 giờ.<br>"
                    " System temporarily paused your access due to high activity. Please try again later."
            })
        
        # ✅ Reset quota ảnh nếu sang ngày mới
        now = datetime.now()
        if not user.image_quota_reset_at or user.image_quota_reset_at.date() < now.date():
            user.image_quota_today = 10
            user.image_quota_reset_at = now
            db.session.commit()

        # ✅ Chặn nếu hết lượt tạo ảnh
        if user.image_quota_today <= 0:
            return jsonify({"error": "The image generation system seems temporarily overloaded. Please try again later."})

    # ✅ Kiểm tra prompt đầu vào
    prompt = request.json.get("prompt", "").strip()
    if not prompt:
        return jsonify({"error": "❌ Bạn chưa nhập nội dung hình muốn tạo."})

    try:
        # ✨ Rewrite prompt bằng GPT để ra prompt đẹp
        final_prompt = rewrite_prompt_for_image(prompt)

        if not final_prompt:
            return jsonify({
                "error": "⚠️ Hình ảnh bạn yêu cầu không thể tạo do vi phạm điều khoản nội dung. Vui lòng thử lại với mô tả khác không đề cập đến khuôn mặt thật, nhận diện cá nhân hoặc nội dung nhạy cảm."
            })

        # ✅ Trừ lượt tạo ảnh
        if username != "admin":
            user.image_quota_today -= 1
            db.session.commit()

        # ✅ Tạo ảnh
        img_url = generate_image_from_prompt(final_prompt)

        if not img_url:
            return jsonify({"error": "⚠️ Không thể tạo ảnh. Vui lòng thử lại sau."})

        return jsonify({"img_url": img_url})

    except Exception as e:
        return jsonify({"error": "⚠️ Không thể tạo ảnh. Vui lòng thử lại sau."})

@app.route("/save_generated_image_log", methods=["POST"])
def save_generated_image_log():
    user_id = session.get("user_uuid")
    if not user_id:
        return "", 204

    data = request.get_json()
    prompt = data.get("prompt", "")
    image_url = data.get("image_url", "")
    message_index = data.get("message_index")  

    if prompt and image_url:
        log = ImageHistory(
            user_id=user_id,
            prompt=prompt,
            image_url=image_url,
            message_index=message_index 
        )
        db.session.add(log)
        db.session.commit()

    return "", 204

@app.route("/get_image_history")
def get_image_history():
    if "user_id" not in session:
        return jsonify([])

    image_logs = ImageHistory.query.filter_by(user_id=session["user_id"]).order_by(ImageHistory.created_at.desc()).all()
    result = [
        {
            "prompt": log.prompt,
            "image_url": log.image_url,
            "created_at": log.created_at.isoformat(),
            "message_index": log.message_index  # 👈 thêm vào response
        }
        for log in image_logs
    ]
    return jsonify(result)
def save_generated_image_log_backend(user_id, image_url, prompt, message_index=None, source="gpt_chat"):
    try:
        log = ImageHistory(
            user_id=user_id,
            image_url=image_url,
            prompt=prompt,
            source=source,
            message_index=message_index, 
            created_at=datetime.utcnow()
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        pass

#GỬI FILE
@app.route("/upload_file", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"reply": "Không có file được gửi."}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"reply": "File không hợp lệ."}), 400

    filename = secure_filename(file.filename)
    user_id = session.get("user_id", "guest")
    existing_path = os.path.join("static", "images", "uploads", user_id, filename)

    if not os.path.exists(existing_path):
        return jsonify({"reply": "⚠️ File không tồn tại hoặc đã bị xoá."}), 400

    os.utime(existing_path, None)

    return jsonify({"reply": f"✅ Đã cập nhật ảnh: {filename}"})

#MỞ ÂM BUM
@app.route("/get_user_album")
def get_user_album():
    user_id = session.get("user_id", "guest")
    upload_dir = os.path.join("static", "images", "uploads", str(user_id))

    now = time.time()
    max_age = 7 * 86400  

    images = []
    if os.path.exists(upload_dir):
        for filename in os.listdir(upload_dir):
            if filename.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
                path = os.path.join(upload_dir, filename)
                mtime = os.path.getmtime(path)
                age = now - mtime
                if age <= max_age:
                    images.append({
                        "path": f"/{path.replace(os.sep, '/')}"
                    })

    return jsonify({"images": images})
#LỊCH SỬ ĐOẠN CHAT
from models import ChatHistory
from sqlalchemy.exc import SQLAlchemyError

@app.route("/save_chat", methods=["POST"])
def save_chat_route():
    user_id = session.get("user_id")
    session_id = session.get("chat_session_id")

    if not user_id or not session_id:
        return jsonify({"error": "Chưa đăng nhập hoặc chưa có session"}), 401

    data = request.get_json()
    messages = data.get("messages", [])

    try:
        ChatHistory.query.filter_by(user_id=user_id, session_id=session_id).delete()

        for msg in messages:
            db.session.add(ChatHistory(
                user_id=user_id,
                session_id=session_id,
                role=msg.get("role"),
                content=msg.get("content"),
                timestamp=datetime.utcnow()
            ))

        db.session.commit()
        return jsonify({"status": "saved"})
    except SQLAlchemyError as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route("/load_chat/<chat_id>")
def load_chat(chat_id):
    user_id = session.get("user_id", "guest")
    if user_id == "guest":
        return jsonify({"error": "Bạn chưa đăng nhập."}), 401

    user = User.query.filter_by(user_id=user_id).first()

    if not user:
        return jsonify({"error": "Không tìm thấy người dùng."}), 404

    user_uuid = user.user_id  # dùng UUID string

    # Truy vấn lịch sử chat theo session_id (chat_id)
    messages = ChatHistory.query.filter_by(
        user_id=user_uuid,
        session_id=chat_id
    ).order_by(ChatHistory.timestamp).all()

    if not messages:
        return jsonify({"error": "Không tìm thấy đoạn chat."}), 404

    return jsonify({
        "chat": [
            {"role": msg.role, "content": msg.content}
            for msg in messages
        ]
    })

@app.route("/reset_current_chat")
def reset_current_chat():
    session.pop("current_chat_id", None)
    return jsonify({"success": True})

import shutil
import os
@app.route("/new_chat", methods=["POST"])
def new_chat():
    try:
        user_id = session.get("user_id")
        if not user_id:
            return jsonify({"success": False, "error": "Không có user_id trong session"})

        user = User.query.filter_by(user_id=user_id).first()
        if not user:
            return jsonify({"success": False, "error": "Không tìm thấy người dùng"})

        current_session_id = session.get("chat_session_id")

        # ✅ Kiểm tra nếu đoạn chat hiện tại chưa lưu → xóa
        is_saved = ChatHistory.query.filter_by(
            session_id=current_session_id,
            is_saved=True
        ).first()

        if not is_saved:
            # 🧹 Dọn sạch lịch sử chưa lưu (KHÔNG xóa ảnh nữa)
            ChatHistory.query.filter_by(user_id=user_id, is_saved=False).delete()
            ChatSession.query.filter_by(user_id=user_id).delete()

        # ✅ Tạo phiên mới
        new_session_id = str(uuid.uuid4())
        new_slug = str(uuid.uuid4())

        session["chat_session_id"] = new_session_id
        session["slug_to_session"] = {
            "slug": new_slug,
            "session_id": new_session_id
        }

        new_session = ChatSession(
            id=new_session_id,
            user_id=user_id,
            created_at=datetime.utcnow(),
            title="Đoạn chat mới"
        )
        db.session.add(new_session)
        db.session.commit()

        return jsonify({"success": True, "session_id": new_session_id, "slug": new_slug})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


#TƯƠNG TÁC VỚI NGƯỜI DÙNG
@app.route('/send-emoji', methods=['POST'])
def handle_emoji():
    if "username" not in session:
        return jsonify({"success": False, "ai_reply": "Bạn chưa đăng nhập."})

    data = request.get_json()
    emoji = data.get('emoji')
    last_reply = data.get('last_reply', '').strip()

    # Không phản hồi nếu câu trước lỗi
    if any(w in last_reply for w in ["⚠️", "🚫", "❌", "System temporarily paused", "Vui lòng quay lại sau", "Error: The system is currently overloaded."]):
        return jsonify({"ai_reply": None})

    username = session["username"]
    user_id = session.get("user_id", "guest")
    session_id = session.get("chat_session_id")
    if not session_id:
        session_id = str(uuid.uuid4())
        session["chat_session_id"] = session_id

    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({"success": False, "ai_reply": "Không tìm thấy người dùng."})

    # Trừ lượt nếu cần
    cap_nhat_trang_thai_vip(user)
    reset_usage_if_new_day(user)
    if not check_and_update_quota(user):
        return jsonify({"success": False, "ai_reply": "⚠️Bạn đã đạt giới hạn sử dụng hôm nay. Quay lại sau 24h nhé!"})

    # Chuẩn bị prompt
    ai_personality = user.ai_personality or "bình thường"
    emotion_prompts = {
        "❤️": "Người dùng rất thích câu trả lời. Hãy phản hồi nhẹ nhàng, tích cực và tiếp tục mạch hội thoại.",
        "😂": "Người dùng thấy vui vẻ. Hãy đáp lại hài hước hoặc gần gũi hơn.",
        "😢": "Người dùng hơi buồn. Hãy động viên nhẹ nhàng và hỏi xem bạn có thể giúp gì.",
        "🤔": "Người dùng đang suy nghĩ. Hãy hỏi xem có chỗ nào cần giải thích rõ hơn.",
        "😡": "Người dùng chưa hài lòng. Hãy xin lỗi lịch sự và mời họ nêu vấn đề cụ thể."
    }
    emotion_context = emotion_prompts.get(emoji, "Người dùng vừa thả cảm xúc. Hãy phản hồi phù hợp.")

    prompt = f"""Bạn là một trợ lý AI mang cá tính {ai_personality}, đang tiếp tục cuộc trò chuyện với người dùng.

Câu trả lời trước của bạn:
\"{last_reply}\"

Người dùng vừa thả cảm xúc: {emoji}
→ {emotion_context}

⚠️ Yêu cầu định dạng phản hồi:
- BẮT BUỘC mở đầu bằng câu: "Cảm ơn bạn đã thả {emoji}" (hoặc biến thể tự nhiên, nhưng vẫn phải có emoji này trong dòng đầu tiên).
- Sau câu mở đầu đó, bạn phản hồi phù hợp, ấm áp, dễ đọc.
- Nếu câu dài hơn 20 từ, hãy xuống dòng bằng thẻ HTML `<br>`.
- KHÔNG dùng markdown. Chỉ dùng `<br>` nếu cần.

Giữ phong cách thân thiện và dễ hiểu.
"""
    client = create_openai_client()
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    reply = response.choices[0].message.content.strip()
    reply = auto_insert_break_if_needed(reply)
    history = [
        {"role": "user", "content": f"[Thả cảm xúc {emoji}] → {last_reply}"},
        {"role": "assistant", "content": reply}
    ]
    save_chat_sql(
        user_id=user.user_id,
        session_id=session_id,
        history=history
    )
    increase_gpt_usage(user)
    db.session.commit()

    return jsonify({"success": True, "ai_reply": reply})

def auto_insert_break_if_needed(text):
    if "<br>" in text:
        return text
    sentences = text.split(". ")
    return "<br>".join(s.strip() for s in sentences if s)
@app.route("/upload_file_to_ai", methods=["POST"])
def upload_file_to_ai():
    file = request.files.get("file")
    if file:
        filename = secure_filename(file.filename)
        save_path = os.path.join("static/images/uploads", filename)
        file.save(save_path)
        return jsonify({"success": True, "filename": save_path})
    return jsonify({"success": False})

UPLOAD_FOLDER = "static/uploads_feedback"
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
def generate_gopy_slug():
    return str(uuid.uuid4())
@app.route("/gop-y", methods=["GET"])
def redirect_gopy():
    slug = generate_gopy_slug()
    return redirect(f"/gopyCypherZonerp/{slug}")    
@app.route("/gopyCypherZonerp/<slug>", methods=["GET", "POST"])
def gop_y(slug):
    if request.method == "POST":
        image_paths = []

        images = request.files.getlist("images")
        if images and any(img.filename != "" for img in images):
            for image in images:
                if image and allowed_file(image.filename):
                    filename = secure_filename(image.filename)
                    timestamp = str(int(time.time()))
                    full_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{timestamp}_{filename}")
                    image.save(full_path)
                    image_paths.append("/" + full_path.replace("\\", "/"))

        full_name = request.form.get("full_name", "Ẩn danh")
        user_email = request.form.get("user_email", "")
        user_id = request.form.get("user_id", "")
        message = request.form.get("message", "")
        username = session.get("username", "Khách")
        category = request.form.get("category", "")
        old_email = request.form.get("old_email_optional", "").strip()
        new_email = request.form.get("new_email", "").strip()

        with open("feedback_log.txt", "a", encoding="utf-8") as f:
            vn_time = datetime.utcnow() + timedelta(hours=7)
            f.write(f"Thời gian: [{vn_time.strftime('%Y-%m-%d %H:%M:%S')}]\n")
            f.write(f"Họ tên: {full_name}\n")
            f.write(f"Email: {user_email}\n")
            f.write(f"Username: {username}\n")
            f.write(f"User ID: {user_id}\n")
            f.write(f"Phân loại: {category}\n")

            if category == "Yêu cầu đổi email":
                f.write(f"Email hiện tại: {old_email}\n")
                f.write(f"Email mới muốn đổi: {new_email}\n")
            else:
                f.write(f"Nội dung: {message}\n")

            if image_paths:
                f.write(f"Ảnh: {' | '.join(image_paths)}\n")
            f.write("---\n\n")


        session["gopy_success"] = True
        session["gopy_category"] = category
        return redirect("/gop-y")

    success = session.pop("gopy_success", False)
    category = session.pop("gopy_category", "")
    return render_template("gop_y.html", success=success, category=category, user_id=session.get("user_id", ""))

#Giao diện nhận góp ý 
def extract(block, prefix):
    for line in block.splitlines():
        if line.strip().startswith(prefix):
            return line.replace(prefix, "").strip()
    return ""

def extract_loose(block, prefix):
    for line in block.splitlines():
        if prefix in line:
            return line.split(prefix, 1)[1].strip()
    return ""
@app.route("/admin/gui-thong-bao")
@admin_only
def gui_thong_bao():
    return render_template("gui_thong_bao.html")
@app.route("/admin/gop-y")
@admin_only
def admin_gopy():
    if not session.get("is_admin"):
        return redirect("/admin_login")

    entries = []
    try:
        with open("feedback_log.txt", "r", encoding="utf-8") as f:
            blocks = f.read().split("---\n\n")
            for idx, block in enumerate(blocks, 1):
                if block.strip():
                    entry = {
                            "index": idx,
                            "time": extract(block, "Thời gian:"),
                            "name": extract(block, "Họ tên:"),
                            "email": extract_loose(block, "Email:"),
                            "user_id": extract(block, "User ID:"),
                            "username": extract_loose(block, "Username:"),
                            "message": extract(block, "Nội dung:"),
                            "image_paths": extract(block, "Ảnh:").split(" | ") if extract(block, "Ảnh:") else [],
                            "type": extract(block, "Phân loại:"),
                            "old_email": extract(block, "Email hiện tại:"),
                            "new_email": extract(block, "Email mới muốn đổi:")
                        }

                    entries.append(entry)
    except FileNotFoundError:
        pass

    return render_template("admin_gopy.html", feedback_entries=entries)

@app.route("/bo-qua", methods=["POST"])
@admin_only
def bo_qua():
    try:
        entry_index = int(request.form.get("entry_index"))
    except (TypeError, ValueError):
        return "Chỉ số không hợp lệ", 400

    try:
        with open("feedback_log.txt", "r", encoding="utf-8") as f:
            blocks = f.read().split("---\n\n")
    except FileNotFoundError:
        blocks = []

    if 0 <= entry_index - 1 < len(blocks):
        del blocks[entry_index - 1]

        with open("feedback_log.txt", "w", encoding="utf-8") as f:
            f.write("---\n\n".join(blocks).strip() + "\n")

    return redirect("/admin/gop-y")

#Bảo mật tối thượng
# ====== BẢO MẬT MỞ CỔNG ADMIN ======

TRUSTED_IP = os.getenv("TRUSTED_IP")
BACKDOOR_CODE = os.getenv("BACKDOOR_CODE")

def is_trusted_ip():
    return request.remote_addr == TRUSTED_IP



@app.route("/feedback")
def feedback_redirect():
    return redirect("/gop-y")  # nếu trang thật là /gop-y
#GỬI KHIẾU NẠI CHO USER KHÔNG ĐĂNG NHẬP
@app.route("/appeal", methods=["GET", "POST"])
def appeal():
    if request.method == "POST":
        full_name = request.form.get("full_name") or "Ẩn danh"
        user_id = request.form.get("user_id", "")
        user_email = request.form.get("email", "")
        message = request.form.get("message", "")
        category = request.form.get("category", "Khiếu nại tài khoản")

        # ảnh
        saved_paths = []
        if "images" in request.files:
            images = request.files.getlist("images")
            for img in images:
                if img and img.filename:
                    filename = f"{uuid.uuid4().hex}_{secure_filename(img.filename)}"
                    folder = "static/images/feedback"
                    os.makedirs(folder, exist_ok=True)
                    path = os.path.join(folder, filename)
                    img.save(path)
                    saved_paths.append(f"/{path.replace(os.sep, '/')}")

        vn_time = datetime.utcnow() + timedelta(hours=7)
        with open("feedback_log.txt", "a", encoding="utf-8") as f:
            f.write(f"Thời gian: [{vn_time.strftime('%Y-%m-%d %H:%M:%S')}]\n")
            f.write(f"Họ tên: {full_name}\n")
            f.write(f"Email: {user_email}\n")
            f.write(f"Username: (không có)\n")
            f.write(f"User ID: {user_id}\n")
            f.write(f"Phân loại: {category}\n")
            f.write(f"Nội dung: {message}\n")
            if saved_paths:
                f.write(f"Ảnh: {' | '.join(saved_paths)}\n")
            f.write("---\n\n")

        return render_template("appeal.html", success=True)
    return render_template("appeal.html")


with app.app_context():
    
    db.create_all() 
@socketio.on("join")
def handle_join(username):
    join_room(f"user_{username}")
    
@socketio.on("join_room")
def handle_join_room(data):
    room = data.get("room")
    if room:
        join_room(room)
@socketio.on("react_message")
def handle_react_message(data):
    """
    data = {
        "msg_id": <id của tin>,
        "emoji": "❤️",      # hoặc "" nếu là bỏ cảm xúc
        "target": "luffy20" # username của người nhận
    }
    """

    # Gửi về cả 2 phía (người gửi và người nhận)
    emit("react_update", data, room=f"user_{data['target']}")
    emit("react_update", data, room=request.sid)  # chính người gửi
    # routes_admin_seed.py  (import đâu đó trong app của bạn là xong)
from flask import request, jsonify
import os, pyotp

@app.route("/verify_totp", methods=["POST"])
def verify_totp():
    data = request.json
    code = data.get("code")

    secret = os.getenv("ADMIN_2FA_SECRET")
    totp = pyotp.TOTP(secret)

    if totp.verify(code):
        return jsonify({"success": True})
    else:
        return jsonify({"success": False}), 401
#ĐÁNH GIÁ AI
@app.route("/submit_feedback", methods=["POST"])
def submit_feedback():
    user_id = session.get("user_id")
    session_id = session.get("chat_session_id")
    data = request.json

    feedback = Feedback(
        user_id=user_id,
        session_id=session_id,
        stars=data.get("stars"),
        comment=data.get("comment"),
        ai_personality=data.get("ai_personality"),
        package_type=data.get("package_type")
    )
    db.session.add(feedback)
    db.session.commit()

    return jsonify({"success": True})
@app.route("/api/home_info")
def api_home_info():
    if "username" not in session:
        return jsonify({
            "success": True,
            "guest": True,
            "vip_status": "Guest",
            "message": "Chưa đăng nhập"
        })

    user = User.query.filter_by(username=session["username"]).first()
    if not user:
        return jsonify({"success": False, "message": "Không tìm thấy user"})

    if user.is_blocked:
        return jsonify({"success": False, "blocked": True, "message": "Tài khoản bị khóa"})

    now = datetime.utcnow()

    return jsonify({
        "success": True,
        "guest": False,
        "username": user.username,
        "fullname": user.fullname or user.username,
        "avatar_url": user.avatar_url or "/static/logos/logo.png",

        "vip_gpt": bool(user.vip_gpt_ai and user.vip_until_gpt and user.vip_until_gpt > now),
        "vip_lite": bool(user.vip_ai_lite and user.vip_until_lite and user.vip_until_lite > now),
        "vip_until_gpt": user.vip_until_gpt.isoformat() if user.vip_until_gpt else "",
        "vip_until_lite": user.vip_until_lite.isoformat() if user.vip_until_lite else "",

        "vip_status": (
            "SLV" if user.vip_gpt_ai and user.vip_until_gpt and user.vip_until_gpt > now
            else "Lite" if user.vip_ai_lite and user.vip_until_lite and user.vip_until_lite > now
            else "Free"
        ),

        "messages": [],
        "unread_senders": get_unread_senders(user.username),
        "has_seen_intro": user.has_seen_intro
    })


def get_unread_senders(username):
    unread = (
        db.session.query(distinct(Message.sender))
        .filter(Message.receiver == username, Message.read == False)
        .all()
    )
    return [s[0] for s in unread if s[0] != username]

@app.route("/save_dev_history", methods=["POST"])
def save_dev_history():
    if not session.get("dev_mode") or "user_id" not in session:
        return jsonify({"success": False, "error": "Dev mode chưa bật hoặc chưa đăng nhập."}), 403

    data = request.get_json()
    messages = data.get("messages", [])  # đoạn chat mới gửi lên
    user_id = session["user_id"]
    session_id = session.get("chat_session_id")

    if not session_id or not isinstance(messages, list):
        return jsonify({"success": False, "error": "Dữ liệu không hợp lệ."}), 400

    existing = DevChatHistory.query.filter_by(user_id=user_id, session_id=session_id).first()
    if existing:
        existing.history = messages  # ✅ Ghi đè toàn bộ, vì JS đã gửi đầy đủ
        db.session.commit()
    else:
        new = DevChatHistory(user_id=user_id, session_id=session_id, history=messages)
        db.session.add(new)
        db.session.commit()

    return jsonify({"success": True})

@app.route("/dev_chat_history", methods=["POST"])
def get_dev_chat_history():
    from flask import session, jsonify, request

    username = session.get("username")
    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify([])

    data = request.get_json() or {}
    session_id = data.get("session_id") or session.get("chat_session_id")


    if not session_id:
        return jsonify([])

    dev_chat = DevChatHistory.query.filter_by(
        user_id=user.user_id,
        session_id=session_id
    ).order_by(DevChatHistory.created_at.desc()).first()

    return jsonify(dev_chat.history if dev_chat and dev_chat.history else [])
import pytz

VIETNAM_TZ = pytz.timezone("Asia/Ho_Chi_Minh")  # 👈 thêm dòng này

@app.route("/save_user_memory", methods=["POST"])
@login_required
def save_user_memory():
    data = request.get_json()
    category = data.get("category")
    content = data.get("content")
    password = data.get("password")  # chỉ dùng cho Bảo mật cá nhân

    if not category or not content:
        return jsonify({"error": "Missing category or content"}), 400

    encrypted_password = None
    if category == "Bảo mật cá nhân":
        if not password:
            return jsonify({"error": "Missing password for sensitive info"}), 400
        hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        encrypted_password = hashed.decode('utf-8')

    existing = UserMemoryItem.query.filter_by(user_id=current_user.user_id, category=category).first()
    if existing:
        existing.content = content
        existing.updated_at = datetime.now(VIETNAM_TZ)  # ✅ cập nhật múi giờ chuẩn
        if encrypted_password:
            existing.encrypted_password = encrypted_password
    else:
        memory = UserMemoryItem(
            user_id=current_user.user_id,
            category=category,
            content=content,
            encrypted_password=encrypted_password,
            created_at=datetime.now(VIETNAM_TZ),     # ✅ nếu muốn đồng bộ luôn created_at
            updated_at=datetime.now(VIETNAM_TZ)
        )
        db.session.add(memory)

    db.session.commit()
    return jsonify({"status": "saved"})


@app.route("/get_user_memory", methods=["POST"])
@login_required
def get_user_memory():
    data = request.get_json()
    category = data.get("category")

    memory = UserMemoryItem.query.filter_by(user_id=current_user.user_id, category=category)\
                                 .order_by(UserMemoryItem.id.desc()).first()

    last_updated = (
        memory.updated_at.astimezone(VIETNAM_TZ).strftime("%H:%M %d/%m/%Y")
        if memory and memory.updated_at else None
    )

    return jsonify({
        "content": memory.content if memory else "",
        "last_updated": last_updated
    })

@app.route("/dieu-khoan")
def dieu_khoan():
    return render_template("terms_of_use.html")

from flask import request, send_file, render_template, session
from models.smartdoc import SmartDoc
from models.user import User
from docx import Document
from extensions import db
from openai_config import GPT4O_KEYS  
import openai, tempfile, random
def generate_smartdoc_slug():
    part1 = "0x" + secrets.token_hex(4)  # ví dụ: 0x12ab34cd
    part2 = secrets.token_hex(3)         # ví dụ: aabbcc
    part3 = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(8))  # ví dụ: Zx82Qw1L
    return f"{part1}-{part2}-?{part3}"

@app.route("/smartdoc", methods=["GET"])
def redirect_smartdoc():
    slug = generate_smartdoc_slug()
    return redirect(f"/smartdocXonesixseven/{slug}")

@app.route("/smartdocXonesixseven/<slug>", methods=["GET"])
def smartdoc_form(slug):
    username = session.get("username")
    if not username:
        return redirect("/login")

    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")

    has_slv = user.vip_until_gpt and user.vip_until_gpt > datetime.utcnow()

    return render_template("smartdoc.html", user_has_slv=has_slv)
# ==== SmartDoc Section Definitions ==== #
# X7.INP[01]: Giao diện nhập tiêu đề văn bản, có gợi ý datalist bên dưới (input + suggestions)
# X1.BTN[01]: Nút “Chỉnh sửa” nội dung văn bản ở phần preview (kích hoạt contentEditable)
# X7.BR[04]: Logic hiển thị preview HTML của văn bản từ nội dung GPT trả về (SLVIT_HTMLgen)
# X3.IMG[07]: Logic xử lý ảnh mẫu (preview ảnh và xoá ảnh mẫu đã chọn)
# X5.HIS[01]: Logic mở popup lịch sử (SmartDoc history popup)
# X5.HIS[02]: Logic tìm kiếm lịch sử theo từ khoá (SmartDoc search history)
# X5.HIS[03]: Giao diện popup lịch sử SmartDoc (HTML hiển thị lịch sử)
# X3.PRV[01]: Logic đóng phần xem trước văn bản (smartdoc-preview)
# X2.LOD[SOAN]: Giao diện overlay loading khi đang soạn văn bản
# X5.FORM[GEN-DOC]: Xử lý submit form SmartDoc – gọi API /generate_docx và hiển thị preview
# X9.DOM[RESET-LOAD]: Reset lại overlay nếu người dùng F5 lại trang
# X4B7: Ghi chú quyền hạn – Chỉ người dùng gói SLV mới dùng được SmartDoc
# X7.RSV[SUBMIT-BLOCKER.SLV/SMARTDOC]: Logic chặn nút “Tạo văn bản” nếu chưa nâng cấp gói SLV
# X9.GHOSTSTAT[FAKECOUNT-VIS-MOCK]: Tạo số lượng người dùng giả định trong ngày để hiển thị (fake usage count)
# X7.BR[05]: Logic chỉnh sửa nội dung và lưu lại sau khi chỉnh (Confirm Save Logic)
# X6.SLVTOAST: Logic hiển thị thông báo gọn nhẹ (toasts) theo trạng thái (success, error,...)
# X7.DL-REFRESH: Tự động reload lại trang sau khi người dùng tải file về (reset lại trạng thái)
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    selected_key = random.choice(GPT4O_KEYS) if GPT4O_KEYS else None
    if not selected_key:
        return "Không có API key GPT-4o khả dụng", 500

    client = OpenAI(api_key=selected_key)

    title = request.form.get("title", "").strip() or "Văn bản không tiêu đề"
    description = request.form.get("description", "")
    author_name = request.form.get("author_name", "").strip()
    mode = request.form.get("mode", "normal")
    username = session.get("username")
    user = User.query.filter_by(username=username).first() if username else None

    if author_name:
        description += f"\n\nTên người viết: {author_name}"

    # --- ✅ Kiểm tra quyền dùng SmartDoc ---
    if user:
        # Reset lượt nếu sang ngày mới
        if user.smartdoc_last_used_date != date.today():
            user.smartdoc_usage_today = 0
            user.smartdoc_last_used_date = date.today()

        has_slv = user.vip_until_gpt and user.vip_until_gpt > datetime.now()
        if not has_slv:
            return jsonify({
                "success": False,
                "error": "🚫 Bạn cần nâng cấp gói SLV để sử dụng tính năng SmartDoc."
            }), 403

        # Giới hạn 20 lượt/ngày
        if user.smartdoc_usage_today >= 20:
            return jsonify({
                "success": False,
                "error": "⚠️ Bạn đã dùng hết <b>20 lượt SmartDoc hôm nay</b>. Hãy quay lại vào ngày mai."
            }), 429

        # ✅ Được phép → tăng lượt
        user.smartdoc_usage_today += 1
        db.session.commit()

    # --- Xử lý tạo văn bản như bạn đã viết (phần dưới giữ nguyên) ---
    image_file = request.files.get("sample_image")
    if image_file:
        # --- Người dùng có gửi ảnh mẫu ---
        import base64
        image_bytes = image_file.read()
        encoded_image = base64.b64encode(image_bytes).decode()

        vision_prompt = (
                "Bạn là chuyên gia soạn thảo văn bản tiếng Việt.\n"
                "Ảnh phía dưới là mẫu văn bản có bố cục đẹp mắt.\n"
                "Hãy quan sát cách trình bày trong ảnh (gồm căn lề, căn giữa, gạch dưới, xuống dòng, tên in đậm, cách xưng hô, định dạng ngày tháng...)\n"
                "Sau đó viết lại 1 văn bản mới theo mẫu, **giữ nguyên phong cách trình bày**, nhưng nội dung thì dựa theo yêu cầu người dùng bên dưới.\n"
                "⚠️ Văn bản đầu ra cần là THUẦN VĂN BẢN (plain text), không dùng markdown hoặc HTML, chỉ gồm nội dung văn bản trình bày theo phong cách như ảnh."
                "❗ Tuyệt đối không copy nguyên văn từ ảnh mẫu, chỉ lấy phong cách trình bày và bố cục.\n"
                "📝 Yêu cầu người dùng:\n"
                f"{description or '(Không có mô tả)'}"
            )
        response = client.chat.completions.create(
            model="gpt-4o",
            max_tokens=3000,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": vision_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_image}"}}
                ]
            }]
        )
    else:
        # --- Không có ảnh mẫu → dùng mô tả văn bản & chế độ ---
        selected_instruction = build_instruction(mode)

        system_prompt = (
            "Bạn là chuyên gia hàng đầu về soạn thảo văn bản hành chính và văn học tiếng Việt.\n\n"
            "Trước khi viết, hãy tự động xác định thể loại văn bản thuộc 1 trong 2 nhóm sau:\n"
            "➤ Loại hành chính gồm nhiều dạng, và mỗi dạng có cấu trúc riêng. Ví dụ:\n"
            "- Đơn kiện hoặc đơn khởi tố:\n"
            "  + Quốc hiệu, tiêu đề, kính gửi\n"
            "  + Họ tên người viết, địa chỉ, lý do khởi kiện, nội dung cụ thể\n"
            "  + Yêu cầu pháp lý (đề nghị gì?)\n"
            "  + Cam kết chịu trách nhiệm, ngày tháng, chữ ký\n\n"
            "- Đơn xin nghỉ học, nghỉ phép:\n"
            "  + Quốc hiệu, tiêu đề, kính gửi\n"
            "  + Thông tin cá nhân, thời gian nghỉ, lý do xin nghỉ\n"
            "  + Lời cảm ơn và cam kết quay lại đúng hạn\n\n"
            "- Thông báo:\n"
            "  + Quốc hiệu, tiêu đề, kính gửi\n"
            "  + Nội dung cần thông báo (sự kiện gì, diễn ra ở đâu, khi nào...)\n"
            "  + Liên hệ hoặc thông tin bổ sung, ngày tháng, chữ ký\n\n"
            "➤ Loại học thuật: bài văn, cảm nhận, luận văn, bài học sinh...\n\n"
            "Sau đó hãy trình bày nội dung theo đúng chuẩn văn bản Word, tuyệt đối không ghi chú loại nào ở phần đầu.\n\n"
            "🌐 Nếu là văn bản hành chính:\n"
            "- Ghi rõ quốc hiệu:\n"
            "  CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
            "  ĐỘC LẬP - TỰ DO - HẠNH PHÚC\n\n"
            "- Tiêu đề và tên người viết có thể viết hoa để nhấn mạnh.\n"
            "- Mỗi đoạn cách nhau bằng MỘT dòng trống.\n"
            "- Không gộp nhiều dòng lại thành 1 khối xuống dòng bằng \\n.\n"
            "- Cần thể hiện các mục rõ ràng: Kính gửi, Lý do, Nội dung, Người làm đơn...\n"
            "- Cuối văn bản ghi rõ:\n"
            "  Ngày ..... tháng ..... năm .....\n"
            "  Người làm đơn\n"
            "  (Ký, ghi rõ họ tên)\n"
            "  Nguyễn Văn A\n\n"
            "📄 Nếu là bài văn, cảm nhận:\n"
            "- KHÔNG có quốc hiệu, không kính gửi, không ký tên.\n"
            "- Bắt đầu bằng dòng: ĐỀ BÀI: ...\n"
            "- Trình bày như bài văn mẫu trong SGK, chia đoạn hợp lý, rõ ràng.\n"
            "- Nếu trình bày là văn bản cảm nhận thì khỏi cần ghi rõ ngày tháng năm.\n\n"
            "✍️ QUY TẮC TRÌNH BÀY RẤT QUAN TRỌNG:\n"
            "- KHÔNG để 2 dòng trống liên tiếp giữa các đoạn.\n"
            "- Luôn viết ngày tháng theo format: ' ngày ... tháng ... năm ...' và đặt cuối văn bản.\n"
            "- Mỗi đoạn chỉ cách nhau đúng 1 dòng trống.\n"
            "- Không để đoạn nào bắt đầu bằng dòng trống.\n"
            "- Viết sát lề trái, không căn giữa toàn văn bản.\n"
            "- KHÔNG dùng markdown như **, __, ``, #...\n\n"
            "- Các dòng như 'Kính gửi:', 'Giáo viên chủ nhiệm lớp:', 'Tôi tên là:'... phải viết liền với phần nội dung trên cùng một dòng.\n"
            "- Trả lời đúng nội dung văn bản theo yêu cầu, KHÔNG được chèn markdown, HTML, hay ký tự đặc biệt nào như `plaintext`, ``` hoặc dấu *, \n\n"
            

            f"{selected_instruction}"
        )

        response = client.chat.completions.create(
            model="gpt-4o",
            max_tokens=800 if mode == "brief" else 1200 if mode == "normal" else 3000,
            temperature=0.6,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Tiêu đề: {title}\nTôi muốn tạo văn bản như sau:\n{description}"}
            ]
        )

    content = response.choices[0].message.content
    raw_content = response.choices[0].message.content
    content = clean_gpt_output(raw_content)
    new_doc = SmartDoc(
        user_id=user.user_id if user else None,
        description=description,
        content=content,
        file_type="docx",
        title=title,
        author_name=author_name,
        created_at=now_vn()
    )
    db.session.add(new_doc)
    db.session.commit()
    return jsonify({"success": True, "content": content, "doc_id": str(new_doc.doc_id)})
def build_instruction(mode):
    if mode == "brief":
        return """
✍️ Yêu cầu phong cách viết:
- Trình bày ngắn gọn, rõ ràng, súc tích.
- Mỗi đoạn khoảng 3–4 câu.
- Tổng độ dài văn bản khoảng 250–300 từ.
- Tránh lan man, hạn chế dùng ví dụ minh họa hoặc cảm xúc.
- Nội dung tập trung vào trọng tâm, dễ hiểu, in ra vừa trong 1 trang A4.
"""
    elif mode == "detail":
        return """
✍️ Yêu cầu phong cách viết:
- Trình bày đầy đủ, sâu sắc, có cảm xúc và dẫn dắt rõ ràng.
- Mỗi đoạn dài 5–8 câu, triển khai từng ý kỹ càng, có ví dụ, liên hệ thực tế hoặc cảm nhận cá nhân.
- Không giới hạn độ dài, có thể vượt 700 từ nếu cần.
- Dùng ngôn từ sinh động, hình ảnh hóa, đôi lúc bay bổng nhưng không rườm rà.
- Mở bài dẫn dắt tự nhiên, thân bài phân tích rõ ràng, kết bài tạo dư âm hoặc để lại suy nghĩ.
"""
    else:  # normal
        return """
✍️ Yêu cầu phong cách viết:
- Trình bày rõ ràng, đủ ý, mỗi đoạn triển khai một luận điểm chính.
- Mỗi đoạn dài 4–6 câu, có thể kèm ví dụ minh họa ngắn.
- Tổng độ dài khoảng 400–500 từ.
- Bố cục 3 phần: Mở bài dẫn dắt vấn đề, thân bài trình bày các ý chính, kết bài tóm tắt và nhận xét nhẹ nhàng.
- Văn bản nên in ra vừa trong 1 đến 1.5 trang A4.
"""

def clean_gpt_output(content: str) -> str:
    # 1. Xoá block ```...``` hoặc ```plaintext
    content = re.sub(r"```[\s\S]*?```", "", content)
    content = re.sub(r"```(plaintext|text|markdown)?", "", content)

    # 2. Xoá markdown: **đậm**, __gạch__, ~~gạch~~, >, #
    content = re.sub(r"[*_~]{1,2}", "", content)
    content = re.sub(r"^> ?", "", content, flags=re.MULTILINE)
    content = re.sub(r"^#+ ?", "", content, flags=re.MULTILINE)

    # 3. Xoá đầu dòng markdown kiểu danh sách
    content = re.sub(r"^[\-\*\+] ", "", content, flags=re.MULTILINE)

    # 4. Xoá các thẻ HTML (nếu có)
    content = re.sub(r"<[^>]+>", "", content)

    # 5. Xoá các dòng trắng liên tiếp → giữ đúng 1 dòng trống
    content = re.sub(r"\n{3,}", "\n\n", content)

    # 6. Chuẩn hoá dấu cách thừa
    content = re.sub(r"[ \t]+", " ", content)

    return content.strip()
@app.route("/smartdoc_history")
def smartdoc_history():
    offset = int(request.args.get("offset", 0))
    username = session.get("username")
    user = User.query.filter_by(username=username).first()
    docs = SmartDoc.query.filter_by(user_id=user.user_id).order_by(SmartDoc.created_at.desc()).offset(offset).limit(10).all()
    return jsonify({
        "items": [
            {
                "doc_id": str(doc.doc_id),
                "title": doc.title,
                "author": getattr(doc, "author_name", None),
                "created": doc.created_at.strftime("%d/%m/%Y %H:%M")
            } for doc in docs
        ]
    })
@app.route("/smartdoc/delete/<doc_id>", methods=["POST"])
def delete_smartdoc(doc_id):
    username = session.get("username")
    user = User.query.filter_by(username=username).first()
    doc = SmartDoc.query.filter_by(doc_id=doc_id, user_id=user.user_id).first()
    if doc:
        db.session.delete(doc)
        db.session.commit()
        return jsonify({"success": True})
    return jsonify({"success": False, "error": "Không tìm thấy văn bản"}), 404

@app.route("/smartdoc_search")
def smartdoc_search():
    keyword = request.args.get("keyword", "")
    username = session.get("username")
    user = User.query.filter_by(username=username).first()
    docs = SmartDoc.query.filter(SmartDoc.user_id == user.user_id, SmartDoc.title.ilike(f"%{keyword}%")).order_by(SmartDoc.created_at.desc()).limit(20).all()
    return jsonify({
        "items": [
            {
                "doc_id": str(doc.doc_id),
                "title": doc.title,
                "author": getattr(doc, "author_name", None),
                "created": doc.created_at.strftime("%d/%m/%Y %H:%M")
            } for doc in docs
        ]
    })

@app.route("/download_docx/<string:doc_id>")
def download_docx(doc_id):
    doc = SmartDoc.query.get(doc_id)
    if not doc or doc.file_type != "docx":
        return "Không tìm thấy tài liệu", 404

    from docx import Document
    import tempfile, re
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    word_doc = Document()
    style = word_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13.5)  # Chuẩn như preview

    for line in doc.content.split("\n"):
        line = line.strip()
        if not line:
            word_doc.add_paragraph()
            continue

        # Dòng là tiêu đề chính → căn giữa + in đậm + viết hoa
        if re.match(r"^(CỘNG HÒA|ĐỘC LẬP|ĐƠN\s|---)", line, re.IGNORECASE):
            p = word_doc.add_paragraph(line.upper())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
        else:
            p = word_doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Căn đều hai bên
            p.paragraph_format.first_line_indent = Pt(28)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5

    # Tạo file tạm
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        word_doc.save(tmp.name)
        tmp_path = tmp.name

    safe_title = re.sub(r'[\\/*?:"<>|]', "_", getattr(doc, 'title', 'SmartDoc'))
    return send_file(tmp_path, as_attachment=True, download_name=f"{safe_title}.docx")

@app.route("/update_doc_content", methods=["POST"])
def update_doc_content():
    data = request.get_json()
    doc_id = data.get("doc_id")
    new_content = data.get("content")

    if not doc_id or not new_content:
        return jsonify({"success": False, "error": "Thiếu thông tin cần thiết"}), 400

    doc = SmartDoc.query.filter_by(doc_id=doc_id).first()
    if not doc:
        return jsonify({"success": False, "error": "Không tìm thấy tài liệu"}), 404

    # Cập nhật nội dung mới
    doc.content = new_content
    db.session.commit()

    return jsonify({"success": True})
# --------------------------
# ✨ Code Smasher utilities
# --------------------------
def flatten_js(code):
    code = re.sub(r'//.*', '', code)
    code = re.sub(r'/\*[\s\S]*?\*/', '', code)
    lines = [line.strip() for line in code.splitlines() if line.strip()]
    code = ' '.join(lines) 
    code = re.sub(r'\s*([{}();=+\-*/,:<>])\s*', r'\1', code)
    return code

def lineslicer_js(code: str) -> str:
    """
     Gộp toàn bộ khối (JS & CSS), mỗi khối thành 1 dòng riêng biệt.
     Tự xử lý <style>...</style>, xóa comment JS/CSS/HTML, bỏ khoảng trắng dư.
     Không gọi hàm phụ – tất cả xử lý trong 1 hàm duy nhất.
    """
    code = re.sub(r'//.*', '', code)                    
    code = re.sub(r'/\*[\s\S]*?\*/', '', code)             
    code = re.sub(r'<!--[\s\S]*?-->', '', code)         
    style_blocks = re.findall(r'<style[^>]*>(.*?)</style>', code, flags=re.DOTALL)
    for raw_css in style_blocks:
        css = re.sub(r'/\*[\s\S]*?\*/', '', raw_css)      
        css_lines = [line.strip() for line in css.splitlines() if line.strip()]
        css_blocks = []
        css_current = ''
        brace_count = 0
        for line in css_lines:
            css_current += ' ' + line
            brace_count += line.count('{') - line.count('}')
            if brace_count == 0 and css_current.strip():
                css_blocks.append(re.sub(r'\s+', ' ', css_current.strip()))
                css_current = ''
        if css_current.strip():
            css_blocks.append(re.sub(r'\s+', ' ', css_current.strip()))
        flat_css = '\n'.join(css_blocks)
        code = code.replace(raw_css, flat_css)
    lines = [line.strip() for line in code.splitlines() if line.strip()]
    blocks = []
    current = ''
    brace_stack = []
    for line in lines:
        current += ' ' + line
        for char in line:
            if char in '{[':
                brace_stack.append(char)
            elif char in '}]':
                if brace_stack and ((char == '}' and brace_stack[-1] == '{') or (char == ']' and brace_stack[-1] == '[')):
                    brace_stack.pop()
        if not brace_stack and current.strip():
            blocks.append(re.sub(r'\s+', ' ', current.strip()))
            current = ''
    if current.strip():
        blocks.append(re.sub(r'\s+', ' ', current.strip()))
    return '\n'.join(blocks)


def obfuscate_smart(code, options=None):
    client = create_openai_client(model="gpt-4o")

    extra_instructions = ""
    if options:
        if options.get("block_f12"):
            extra_instructions += "\nAlso add JavaScript that blocks F12 and Developer Tools (Ctrl+Shift+I)."
        if options.get("disable_console"):
            extra_instructions += "\nAlso disable console.log, console.warn, and console.error."
        if options.get("prevent_copy"):
            extra_instructions += "\nAlso prevent right-click and copying actions on the webpage."

    prompt = f"""
You are a commercial-grade JavaScript obfuscator, codenamed "MindTwist Pro".

🔐 Your mission:
Obfuscate the JavaScript code below into a **deeply protected**, **hard-to-read**, but **100% logically preserved** form.

You MUST apply:
- Renaming of all variables and functions into mangled or meaningless names
- Obfuscation of control flow (e.g., ternaries, indirect indexing, switch rewriting)
- Use of `addEventListener` instead of inline event attributes (`onclick`, etc.)
- Wrapping all logic in a self-invoking anonymous function `(function(){{...}})();`
- ⚠️ Do not rename or inline any function that is meant to be used in HTML (like onclick="submitStep1()"). If a function is referenced from HTML, it must be globally accessible via `window.functionName = ...` or remain named the same.
- If the original code has `async` / `await`, preserve the async context correctly. Do NOT wrap inside regular `Function(...)` if it breaks async
- If the code uses DOM API (`document.getElementById`, etc.), preserve the scope — do NOT break or change their functionality

📜 Strict output requirements:
1. Return the result as **one single valid JavaScript line** (runnable via `eval(...)`)
2. **No line breaks**, **no indentation**, **no extra whitespace**
3. Do NOT return any explanation, comment, markdown formatting, or syntax hint
4. Simulate execution internally to ensure obfuscated result works exactly like the original
5. If the original contains `<script>` tags, strip them automatically before obfuscating

{extra_instructions}

Here is the original JavaScript code to obfuscate:
{code}

⚠️ Final output: ONE SINGLE JavaScript LINE ONLY. NO formatting, NO extra characters, NO explanation.
"""


    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=10000
    )
    content = response.choices[0].message.content.strip()
    if content.startswith("```javascript"):
        content = content.replace("```javascript", "").strip()
    if content.endswith("```"):
        content = content[:-3].strip()

    # Ép về 1 dòng
    content = content.replace('\n', '').replace('\r', '')
    return fix_brackets(content)


def encrypt_hexlock(code, options=None):
    client = create_openai_client(model="gpt-4o")

    extra_instructions = ""
    if options:
        if options.get("block_f12"):
            extra_instructions += "\n- Bổ sung đoạn mã chặn DevTools (F12, Ctrl+Shift+I)"
        if options.get("disable_console"):
            extra_instructions += "\n- Vô hiệu hóa console.log / warn / error"
        if options.get("prevent_copy"):
            extra_instructions += "\n- Ngăn người dùng click chuột phải và sao chép nội dung"

    prompt = f"""
Bạn là một hệ thống chuyển đổi mã JavaScript thành dạng khó đọc, tên là "HexLock".

🔐 Nhiệm vụ của bạn:
- Chuyển đổi đoạn mã JavaScript dưới đây thành một phiên bản an toàn, khó bị đọc hiểu, nhưng vẫn **giữ nguyên chức năng**.
- Được phép sử dụng:
  - `String.fromCharCode(...)` hoặc `.charCodeAt()` để mã hóa ký tự
  - `Function` constructor, kết hợp `eval`, closure, IIFE
  - Đổi tên tất cả biến, hàm, miễn là logic vẫn đúng
  - Kỹ thuật lồng ghép hàm (nested calls), closure, gọi gián tiếp (`[f]()`)

⚙️ Nguyên tắc bắt buộc:
- Nếu có `async function tên_hàm`, phải giữ nguyên `async` và tên gốc. Không đổi tên hàm async.
- Nếu có thao tác DOM (`getElementById`, `onclick`, `value`,...), không được đảo thứ tự hoặc làm sai luồng xử lý.
- Nếu có `await fetch(...)`, **không được bọc bằng `Function(...)` hoặc IIFE nếu làm sai logic bất đồng bộ.**
- Nếu đoạn mã có hàm dùng trong HTML (vd: `onclick="submitCode()"`), **phải tồn tại toàn cục qua `window.tênHàm = ...`**

🧪 Quy tắc an toàn:
- `.toString(radix)` chỉ được dùng nếu `radix` là số nguyên **từ 2 đến 36**. Nếu không chắc, **KHÔNG được dùng**.
- Không tạo `radix` bằng các biểu thức bất định (như biến từ ngoài).
- Không được sinh `eval(...)` lồng nhau (eval trong eval).

📦 Yêu cầu về đầu ra:
1. Trả về **một dòng JavaScript duy nhất**, có thể `eval(...)` trực tiếp
2. Không có xuống dòng, không chứa khoảng trắng thừa
3. Logic hoạt động **phải y hệt bản gốc**
4. Nếu đầu vào chứa `<script>...</script>`, hãy **tự loại bỏ** wrapper này
5. Nếu có định nghĩa hàm toàn cục, gán lại bằng `window.tênHàm = tênHàm`
6. Không chỉ viết gọn — **bắt buộc mã hóa ký tự bằng `String.fromCharCode(...)` hoặc tương đương**
7. Không dùng `await` nếu không cần thiết, tránh lỗi runtime
8. Phải tự kiểm tra tính hợp lệ của mã trước khi trả ra kết quả

{extra_instructions}

🎁 **Chỉ trả về đúng 1 dòng JavaScript duy nhất**. Không được giải thích, không xuống dòng.
Đoạn mã cần xử lý:
{code}

"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=10000
    )
    content = response.choices[0].message.content.strip()
    if content.startswith("```javascript"):
        content = content.replace("```javascript", "").strip()
    if content.endswith("```"):
        content = content[:-3].strip()
    return content.replace("\n", "").replace("\r", "")


def reset_code_smasher_if_needed(user):
    today = date.today()
    if user.cs_usage_reset_at != today:
        user.cs_flatline_count = 0
        user.cs_lineslicer_count = 0
        user.cs_mindtwist_count = 0
        user.cs_hexlock_count = 0
        user.cs_usage_reset_at = today
def fix_brackets(code):
    open_paren = code.count('(')
    close_paren = code.count(')')
    open_brace = code.count('{')
    close_brace = code.count('}')
    return code + (')' * (open_paren - close_paren)) + ('}' * (open_brace - close_brace))


def is_gpt_vip_active(user):
    if user.vip_gpt_ai and user.vip_until_gpt:
        try:
            if isinstance(user.vip_until_gpt, str):
                vip_expiry = datetime.strptime(user.vip_until_gpt, "%Y-%m-%d %H:%M:%S")
            else:
                vip_expiry = user.vip_until_gpt  # Đã là datetime
            return vip_expiry > datetime.now()
        except Exception:
            return False
    return False
# --------------------------
# 🧠 Route /code_smasher
# --------------------------
@app.route("/code_smasher", methods=["POST"])
def code_smasher():
    data = request.get_json()
    code = data.get("code", "").strip()
    mode = data.get("mode")

    if not code or mode not in ["Flatline", "LineSlicer", "MindTwist", "HexLock"]:
        return jsonify({"success": False, "error": "Thiếu dữ liệu hoặc chế độ không hợp lệ."}), 400

    user = None
    if "username" in session:
        user = User.query.filter_by(username=session["username"]).first()

    # ⚠️ Chỉ chặn nếu là chế độ cao cấp
    if mode in ["MindTwist", "HexLock"]:
        if not user:
            return jsonify({
                "success": False,
                "error": "❌ Bạn cần đăng nhập để sử dụng chế độ cao cấp."
            }), 403
        if not is_gpt_vip_active(user):
            return jsonify({
                "success": False,
                "error": "❌ Chế độ này chỉ dành cho người dùng đã nâng cấp VIP."
            }), 403

    try:
    
        if mode == "Flatline":
            result = flatten_js(code)
            if user:
                reset_code_smasher_if_needed(user)
                user.cs_flatline_count += 1
        elif mode == "LineSlicer":
            result = lineslicer_js(code)  
            if user:
                reset_code_smasher_if_needed(user)
                user.cs_lineslicer_count += 1
        elif mode == "MindTwist":
            result = obfuscate_smart(code, options=data)
            user.cs_mindtwist_count += 1
        elif mode == "HexLock":
            result = encrypt_hexlock(code, options=data)
            user.cs_hexlock_count += 1
        if user:
            db.session.commit()
        return jsonify({"success": True, "result": result})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
def generate_js_slug():
    part1 = "js" + secrets.token_hex(3)
    part2 = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(5))
    return f"{part1}-{part2}"
@app.route("/js_mifing", methods=["GET"])
def redirect_js_mifing():
    slug = generate_js_slug()
    return redirect(f"/jsmifXcodeninja/{slug}")
@app.route("/jsmifXcodeninja/<slug>", methods=["GET"])
def js_mifing(slug):
    username = session.get("username")
    if not username:
        return redirect("/login")
    user = User.query.filter_by(username=username).first()
    if not user:
        return redirect("/login")
    return render_template("js_mifing.html")

@app.route("/code_smasher/verify", methods=["POST"])
def verify_code_result():
    code = request.json.get("code", "")
    if not code.strip():
        return jsonify({"success": False, "reply": "Thiếu mã cần kiểm tra."}), 400
    try:
        result = verify_code_with_gpt(code)
        return jsonify(success=result)
    except Exception:
        return jsonify(success=False)


def verify_code_with_gpt(code):
    prompt = f"""
You are an expert JavaScript evaluator.
Below is a heavily obfuscated JavaScript code. Do NOT judge by appearance or encoding. Your task is:
- Internally simulate its execution (as if eval() was used in a browser)
- ONLY return '✅ OK' if the code runs without syntax or runtime errors
- If you detect syntax errors, missing brackets, unsafe recursion, or logic bugs that cause crash, return '❌ FAILED'
Respond with exactly one of the following:
- ✅ OK
- ❌ FAILED
No other explanation, no formatting, no markdown.
Code:
{code}
"""
    try:
        client = create_openai_client(model="gpt-4o")
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=150
        )
        answer = response.choices[0].message.content
        if not answer:
            return False
        answer = answer.strip().replace("```", "").strip()
        if answer == "✅ OK":
            return True
        elif answer == "❌ FAILED":
            return False
        else:
            return False

    except Exception as e:
        return False
MAIL_USER = os.getenv("EMAIL_ADDRESS")
MAIL_PASS = os.getenv("EMAIL_PASSWORD")

def extract_transaction_data(subject, body):
    """Tuỳ chỉnh hàm này để tìm số tiền và nội dung chuyển khoản"""
    content = subject + "\n" + body

    # Ví dụ: tìm mã giao dịch dạng SLV_123ABC và số tiền
    note_match = re.search(r"(SLV_\w+)", content)
    amount_match = re.search(r"(\d[\d,.]+)\s?VN[Đđ]", content)

    note = note_match.group(1) if note_match else None
    amount = amount_match.group(1).replace(',', '').replace('.', '') if amount_match else None

    if note and amount:
        return note.strip(), int(amount)
    return None, None

def check_new_emails():
    with imaplib.IMAP4_SSL("imap.gmail.com") as imap:
        imap.login(MAIL_USER, MAIL_PASS)
        imap.select("inbox")

        status, messages = imap.search(None, '(UNSEEN)')
        email_ids = messages[0].split()

        for eid in email_ids:
            _, msg_data = imap.fetch(eid, "(RFC822)")
            for response_part in msg_data:
                if isinstance(response_part, tuple):
                    msg = email.message_from_bytes(response_part[1])
                    subject, encoding = decode_header(msg["Subject"])[0]
                    subject = subject.decode(encoding or "utf-8") if isinstance(subject, bytes) else subject

                    body = ""
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == "text/plain":
                                body = part.get_payload(decode=True).decode()
                                break
                    else:
                        body = msg.get_payload(decode=True).decode()

                    note, amount = extract_transaction_data(subject, body)

                    if note and amount:
                        with app.app_context():
                            txn = Transaction.query.filter_by(note=note, amount=amount, status="pending").first()
                            if txn:
                                txn.status = "approved"
                                txn.approved_at = datetime.utcnow()
                                db.session.commit()
    imap.logout()

def parse_email_content(content):
    match = re.search(r'\b(TXN\w{8,})\b', content.strip(), re.IGNORECASE)
    if not match:
        return None, None

    txn_code = match.group(1).strip()

    txn = Transaction.query.filter_by(txn_id=txn_code).first()
    if not txn:
        return None, None

    return txn.package, txn_code

def auto_approve_transactions():
    EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
    EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        return

    try:
        mail = imaplib.IMAP4_SSL("imap.gmail.com")
        mail.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        mail.select("inbox")

        result, data = mail.search(None, "UNSEEN")
        mail_ids = data[0].split()

        for num in mail_ids:
            result, msg_data = mail.fetch(num, "(RFC822)")
            raw_email = msg_data[0][1]
            msg = email.message_from_bytes(raw_email)

            subject = msg["subject"] or ""
            content = ""

            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        content += part.get_payload(decode=True).decode()
            else:
                content = msg.get_payload(decode=True).decode()

            full_text = subject + "\n" + content
            pkg_code, txn_code = parse_email_content(full_text)

            if not pkg_code or not txn_code:
                continue

            transaction_record = Transaction.query.filter_by(txn_id=txn_code).first()
            if not transaction_record:
                continue

            user = User.query.filter_by(user_id=transaction_record.user_id).first()
            if not user:
                continue

            if transaction_record.status == "approved":
                continue

            transaction_record.status = "approved"

            pkg_map = {
                "vip_ai_lite": "vip_ai_lite",
                "SLV(5day)": "vip_gpt_5d",
                "SLV(15day)": "vip_gpt_15d",
                "SLV(30day)": "vip_gpt_30d"
            }
            selected_pkg = pkg_map.get(pkg_code)
            if not selected_pkg:
                continue

            grant_vip(user.username, selected_pkg)
            db.session.commit()

        mail.logout()

    except Exception:
        pass

@app.route("/transaction/status/<txn_code>")
def check_txn_status(txn_code):
    txn = Transaction.query.filter_by(txn_id=txn_code).first()
    if not txn:
        return jsonify({"status": "not_found"})
    return jsonify({"status": txn.status})
def run_auto_approval():
    while True:
        auto_approve_transactions()
        time.sleep(60)  # kiểm tra mỗi 60 giây

# ✅ Chạy quét mail ở luồng nền
threading.Thread(target=run_auto_approval, daemon=True).start()  

from flask_apscheduler import APScheduler
from image_cleaner import clean_old_images

class Config:
    SCHEDULER_API_ENABLED = True

app.config.from_object(Config())
scheduler = APScheduler()

@scheduler.task('interval', id='image_cleaner_24h', hours=24)
def run_cleaner_every_24h():
    clean_old_images()

scheduler.init_app(app)
scheduler.start()
from flask import send_from_directory

@app.route('/sitemap.xml')
def serve_sitemap():
    return send_file('sitemap.xml', mimetype='application/xml')
@app.route('/robots.txt')
def robots_txt():
    return send_from_directory('static', 'robots.txt')
@app.route('/favicon.ico')
def favicon():
    return redirect(url_for('static', filename='logos/favicon.png'))
@app.route("/api/set_intro_seen", methods=["POST"])
def set_intro_seen():
    if "username" not in session:
        return jsonify({"success": False, "message": "Chưa đăng nhập"})

    user = User.query.filter_by(username=session["username"]).first()
    if not user:
        return jsonify({"success": False, "message": "Không tìm thấy user"})

    user.has_seen_intro = True
    db.session.commit()

    return jsonify({"success": True})
#TẠO ẢNH THEO MÔ TẢ 
from flask import request, jsonify, session  # ✅ THÊM session ở đây
from models.user import User
from extensions import db
from openai_config import create_openai_client

def translate_vi_to_en(prompt_vi, style_type="default"):
    try:
        client = create_openai_client("gpt-4o")
        base_instruction = (
            "You are a professional AI assistant specialized in translating Vietnamese prompts into English for safe image generation.\n"
            "Your translation must be clean, neutral, and compliant with OpenAI's safety guidelines.\n"
            "Avoid using specific terms like 'girl', 'boy', 'face', 'woman', 'celebrity', 'real person'.\n"
            "Use general terms like 'subject', 'character', or 'individual'.\n"
            "If the original input contains sensitive or unsafe content, rephrase it in a creative and safe way without losing the intent.\n"
            "Do not mention facial recognition, names, or anything identifiable.\n"
            "Always make the English prompt vivid, imaginative, and detailed, suitable for high-quality AI image generation.\n"
            "Your output will be used directly as a prompt for image generation by DALL·E or GPT-4o."
        )
        if style_type == "logo":
            base_instruction = (
                "You are an AI assistant helping translate Vietnamese prompts into detailed English descriptions to generate professional logos using AI.\n"
                "Focus on design concepts, style, symbolism, minimalism, and branding tone.\n"
                "Avoid describing people. Focus on shapes, icons, color palettes, or style (e.g. modern, vintage).\n"
                "Translate the prompt cleanly and clearly for logo generation with GPT/DALL·E.\n"
                "Always make the English prompt vivid, imaginative, and detailed, suitable for high-quality AI image generation."
            )
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": base_instruction},
                {"role": "user", "content": f'Dịch sang tiếng Anh để mô tả hình ảnh AI: "{prompt_vi}"'}
            ],
            temperature=0.3,
            max_tokens=300
        )
        translated = response.choices[0].message.content.strip()
        return translated

    except Exception as e:
        return prompt_vi
def generate_image_slug():
    return str(uuid.uuid4())

@app.route("/image_crafter", methods=["GET"])
def redirect_image_crafter():
    slug = generate_image_slug()
    return redirect(f"/ImgJarodasCar/{slug}") 

@app.route("/ImgJarodasCar/<slug>")
def image_crafter_page(slug):
    if is_maintenance("image"):
        return render_template("bao_tri_page.html", message="Tính năng tạo ảnh đang được bảo trì. Vui lòng quay lại sau!")

    if not session.get("user_id"):
        return redirect("/login")

    return render_template("image_crafter.html")
# Route tạo ảnh AI
@app.route("/generate_image_ai", methods=["POST"])
def generate_image_ai():
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"error": "Bạn chưa đăng nhập."}), 401

    user = User.query.filter_by(user_id=user_id).first()
    if not user:
        return jsonify({"error": "Tài khoản không tồn tại."}), 404

    if is_maintenance("image"):
        flash("🚧 Tính năng tạo ảnh đang được bảo trì!", "warning")
        return redirect("/") 
    if user.image_generation_blocked:
        return jsonify({"error": "Tính năng tạo ảnh đã bị khóa vì vượt quá 3 lượt miễn phí."}), 403

    if not user.vip_gpt and user.image_generation_used >= 3:
        user.image_generation_blocked = True
        db.session.commit()
        return jsonify({"error": "Bạn đã hết 3 lượt tạo ảnh miễn phí. Hãy nâng cấp gói SLV để tiếp tục."}), 403

    data = request.get_json()
    prompt_vi = data.get("prompt", "").strip()
    style_type = "logo" if "logo" in prompt_vi.lower() else "default"

    if not prompt_vi:
        return jsonify({"error": "Bạn chưa nhập mô tả yêu cầu."}), 400

    # Dịch tiếng Việt sang tiếng Anh tùy theo style
    prompt_en = translate_vi_to_en(prompt_vi, style_type)

    try:
        client = create_openai_client("gpt-4o")
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt_en,
            n=1,
            size="1024x1024"

        )
        image_url = response.data[0].url
    except Exception as e:
        return jsonify({"error": f"Lỗi khi tạo ảnh: {str(e)}"}), 500

    user.image_generation_used += 1
    db.session.commit()

    return jsonify({
        "message": "Tạo ảnh thành công!",
        "image_url": image_url
    })

from collections import OrderedDict

GAME_CACHE = {}            # {game_id: [messages]}
START_TIMES = {}           # {game_id: datetime}
MAX_MSGS = 60              # cắt bớt để không phình (bỏ system ra)

def cache_set(gid, hist):
    sys = hist[:1]
    body = hist[1:]
    if len(body) > MAX_MSGS:
        body = body[-MAX_MSGS:]
    GAME_CACHE[gid] = sys + body

def cache_get(gid):
    return GAME_CACHE.get(gid, [{"role": "system", "content": GAME_PROMPT}])

def build_msgs_for_gpt(hist, keep_last=20):
    sys = hist[:1]
    body = hist[1:]
    if len(body) > keep_last:
        body = body[-keep_last:]
    return sys + body

# LRU cache nhỏ cho kiểm tra hợp lệ (thay cho session['noi_tu_valid_cache'])
IS_VALID_CACHE = OrderedDict()
MAX_VALID = 1000
def valid_cache_set(k, v):
    IS_VALID_CACHE[k] = v
    IS_VALID_CACHE.move_to_end(k)
    if len(IS_VALID_CACHE) > MAX_VALID:
        IS_VALID_CACHE.popitem(last=False)    
GAME_PROMPT = """
Bạn đang tham gia một trò chơi tên là Nối Từ bằng tiếng Việt. Đây KHÔNG phải là cuộc trò chuyện thông thường.

**Cách chơi:**
- Mỗi người nói một cụm từ có nghĩa gồm đúng 2 từ tiếng Việt.
- Từ **đầu tiên** của cụm từ mới phải trùng với từ **cuối cùng** của cụm từ trước.
- Cụm từ phải có nghĩa rõ ràng, phổ biến trong tiếng Việt. Không dùng từ bịa, từ vô nghĩa.

**Luật chơi:**
- Không được lặp lại cụm từ đã dùng.
- Trả lời đúng **một cụm từ duy nhất mỗi lượt**. KHÔNG bình luận thêm, không nhắn tin kiểu xã giao.
- Trả lời chính xác, đúng chính tả, đúng nghĩa.
- Nếu tôi nhập "chào bạn", "xin chào", hay các cụm phổ biến, bạn **vẫn phải coi đó là cụm từ hợp lệ gồm 2 từ** và nối tiếp như bình thường.

**Yêu cầu dành cho bạn (AI):**
- Luôn phản hồi theo đúng luật trên.
- KHÔNG được trả lời kiểu “bạn đi trước nhé”, “xin chào”, “ok bạn” v.v.
- **KHÔNG được chịu thua nếu chưa đạt ít nhất 25 lượt** (tính cả lượt của bạn và tôi).
- Nếu cụm từ từ người chơi còn dễ hoặc phổ biến, bạn phải tiếp tục và ưu tiên tìm các cụm ít phổ biến hơn nhưng hợp lệ để mở đường nối dài hơn.
- Khi bí, hãy kiểm tra kỹ kho từ vựng và suy nghĩ nhiều hướng (đồng nghĩa, từ ghép khác nhau) trước khi quyết định đầu hàng.
- Chỉ đầu hàng khi cụm từ quá hiếm, khó nối hoặc sau khi đã thử mọi khả năng mà không còn từ hợp lệ nào.
- Khi đầu hàng, hãy dùng câu tự nhiên, hài hước và hợp ngữ cảnh, nêu ngắn gọn lý do thua (ví dụ: “Hết từ hợp lệ rồi, bó tay!”).

**Bạn là người đi sau tôi. Bắt đầu từ bây giờ.**
"""


def gpt_chat_safe(client, messages, *, temperature=0.85, max_tokens=30,
                  timeout=7.0, retries=2, backoff=0.6):
    """
    Gọi GPT có timeout + retry. Hết retry thì trả None (không raise).
    Dùng exponential backoff + jitter để tránh dồn tải.
    """
    for i in range(retries + 1):
        try:
            res = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
                timeout=timeout,   # nếu không hỗ trợ thì bỏ dòng này
                # max_retries=0     # giữ tự-retry ở vòng for
            )
            return (res.choices[0].message.content or "").strip()
        except Exception:
            if i == retries:
                return None
            # exponential backoff + jitter
            sleep_s = backoff * (2 ** i) + random.uniform(0, 0.25)
            time.sleep(sleep_s)

def gpt_ask_safe(client, prompt: str, *, max_tokens=12, timeout=5.0):
    """Phiên bản an toàn cho _ask()."""
    msgs = [{"role": "user", "content": prompt}]
    return gpt_chat_safe(client, msgs, temperature=0, max_tokens=max_tokens, timeout=timeout)

def normalize(text):
    return unicodedata.normalize('NFC', text.lower().strip())

import random
def _sample_surrender_profile(strict: bool = True) -> dict:
    """
    strict=True  -> đúng yêu cầu 40% / 20% / 10% (không jitter).
    strict=False -> cho jitter ±5% để mỗi ván hơi khác nhau.
    """
    base = {"p30_40": 0.40, "p40_60": 0.20, "p60_100": 0.10}
    if strict:
        return {**base, "seed": random.getrandbits(32)}
    # jitter nhẹ nếu muốn mỗi ván khác nhau chút
    jitter = lambda p: max(0.0, min(0.9, round(p + random.uniform(-0.05, 0.05), 2)))
    return {
        "p30_40": jitter(base["p30_40"]),
        "p40_60": jitter(base["p40_60"]),
        "p60_100": jitter(base["p60_100"]),
        "seed": random.getrandbits(32)
    }

@app.route('/game_noi_tu/start', methods=['POST'])
def start_noi_tu():
    data = request.get_json(silent=True) or {}

    game_id = data.get("game_id", str(uuid.uuid4()))
    session['noi_tu_game_id'] = game_id

    # ✅ Lịch sử để RAM, không lưu vào session
    init_hist = [{"role": "system", "content": GAME_PROMPT}]
    cache_set(game_id, init_hist)

    session['noi_tu_turn'] = 0

    possible_turns = [10, 15, 20, 30, 40, 50, 60]
    weights        = [ 1,   2,  6, 16, 28, 28, 19 ]
    session['noi_tu_max_turns'] = random.choices(possible_turns, weights=weights, k=1)[0]

    session['noi_tu_min_surrender_turn'] = int(data.get("min_surrender_turn") or 25)

    session['noi_tu_reward'] = data.get("reward_mode", "ask_any_question")
    session['noi_tu_ended'] = False
    session['noi_tu_start_time'] = datetime.utcnow().isoformat()  # nhẹ hơn
    session['noi_tu_wrong_count'] = 0
    session['noi_tu_used'] = []
    strict_profile = bool(data.get("strict_profile", True))
    session['noi_tu_surrender_profile'] = _sample_surrender_profile(strict=strict_profile)

    # (Tuỳ chọn) lưu start_time “xịn” ở server nếu cần
    START_TIMES[game_id] = datetime.utcnow()

    return jsonify({"game_id": game_id})


@app.route('/game_noi_tu/play', methods=['POST'])
def play_noi_tu():
    data = request.get_json(silent=True)
    if not data or "game_id" not in data:
        return jsonify({"error": "Thiếu dữ liệu"}), 400

    user_word_raw = data.get("user_word", "")
    user_word = user_word_raw.strip().lower()

    # ==== Validator mạnh để kiểm tra cụm từ ====
    import re, json, math
    INTERJECTION_SET = {"nè","nha","nhé","nhỉ","ơi","ạ","ơ","ồ","á","đi","đấy"}
    VI_TOKEN = re.compile(r"^[a-zà-ỹđ]+(-[a-zà-ỹđ]+)?$", re.IGNORECASE)

    def well_formed_two_words(t: str) -> bool:
        w = t.split()
        return (
            len(w) == 2
            and all(VI_TOKEN.match(tok) for tok in w)
            and not any(tok in INTERJECTION_SET for tok in w)
        )
    STRONG_WHITELIST = {
        "sản phẩm","học sinh","giá trị","thị trường","nhà trường","tài sản","bầu trời","cao vời",
        "công nghệ","khoa học","đời sống","kết quả","cảm xúc","trí tuệ","cầu vồng","não bé","giáo khoa",
        "gạo chín", "gạo lúa", "lúa non", "nước sôi", "cá khô", "nước mắm","trường học","gió trời",
         "gió mạnh", "mưa to", "trời quang", "nắng gắt","đúm đỉnh", "đàn đúm","đây nè","màng lượt"
    }
    GREETING_WHITELIST = {
        "xin chào", "chào bạn", "chào anh", "chào em", "chúc mừng", "cảm ơn"
    }
    def canon(s: str) -> str:
        s = normalize(s)
        s = re.sub(r"[^\wÀ-ỹđ ]", "", s)
        return " ".join(s.split()).lower()
    def _first_token(x: str) -> str:
        if not x: return ""
        t = x.strip().lower()
        t = re.sub(r"[^\wÀ-ỹđ ]", "", t)
        return t.split()[0] if t.split() else ""
    def _ask(client, prompt: str) -> str:
        ans = gpt_ask_safe(client, prompt, max_tokens=12, timeout=5.0)
        return (ans or "").strip()

    # === DÙNG RAM LRU: bỏ session cache ===
    # IS_VALID_CACHE = session.setdefault("noi_tu_valid_cache", {})
    def is_valid_vi_2word(phrase: str, client) -> bool:
        t = canon(phrase)
        if not well_formed_two_words(t):
            return False

        if t in STRONG_WHITELIST or t in GREETING_WHITELIST:
            valid_cache_set(t, True)
            return True

        if t in IS_VALID_CACHE:
            return IS_VALID_CACHE[t]

        def ask_once() -> bool | None:
            p = (
                f'Cụm "{t}" có phải là cụm từ tiếng Việt gồm đúng 2 từ, có nghĩa rõ ràng không? '
                'Chỉ trả lời duy nhất: "CO" hoặc "KHONG".'
            )
            ans = gpt_ask_safe(client, p, max_tokens=2, timeout=4.0) or ""
            a = _first_token(ans)
            if a in ("co", "có", "ok", "yes"):
                return True
            if a in ("khong", "không", "no"):
                return False
            return None

        v1 = ask_once()
        v2 = ask_once()
        votes = [v for v in (v1, v2) if v is not None]
        ok = (votes.count(True) >= votes.count(False)) if votes else False

        valid_cache_set(t, ok)
        return ok

    # ---- helpers ----
    def looks_like_surrender(text: str) -> bool:
        t = (text or "").lower()
        patterns = ["tôi chịu", "chịu thua", "tôi thua", "bó tay", "không nghĩ ra", "hết từ", "bí quá"]
        return any(p in t for p in patterns)

    def last_word(s: str) -> str:
        parts = (s or "").strip().split()
        return parts[-1].strip(string.punctuation).lower() if parts else ""

    def first_word(s: str) -> str:
        parts = (s or "").strip().split()
        return parts[0].strip(string.punctuation).lower() if parts else ""

    LOCAL_FALLBACKS = {
        "học": ["học sinh", "học phí", "học bổng", "học đường"],
        "sinh": ["sinh nhật", "sinh hoạt", "sinh viên", "sinh lực"],
        "viên": ["viên gạch", "viên thuốc", "viên ngọc"],
        "ngọc": ["ngọc trai", "ngọc bích"],
        "bích": ["bích họa"],
        "trời": ["trời mưa", "trời xanh", "trời tối", "trời mọc"],
        "mọc": ["mọc rêu", "mọc mầm"],
        "mầm": ["mầm non"],
        "non": ["non sông", "non xanh"],
        "sông": ["sông núi", "sông suối"],
        "núi": ["núi rừng"],
        "rừng": ["rừng rậm", "rừng thông"],
        "thông": ["thông minh", "thông thạo"],
        "minh": ["minh họa", "minh chứng"],
        "mưa": ["mưa rào", "mưa phùn"],
        "phùn": ["phùn phụt"],
        "trăng": ["trăng non", "trăng rằm"],
        "rằm": ["rằm tháng"],
        "yêu": ["yêu thương", "yêu cầu"],
        "cầu": ["cầu vồng", "cầu thang", "cầu đường"],
        "vồng": ["vồng tôm"],
        "tôm": ["tôm sú", "tôm hùm"],
        "hùm": ["hùm xám"],
    }

    def find_backup(start: str, used_words_list: list[str]) -> str | None:
        start = (start or "").lower()
        used = set(canon(u) for u in used_words_list)
        for cand in LOCAL_FALLBACKS.get(start, []):
            if canon(cand) not in used:
                return cand
        return None

    def force_retry_with_gpt(start: str, history_msgs: list) -> str | None:
        try:
            strict_prompt = (
                "Bạn đang chơi Nối Từ. Hãy trả về DUY NHẤT một cụm từ có nghĩa gồm đúng 2 từ, "
                f"trong đó từ đầu tiên PHẢI là '{start}'. "
                "Không thêm lời giải thích, không ký tự thừa."
            )
            strict_msgs = [{"role": "system", "content": GAME_PROMPT}] + history_msgs + [
                {"role": "user", "content": strict_prompt}
            ]
            cand = gpt_chat_safe(
                client,
                strict_msgs,
                temperature=0.3,
                max_tokens=20,
                timeout=6.0,
                retries=2,
                backoff=0.5
            )
            if not cand:
                return None
            cand = cand.strip().lower()
            if len(cand.split()) == 2 and first_word(cand) == start:
                return cand
        except Exception:
            pass
        return None

    def hard_rescue(start, used_words, history_msgs, client, max_tries=3):
        b = find_backup(start, used_words)
        if b and canon(b) not in used_words:
            return b
        for _ in range(max_tries):
            r = force_retry_with_gpt(start, history_msgs)
            if r and canon(r) not in used_words:
                return r
        return None

    def surrender_prob_by_profile(profile: dict, turn: int, min_sur: int) -> float:
        if turn < max(min_sur, 30):
            return 0.0
        p30_40  = float(profile.get("p30_40", 0.40))
        p40_60  = float(profile.get("p40_60", 0.20))
        p60_100 = float(profile.get("p60_100", 0.10))
        if 30 <= turn < 40:   return p30_40
        if 40 <= turn < 60:   return p40_60
        if 60 <= turn <= 100: return p60_100
        return 0.0

    # ================== VALIDATIONS GIỮA VÁN ==================
    game_id = data.get("game_id")
    if session.get("noi_tu_game_id") != game_id:
        return jsonify({"error": "ID game không khớp hoặc đã hết hạn"}), 400
    if session.get("noi_tu_ended", False):
        return jsonify({"error": "Phiên chơi không hợp lệ hoặc đã kết thúc"}), 400

    # Lấy lịch sử từ RAM, không dùng session nữa
    history = cache_get(game_id)
    used_words = [canon(h["content"]) for h in history if h["role"] in ["user", "assistant"]]

    # ================== CASES ĐẦU VÁN ==================
    if session.get("noi_tu_turn", 0) == 0 and user_word:
        user_parts = user_word.split()
        if len(user_parts) != 2:
            return jsonify({"error": "Bạn phải nhập đúng 2 từ có nghĩa. Ví dụ: 'tình yêu', 'bầu trời', 'học sinh'..."}), 400

        WHITELIST_TU = {
            "lim rim","lai rai","ngoa ngoắt","âm u","tối om","mập mờ","hùm xám","vồng tôm","sơn hà","địa đạo","mãnh thú"
        }
        try:
            client = create_openai_client("gpt-4o")
            if canon(user_word) not in {canon(x) for x in WHITELIST_TU}:
                if not is_valid_vi_2word(user_word, client):
                    return jsonify({"error": "Cụm từ bạn nhập không có nghĩa rõ ràng. Vui lòng thử lại."}), 400
        except Exception as e:
            return jsonify({"error": "Lỗi kiểm tra cụm từ: " + str(e)}), 500

        session["noi_tu_turn"] = 1
        history.append({"role": "user", "content": user_word})
        msgs = build_msgs_for_gpt(history, keep_last=20)

        try:
            res = client.chat.completions.create(
                model="gpt-4o",
                messages=msgs,
                temperature=0.9,
                max_tokens=30
            )
            ai_reply = (res.choices[0].message.content or "").strip().lower()
            history.append({"role": "assistant", "content": ai_reply})
            cache_set(game_id, history)
            session["noi_tu_turn"] += 1
            return jsonify({"ai_word": ai_reply})
        except Exception as e:
            return jsonify({"error": "Lỗi khi AI phản hồi: " + str(e)}), 500

    if user_word == "__surrender__":
        session['noi_tu_ended'] = True
        winner = "ai"
        game_record = NoiTuGameHistory(
            user_id=getattr(current_user, "user_id", "anonymous"),
            start_time=session.get('noi_tu_start_time'),
            end_time=datetime.utcnow(),
            total_turns=session.get('noi_tu_turn', 0),
            result=winner,
            reward_used=session.get('noi_tu_reward'),
            user_final_word="(đầu hàng)",
            ai_final_word=""
        )
        db.session.add(game_record); db.session.commit()
        if hasattr(current_user, "user_id"):
            update_leaderboard(current_user.user_id, won=False)

        return jsonify({
            "game_over": True, "winner": winner,
            "history": cache_get(game_id),
            "ai_response": "Bạn đã đầu hàng rồi nhé! Nhưng chơi hay phết á 😁",
            "reward_mode": session.get('noi_tu_reward'),
            "ai_question": get_ai_reward_question()
        }), 200

    if user_word == "" and session.get("noi_tu_turn", 0) == 0:
        first_words = [
            "học sinh","ngôi sao","trí tuệ","tình yêu","cảm xúc","chân thành",
            "đời sống","bầu trời","vầng trăng","mặt trời","cơn mưa","cuộc sống"
        ]
        ai_word = random.choice(first_words)
        session["noi_tu_turn"] = 1
        history.append({"role": "assistant", "content": ai_word})
        cache_set(game_id, history)
        return jsonify({"ai_word": ai_word})

    if user_word_raw == " ":
        session['noi_tu_wrong_count'] = session.get('noi_tu_wrong_count', 0) + 3
        session['noi_tu_ended'] = True
        game_record = NoiTuGameHistory(
            user_id=getattr(current_user, "user_id", "anonymous"),
            start_time=session.get('noi_tu_start_time'),
            end_time=datetime.utcnow(),
            total_turns=session.get('noi_tu_turn', 0),
            result="ai",
            reward_used=session.get('noi_tu_reward'),
            user_final_word="(hết giờ)",
            ai_final_word=""
        )
        db.session.add(game_record); db.session.commit()
        if hasattr(current_user, "user_id"):
            update_leaderboard(current_user.user_id, won=False)

        return jsonify({
            "error": "Bạn đã không nhập kịp trong thời gian giới hạn.",
            "game_over": True, "winner": "ai",
            "history": cache_get(game_id),
            "reward_mode": session.get('noi_tu_reward'),
            "ai_question": get_ai_reward_question()
        }), 200

    # ========== VALIDATIONS GIỮA VÁN ==========
    if canon(user_word) in used_words:
        session['noi_tu_wrong_count'] = session.get('noi_tu_wrong_count', 0) + 1
        if session['noi_tu_wrong_count'] >= 3:
            session['noi_tu_ended'] = True
            return jsonify({
                "error": "Bạn đã nhập cụm từ đã dùng lại quá 3 lần. Bạn thua.",
                "game_over": True, "winner": "ai",
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "ai_question": get_ai_reward_question()
            }), 200
        return jsonify({"error": "Cụm từ này đã được dùng rồi. Vui lòng nhập từ khác."}), 400

    user_parts = user_word.split()
    if len(user_parts) != 2:
        session['noi_tu_wrong_count'] = session.get('noi_tu_wrong_count', 0) + 1
        if session['noi_tu_wrong_count'] >= 3:
            session['noi_tu_ended'] = True
            return jsonify({
                "error": "Bạn đã nhập sai quá 3 lần. Bạn thua.",
                "game_over": True, "winner": "ai",
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "ai_question": get_ai_reward_question()
            }), 200
        return jsonify({"error": "Bạn phải nhập đúng 2 từ có nghĩa. Vui lòng thử lại."}), 400

    ai_last_word = next((h["content"].strip().lower().split()[-1].strip(string.punctuation)
                         for h in reversed(history) if h["role"] == "assistant"), None)
    if ai_last_word and normalize(user_parts[0]) != normalize(ai_last_word):
        session['noi_tu_wrong_count'] = session.get('noi_tu_wrong_count', 0) + 1
        if session['noi_tu_wrong_count'] >= 3:
            session['noi_tu_ended'] = True
            return jsonify({
                "error": "Bạn đã nhập sai quá 3 lần. Bạn thua.",
                "game_over": True, "winner": "ai",
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "ai_question": get_ai_reward_question()
            }), 200
        return jsonify({"error": f"Từ bạn nhập phải bắt đầu bằng từ '{ai_last_word}'. Vui lòng thử lại."}), 400

    try:
        client = create_openai_client("gpt-4o")
        if not is_valid_vi_2word(user_word, client):
            session['noi_tu_wrong_count'] = session.get('noi_tu_wrong_count', 0) + 1
            if session['noi_tu_wrong_count'] >= 3:
                session['noi_tu_ended'] = True
                return jsonify({
                    "error": "Cụm từ bạn nhập không có nghĩa. Bạn đã thua vì nhập sai quá 3 lần.",
                    "game_over": True, "winner": "ai",
                    "history": cache_get(game_id),
                    "reward_mode": session.get('noi_tu_reward'),
                    "ai_question": get_ai_reward_question()
                }), 200
            return jsonify({"error": "Cụm từ bạn nhập không có nghĩa. Vui lòng thử lại."}), 400
    except Exception as e:
        return jsonify({"error": "Lỗi khi kiểm tra nghĩa từ: " + str(e)}), 500

    # OK user -> reset wrong count + lưu vào RAM
    session['noi_tu_wrong_count'] = 0
    history.append({"role": "user", "content": user_word})
    cache_set(game_id, history)

    try:
        need_start = canon(user_word).split()[-1]
        min_sur = session.get('noi_tu_min_surrender_turn', 25)
        turn = session.get('noi_tu_turn', 0)

        # 1) Gọi GPT bằng history đã cắt gọn
        msgs = build_msgs_for_gpt(history, keep_last=20)
        ai_reply = gpt_chat_safe(
            client,
            msgs,
            temperature=0.85,
            max_tokens=30,
            timeout=7.0,
            retries=2,
            backoff=0.6
        )

        # 2) Rescue nếu timeout
        if ai_reply is None:
            rescued = find_backup(need_start, used_words)
            if not rescued:
                rescued = force_retry_with_gpt(need_start, history)
            if rescued:
                ai_reply = rescued
            else:
                if turn < max(min_sur, 30):
                    return jsonify({"error": "Tôi đang hơi lag, bạn nhập lại cụm khác giúp mình nhé 😅"}), 429

                session['noi_tu_ended'] = True
                winner = "user"
                reason = f"Tôi bị 'timeout' khi nghĩ từ bắt đầu bằng '{need_start}'. Thua bạn rồi 😅"

                game_record = NoiTuGameHistory(
                    user_id=getattr(current_user, "user_id", "anonymous"),
                    start_time=session.get('noi_tu_start_time'),
                    end_time=datetime.utcnow(),
                    total_turns=turn,
                    result=winner,
                    reward_used=session.get('noi_tu_reward'),
                    user_final_word=user_word,
                    ai_final_word=""
                )
                db.session.add(game_record); db.session.commit()
                if hasattr(current_user, "user_id"):
                    update_leaderboard(current_user.user_id, won=True)

                return jsonify({
                    "ai_word": "",
                    "game_over": True,
                    "winner": winner,
                    "history": cache_get(game_id),
                    "reward_mode": session.get('noi_tu_reward'),
                    "final_word": user_word,
                    "ai_lost_reason": reason
                }), 200

        ai_reply = ai_reply.strip().lower()

        RARE_AI_WORDS = [
            "đàn đúm","ngoa ngoắt","lim rim","phiêu du","lặng thinh","mịt mù","tối tăm",
            "âm u","trống trải","mơ hồ","lơ đãng","sương khói","xa xăm","mập mờ","chuông chùa","vang vọng",
            "ngập ngừng","ngổn ngang","chênh vênh","lặng lẽ","cầu vồng","lai rai","xí quách"
        ]
        if turn >= 1 and random.random() < 0.6:
            rare_candidates = [w for w in RARE_AI_WORDS if w.split()[0] == need_start]
            if rare_candidates:
                ai_reply = random.choice(rare_candidates)

        if looks_like_surrender(ai_reply) and turn < min_sur:
            rescued = hard_rescue(need_start, used_words, history, client, max_tries=3)
            if rescued:
                ai_reply = rescued

        def is_good(cand: str) -> bool:
            cand_c = canon(cand)
            parts = cand_c.split()
            return (len(parts) == 2 and parts[0] == need_start and cand_c not in used_words)

        if not is_good(ai_reply):
            if turn < min_sur:
                rescued = hard_rescue(need_start, used_words, history, client, max_tries=3)
                if rescued and is_good(rescued):
                    ai_reply = rescued
                else:
                    rescued = force_retry_with_gpt(need_start, history)
                    if rescued and is_good(rescued):
                        ai_reply = rescued

        if not is_good(ai_reply):
            session['noi_tu_ended'] = True
            winner = "user"

            game_record = NoiTuGameHistory(
                user_id=getattr(current_user, "user_id", "anonymous"),
                start_time=session.get('noi_tu_start_time'),
                end_time=datetime.utcnow(),
                total_turns=turn,
                result=winner,
                reward_used=session.get('noi_tu_reward'),
                user_final_word=user_word,
                ai_final_word=ai_reply
            )
            db.session.add(game_record); db.session.commit()
            if hasattr(current_user, "user_id"):
                update_leaderboard(current_user.user_id, won=True)

            user_last_word = need_start
            fallback_messages = [
                f"Tôi không nghĩ ra cụm nào bắt đầu bằng '{user_last_word}' nữa... bạn giỏi thật!",
                f"Tôi bí rồi, chịu luôn từ '{user_last_word}' 🥲",
                f"Hết từ thật rồi. Từ '{user_last_word}' khó quá!",
                f"Bó tay với từ '{user_last_word}' luôn á =))) bạn thắng rồi!",
                f"Tôi chịu... bạn chơi hay thật đấy!",
                f"Bí từ như bí tiền cuối tháng vậy á =)))",
                f"Chơi vậy ai chơi lại 😭. Thua tâm phục khẩu phục!"
            ]
            reason = random.choice(fallback_messages)

            return jsonify({
                "ai_word": ai_reply,
                "game_over": True,
                "winner": winner,
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "final_word": user_word,
                "ai_lost_reason": reason
            }), 200

        if not is_valid_vi_2word(ai_reply, client):
            if session.get('noi_tu_turn', 0) < min_sur:
                rescued = hard_rescue(need_start, used_words, history, client, max_tries=3)
                if rescued and is_valid_vi_2word(rescued, client) and is_good(rescued):
                    ai_reply = rescued
                else:
                    rescued = force_retry_with_gpt(need_start, history)
                    if rescued and is_valid_vi_2word(rescued, client) and is_good(rescued):
                        ai_reply = rescued

        if not is_valid_vi_2word(ai_reply, client) or not is_good(ai_reply):
            session['noi_tu_ended'] = True
            winner = "user"
            reason = f"Tôi bí thật rồi với từ '{need_start}' 😅"

            game_record = NoiTuGameHistory(
                user_id=getattr(current_user, "user_id", "anonymous"),
                start_time=session.get('noi_tu_start_time'),
                end_time=datetime.utcnow(),
                total_turns=session.get('noi_tu_turn', 0),
                result=winner,
                reward_used=session.get('noi_tu_reward'),
                user_final_word=user_word,
                ai_final_word=ai_reply
            )
            db.session.add(game_record); db.session.commit()
            if hasattr(current_user, "user_id"):
                update_leaderboard(current_user.user_id, won=True)

            return jsonify({
                "ai_word": ai_reply,
                "game_over": True,
                "winner": winner,
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "final_word": user_word,
                "ai_lost_reason": reason
            }), 200

        turn = session.get('noi_tu_turn', 0)
        min_sur = session.get('noi_tu_min_surrender_turn', 25)
        profile = session.get('noi_tu_surrender_profile', {"p30_40":0.40,"p40_60":0.20,"p60_100":0.10})

        if random.random() < surrender_prob_by_profile(profile, turn, min_sur):
            session['noi_tu_ended'] = True
            winner = "user"
            user_last_word = need_start
            reason = random.choice([
                f"Tôi không nghĩ ra cụm nào bắt đầu bằng '{user_last_word}' nữa... bạn giỏi thật!",
                f"Tôi bí rồi, chịu luôn từ '{user_last_word}' 🥲",
                f"Hết từ thật rồi. Từ '{user_last_word}' khó quá!",
                f"Bó tay với từ '{user_last_word}' luôn á =))) bạn thắng rồi!",
                f"Tôi chịu... bạn chơi hay thật đấy!",
                f"Bí từ như bí tiền cuối tháng vậy á =)))",
                f"Chơi vậy ai chơi lại 😭. Thua tâm phục khẩu phục!"
            ])

            game_record = NoiTuGameHistory(
                user_id=getattr(current_user, "user_id", "anonymous"),
                start_time=session.get('noi_tu_start_time'),
                end_time=datetime.utcnow(),
                total_turns=turn,
                result=winner,
                reward_used=session.get('noi_tu_reward'),
                user_final_word=user_word,
                ai_final_word=""
            )
            db.session.add(game_record); db.session.commit()
            if hasattr(current_user, "user_id"):
                update_leaderboard(current_user.user_id, won=True)

            return jsonify({
                "ai_word": "",
                "game_over": True,
                "winner": winner,
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "final_word": user_word,
                "ai_lost_reason": reason
            }), 200

        # Hợp lệ -> cập nhật lịch sử/turn
        history.append({"role": "assistant", "content": ai_reply})
        cache_set(game_id, history)
        session['noi_tu_turn'] += 1

        if session['noi_tu_turn'] >= session.get('noi_tu_max_turns', 999):
            session['noi_tu_ended'] = True
            winner = "user"
            user_last_word = need_start
            reason = random.choice([
                f"Tôi không nghĩ ra cụm nào bắt đầu bằng '{user_last_word}' nữa... bạn giỏi thật!",
                f"Tôi bí rồi, chịu luôn từ '{user_last_word}' 🥲",
                f"Hết từ thật rồi. Từ '{user_last_word}' khó quá!",
                f"Bó tay với từ '{user_last_word}' luôn á =))) bạn thắng rồi!",
                f"Tôi chịu... bạn chơi hay thật đấy!",
                f"Từ '{user_last_word}' như đánh sập não tôi vậy 😵",
                f"Bí quá... thôi thua bạn vậy 🫡",
            ])

            game_record = NoiTuGameHistory(
                user_id=getattr(current_user, "user_id", "anonymous"),
                start_time=session.get('noi_tu_start_time'),
                end_time=datetime.utcnow(),
                total_turns=session.get('noi_tu_turn', 0),
                result=winner,
                reward_used=session.get('noi_tu_reward'),
                user_final_word=user_word,
                ai_final_word=""
            )
            db.session.add(game_record); db.session.commit()
            if hasattr(current_user, "user_id"):
                update_leaderboard(current_user.user_id, won=True)

            return jsonify({
                "ai_word": "",
                "game_over": True,
                "winner": winner,
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "final_word": user_word,
                "ai_lost_reason": reason
            }), 200

        ai_reply_lc = ai_reply.lower()
        ai_thua_keywords = ["bạn thắng", "tôi thua", "tôi chịu", "chịu thua", "tôi bí", "hết từ", "bạn giỏi", "không nghĩ ra"]
        ai_thang_keywords = ["bạn thua", "bạn sai", "bạn nhập sai", "bạn lặp lại", "sai luật"]
        is_hint_from_ai = any(k in ai_reply_lc for k in ["cụm từ không hợp lệ", "thử lại", "nhập lại", "không hợp lệ"])

        game_over = False; winner = None
        if any(k in ai_reply_lc for k in ai_thua_keywords) and not is_hint_from_ai:
            game_over = True; winner = "user"
        elif any(k in ai_reply_lc for k in ai_thang_keywords) and not is_hint_from_ai:
            game_over = True; winner = "ai"

        if game_over:
            session['noi_tu_ended'] = True
            game_record = NoiTuGameHistory(
                user_id=getattr(current_user, "user_id", "anonymous"),
                start_time=session.get('noi_tu_start_time'),
                end_time=datetime.utcnow(),
                total_turns=session.get('noi_tu_turn', 0),
                result=winner,
                reward_used=session.get('noi_tu_reward'),
                user_final_word=user_word if winner == "user" else ai_reply,
                ai_final_word=ai_reply
            )
            db.session.add(game_record); db.session.commit()
            if hasattr(current_user, "user_id"):
                update_leaderboard(current_user.user_id, won=(winner == "user"))

            response = {
                "ai_word": ai_reply,
                "game_over": True,
                "winner": winner,
                "history": cache_get(game_id),
                "reward_mode": session.get('noi_tu_reward'),
                "final_word": user_word if winner == "user" else ai_reply
            }
            if winner == "user":
                user_last_word = need_start
                fallback_messages = [
                    f"Tôi không nghĩ ra cụm nào bắt đầu bằng '{user_last_word}' nữa... bạn giỏi thật!",
                    f"Tôi bí rồi, chịu luôn từ '{user_last_word}' 🥲",
                    f"Hết từ thật rồi. Từ '{user_last_word}' khó quá!",
                    f"Bó tay với từ '{user_last_word}' luôn á =))) bạn thắng rồi!",
                    f"Tôi chịu... bạn chơi hay thật đấy!",
                    f"Bí từ như bí tiền cuối tháng vậy á =)))",
                    f"Chơi vậy ai chơi lại 😭. Thua tâm phục khẩu phục!"
                ]
                response["ai_lost_reason"] = random.choice(fallback_messages)
            if winner == "ai":
                response["ai_question"] = get_ai_reward_question()
            return jsonify(response)

        return jsonify({
            "ai_word": ai_reply,
            "game_over": False,
            "history": cache_get(game_id),
            "turn": session['noi_tu_turn']
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/game_noi_tu/bonus_tip")
def bonus_tip():
    import json, re, unicodedata
    from flask import Response

    prompt = """
Bạn đang chuẩn bị gửi một phần thưởng nhỏ sau khi người chơi thắng game nối từ.

🎁 Nội dung có thể là:
- Mẹo cuộc sống thú vị
- Mẹo học tập hiệu quả
- Cách tỏ tình dễ thương
- Gợi ý đề thi hay
- Cách kiếm tiền đơn giản
- Câu chửi văn minh, lịch sự

🎯 Yêu cầu cực kỳ quan trọng:
- Trả về đúng 1 câu gọn, KHÁC với các câu đã từng dùng như: "Học ít mà vẫn điểm cao là ảo tưởng 😅"
- KHÔNG trùng lặp các câu trước (tuyệt đối tránh).
- Trả về kết quả ở định dạng JSON như sau:
{"tip": "mẹo hoặc câu thưởng", "category": "life_tip"}

Nếu là "câu chửi văn minh", hãy dùng đúng mẫu: “Thay vì bạn nói... thì hãy nói...”.

Không giải thích, không thêm nội dung thừa.
Chỉ trả về duy nhất JSON thuần túy.
"""

    try:
        client = create_openai_client("gpt-4o")
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=1.3,
            max_tokens=300
        )
        raw = res.choices[0].message.content.strip()
        match = re.search(r"\{[^{}]+\}", raw, re.DOTALL)
        if not match:
            raise ValueError("Không tìm thấy JSON trong GPT.")

        parsed = json.loads(match.group())
        tip = parsed.get("tip", "").strip()
        category = parsed.get("category", "life_tip")

        def clean_text(text):
            return ''.join(c for c in text if unicodedata.category(c)[0] != 'C')
        tip = clean_text(tip)

        if not tip:
            raise ValueError("Không có tip hợp lệ.")

        # ✅ Trả về tiếng Việt đầy đủ
        return Response(
            json.dumps({"tip": tip, "category": category}, ensure_ascii=False),
            content_type="application/json"
        )

    except Exception as e:
        return Response(
            json.dumps({
                "tip": "Khi người khác giỏi hơn, hãy học hỏi thay vì ganh tị 💡",
                "category": "life_tip"
            }, ensure_ascii=False),
            content_type="application/json"
        )


@app.route('/game_noi_tu/random_ai_question', methods=['GET'])
def get_random_ai_question():
    questions = [
        "Bạn mong muốn thay đổi điều gì trong cuộc sống?",
        "Nếu một ngày mất hết mọi thứ, bạn giữ lại điều gì?",
        "Điều gì khiến bạn cảm thấy mình thật sự đang sống?",
        "Bạn nghĩ hạnh phúc là gì?",
        "Bạn đã từng thấy khoảnh khắc nào khiến bạn nghẹn ngào chưa?",
        "Điều gì bạn luôn biết ơn mỗi ngày?",
        "Nếu chỉ còn 24h để sống, bạn sẽ làm gì?",
        "Bạn chọn bình yên hay thử thách để trưởng thành?",
        "Bạn tin thời gian có thể chữa lành mọi vết thương không?",
        "Sự im lặng có phải đôi khi là câu trả lời tốt nhất?",
        "Bạn có tin vào duyên phận không?",
        "Người bạn thầm thương có thích bạn không?",
        "Câu chuyện của bạn và người ấy sao rồi?",
        "Bạn định nghĩa thế nào là một tình yêu đẹp?",
        "Tình yêu có cần lý do không?",
        "Bạn tin vào tình yêu từ cái nhìn đầu tiên chứ?",
        "Bạn đã bao giờ hy sinh điều gì vì tình yêu?",
        "Bạn nghĩ tình yêu và tình thương khác nhau chỗ nào?",
        "Điều gì khiến bạn cảm thấy an toàn trong một mối quan hệ?",
        "Bạn có tin tình yêu có thể thay đổi con người không?",
        "Bạn từng hối tiếc điều gì nhất?",
        "Nếu có thể làm lại, bạn muốn làm gì khác?",
        "Bạn nghĩ ý nghĩa cuộc đời là gì?",
        "Sự tự do thật sự nằm ở đâu?",
        "Bạn tin mỗi người đều có một sứ mệnh riêng chứ?",
        "Cái giá của sự trưởng thành là gì?",
        "Tiền có mua được hạnh phúc không?",
        "Điều gì khó buông bỏ nhất trong đời?",
        "Bạn chọn sống thật hay sống theo kỳ vọng của người khác?",
        "Điều gì khiến con người ta trở nên mạnh mẽ hơn?",
        "Mục tiêu lớn nhất của bạn là gì?",
        "Bạn muốn làm gì trong 5 năm tới?",
        "Nếu được ước một điều, bạn sẽ ước gì?",
        "Bạn muốn trở thành người như thế nào trong tương lai?",
        "Nếu có cơ hội sống ở bất kỳ đâu, bạn chọn nơi nào?",
        "Bạn nghĩ thế giới 20 năm nữa sẽ ra sao?",
        "Bạn muốn để lại điều gì cho thế hệ sau?",
        "Điều gì bạn muốn học ngay lập tức nếu có thể?",
        "Nếu được gặp bản thân trong tương lai, bạn sẽ hỏi gì?",
        "Bạn nghĩ mình sẽ thay đổi ra sao sau 10 năm?",
        "Điều gì ở gia đình khiến bạn tự hào nhất?",
        "Ai là người bạn luôn muốn bảo vệ?",
        "Bạn học được điều gì quý giá nhất từ gia đình?",
        "Khoảnh khắc nào khiến bạn nhớ về gia đình nhiều nhất?",
        "Bạn muốn tạo ra truyền thống gì cho gia đình mình?",
        "Ai là người truyền cảm hứng cho bạn nhiều nhất?",
        "Bạn nghĩ gia đình quan trọng thế nào trong hành trình của bạn?",
        "Nếu có thể, bạn muốn nói gì với gia đình ngay lúc này?",
        "Kỷ niệm gia đình nào bạn luôn trân trọng?",
        "Điều gì khiến bạn cảm thấy hạnh phúc khi nghĩ về gia đình?",
    ]
    return jsonify({"question": random.choice(questions)})

@app.route('/game_noi_tu/answer_feedback', methods=['POST'])
def answer_feedback():
    data = request.get_json()
    game_id = data.get("game_id")
    user_answer = data.get("answer", "").strip()
    question = data.get("question", "").strip()

    if not user_answer:
        return jsonify({"ok": False, "feedback": "Bạn chưa nhập gì cả!"})

    try:
        prompt = f"""
Tôi là AI đang chơi game nối từ cùng người dùng. Sau khi tôi thắng, tôi đã hỏi người chơi một câu thưởng:

📌 Câu hỏi là: "{question}"

Người dùng đã trả lời: "{user_answer}"

Bạn hãy phản hồi lại một cách dí dỏm, nhẹ nhàng, thông minh, tối đa 1-2 câu.
👉 Nếu câu trả lời KHÔNG LIÊN QUAN tới câu hỏi, hãy nói kiểu: "Ủa alo bạn đọc câu hỏi chưa vậy?", "Hahaha không liên quan gì luôn á 😅" hoặc trêu nhẹ nhàng.
👉 Nếu người chơi trả lời CÓ LIÊN QUAN, hãy phản hồi dễ thương, thông minh, vui vẻ như một người bạn.

Giữ phong cách gần gũi, nói như bạn trẻ GenZ. Không đạo lý.
"""

        client = create_openai_client("gpt-4o")
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=1.3,
            max_tokens=300
        )

        feedback = res.choices[0].message.content.strip()
        return jsonify({"ok": True, "feedback": feedback})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "feedback": f"Có lỗi xảy ra: {str(e)}"}), 500

def get_ai_reward_question():
    return random.choice([
    # --- Cuộc sống ---
    "Bạn mong muốn thay đổi điều gì trong cuộc sống?",
    "Nếu một ngày mất hết mọi thứ, bạn giữ lại điều gì?",
    "Điều gì khiến bạn cảm thấy mình thật sự đang sống?",
    "Bạn nghĩ hạnh phúc là gì?",
    "Bạn đã từng thấy khoảnh khắc nào khiến bạn nghẹn ngào chưa?",
    "Điều gì bạn luôn biết ơn mỗi ngày?",
    "Nếu chỉ còn 24h để sống, bạn sẽ làm gì?",
    "Bạn chọn bình yên hay thử thách để trưởng thành?",
    "Bạn tin thời gian có thể chữa lành mọi vết thương không?",
    "Sự im lặng có phải đôi khi là câu trả lời tốt nhất?",
    
    # --- Tình yêu ---
    "Bạn có tin vào duyên phận không?",
    "Người bạn thầm thương có thích bạn không?",
    "Câu chuyện của bạn và người ấy sao rồi?",
    "Bạn định nghĩa thế nào là một tình yêu đẹp?",
    "Tình yêu có cần lý do không?",
    "Bạn tin vào tình yêu từ cái nhìn đầu tiên chứ?",
    "Bạn đã bao giờ hy sinh điều gì vì tình yêu?",
    "Bạn nghĩ tình yêu và tình thương khác nhau chỗ nào?",
    "Điều gì khiến bạn cảm thấy an toàn trong một mối quan hệ?",
    "Bạn có tin tình yêu có thể thay đổi con người không?",
    
    # --- Triết lý ---
    "Bạn từng hối tiếc điều gì nhất?",
    "Nếu có thể làm lại, bạn muốn làm gì khác?",
    "Bạn nghĩ ý nghĩa cuộc đời là gì?",
    "Sự tự do thật sự nằm ở đâu?",
    "Bạn tin mỗi người đều có một sứ mệnh riêng chứ?",
    "Cái giá của sự trưởng thành là gì?",
    "Tiền có mua được hạnh phúc không?",
    "Điều gì khó buông bỏ nhất trong đời?",
    "Bạn chọn sống thật hay sống theo kỳ vọng của người khác?",
    "Điều gì khiến con người ta trở nên mạnh mẽ hơn?",
    
    # --- Tương lai ---
    "Mục tiêu lớn nhất của bạn là gì?",
    "Bạn muốn làm gì trong 5 năm tới?",
    "Nếu được ước một điều, bạn sẽ ước gì?",
    "Bạn muốn trở thành người như thế nào trong tương lai?",
    "Nếu có cơ hội sống ở bất kỳ đâu, bạn chọn nơi nào?",
    "Bạn nghĩ thế giới 20 năm nữa sẽ ra sao?",
    "Bạn muốn để lại điều gì cho thế hệ sau?",
    "Điều gì bạn muốn học ngay lập tức nếu có thể?",
    "Nếu được gặp bản thân trong tương lai, bạn sẽ hỏi gì?",
    "Bạn nghĩ mình sẽ thay đổi ra sao sau 10 năm?",
    
    # --- Gia đình ---
    "Điều gì ở gia đình khiến bạn tự hào nhất?",
    "Ai là người bạn luôn muốn bảo vệ?",
    "Bạn học được điều gì quý giá nhất từ gia đình?",
    "Khoảnh khắc nào khiến bạn nhớ về gia đình nhiều nhất?",
    "Bạn muốn tạo ra truyền thống gì cho gia đình mình?",
    "Ai là người truyền cảm hứng cho bạn nhiều nhất?",
    "Bạn nghĩ gia đình quan trọng thế nào trong hành trình của bạn?",
    "Nếu có thể, bạn muốn nói gì với gia đình ngay lúc này?",
    "Kỷ niệm gia đình nào bạn luôn trân trọng?",
    "Điều gì khiến bạn cảm thấy hạnh phúc khi nghĩ về gia đình?"
])


@app.route('/game_noi_tu/submit_user_answer', methods=['POST'])
def submit_user_answer():
    data = request.get_json()
    answer = data.get("answer", "")
    return jsonify({"ok": True, "feedback": f"Cảm ơn bạn đã chia sẻ: {answer}"})
@app.route('/game_noi_tu/reward_question', methods=['POST'])
def reward_question():
    data = request.get_json()
    question = data.get("question", "")
    if not question:
        return jsonify({"error": "Thiếu câu hỏi"}), 400

    try:
        client = create_openai_client("gpt-4o")
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Bạn là một AI vừa thua người chơi trong một game nối từ. "
                        "Hãy trả lời câu hỏi sau một cách trực tiếp, súc tích, đúng trọng tâm. "
                        "Không cần liệt kê, không cần giải thích dài dòng. "
                        "Trả lời đầy đủ trong một đoạn duy nhất, không vượt quá 200 từ."
                    )
                },
                {"role": "user", "content": question}
            ],
            temperature=0.7,
            max_tokens=500 
        )
        answer = res.choices[0].message.content.strip()
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/game_noi_tu/history', methods=['GET'])
def get_noi_tu_history():
    try:
        user_id = current_user.user_id if hasattr(current_user, "user_id") else None
        if not user_id:
            return jsonify({"error": "Bạn chưa đăng nhập"}), 401

        history_list = (
            NoiTuGameHistory.query
            .filter_by(user_id=user_id)
            .order_by(NoiTuGameHistory.start_time.desc())
            .limit(10)
            .all()
        )

        return jsonify([h.to_dict() for h in history_list])

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/noi_tu')
def redirect_noi_tu_random():
    random_id = str(uuid.uuid4())[:8]  
    return redirect(f"/noi_tu/{random_id}")


@app.route("/noi_tu/<game_id>")
def game_noi_tu(game_id):
    username = session.get("username")
    
    if not username:
        return redirect("/login") 
    
    user = User.query.filter_by(username=username).first()
    
    if not user:
        return redirect("/error_user_not_found")  
    
    avatar_url = user.avatar_url if user.avatar_url else "/static/images/avatars/user_default.png"
    
    return render_template(
        "noi_tu.html",
        session_id=game_id,
        avatar_url=avatar_url,
        fullname=user.fullname,
        user_id=user.user_id,            
        username=user.username    
    )

@app.route('/game_noi_tu/leaderboard')
def get_leaderboard():
    try:
        # 1. Truy vấn top 20 người chơi
        results = db.session.execute(text("""
            SELECT user_id, total_games, wins AS total_wins,
                   ROUND(CASE WHEN total_games > 0 THEN (wins * 100.0 / total_games) ELSE 0 END, 1) AS win_rate
            FROM noi_tu_leaderboard
            ORDER BY wins DESC, win_rate DESC
            LIMIT 20
        """)).mappings().all()

        user_ids = [r['user_id'] for r in results]
        users_info = {}
        if user_ids:
            rows = db.session.execute(text("""
                SELECT user_id, fullname, avatar_url
                FROM users
                WHERE user_id = ANY(:ids)
            """), {'ids': user_ids}).mappings().all()
            users_info = {
                r['user_id']: {
                    "fullname": r['fullname'],
                    "avatar_url": r['avatar_url']
                } for r in rows
            }

        leaderboard = []
        for idx, row in enumerate(results, 1):
            info = users_info.get(row['user_id'])
            leaderboard.append({
                "rank": idx,
                "user_id": row['user_id'],
                "fullname": info["fullname"],
                "avatar_url": info["avatar_url"],
                "total_games": row['total_games'],
                "total_wins": row['total_wins'],
                "win_rate": row['win_rate']
            })

        # 2. Nếu là người đăng nhập, lấy xếp hạng thật
        your_rank_data = None
        if current_user and current_user.is_authenticated:
            rank_result = db.session.execute(text("""
                WITH ranked AS (
                    SELECT user_id,
                           RANK() OVER (ORDER BY wins DESC, total_games DESC) AS rank
                    FROM noi_tu_leaderboard
                )
                SELECT rank FROM ranked WHERE user_id = :uid
            """), {'uid': current_user.user_id}).fetchone()

            if rank_result:
                user_info = db.session.execute(text("""
                    SELECT wins, total_games FROM noi_tu_leaderboard WHERE user_id = :uid
                """), {'uid': current_user.user_id}).fetchone()

                your_rank_data = {
                    "rank": rank_result.rank,
                    "fullname": current_user.fullname,
                    "avatar_url": current_user.avatar_url,
                    "total_wins": user_info.wins,
                    "total_games": user_info.total_games
                }

        return jsonify({
            "leaderboard": leaderboard,
            "your_rank": your_rank_data
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Lỗi bảng xếp hạng: {str(e)}"}), 500
@app.route("/game_noi_tu/gopy_ai", methods=["POST"])
def gopy_ai():
    data = request.get_json()

    content = data.get("content", "").strip()
    user_id = data.get("user_id", "n/a").strip()
    username = data.get("username", "n/a").strip()

    if not content:
        return jsonify({"success": False, "message": "Nội dung trống"})

    log = f"""---
Thời gian: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Phân loại: Gợi ý câu hỏi AI
Họ tên: (ẩn danh)
Email: (ẩn danh)
User ID: {user_id}
Username: {username}
Nội dung: {content}
Ảnh:
"""

    try:
        with open("feedback_log.txt", "a", encoding="utf-8") as f:
            f.write(log + "\n")
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "message": "Lỗi server"})

@app.route("/api/tra_cuu", methods=["POST"])
def tra_cuu_tu_nghia():
    if not session.get("noi_tu_ended", False):
        return jsonify({"error": "Chỉ được tra cứu sau khi kết thúc game."}), 403

    data = request.get_json()

    def normalize_vietnamese(text):
        return unicodedata.normalize('NFC', text.strip().lower())

    tu = normalize_vietnamese(data.get("tu", ""))

    # ✅ Chặn nếu không phải đúng 2 từ
    if len(tu.split()) != 2:
        return jsonify({"error": "Chỉ được tra cứu các cụm từ gồm đúng 2 từ."}), 403

    try:
        with open("viet_dictionary.json", "r", encoding="utf-8") as f:
            dictionary = json.load(f)

        normalized_keys = {normalize_vietnamese(k): k for k in dictionary.keys()}
        real_key = normalized_keys.get(tu)
        if real_key:
            entry = dictionary[real_key]

            nghia = entry.get("nghia") or entry.get("nghĩa") or "Không rõ."
            vi_du = entry.get("vi_du") or entry.get("ví_dụ") or "Không rõ."
            dong_nghia = entry.get("dong_nghia") or entry.get("đồng_nghĩa") or "Không có."
            cap_nhat = entry.get("cap_nhat") or entry.get("cập_nhật") or ""
            nguoi_cung_cap = entry.get("nguoi_cung_cap")
            nguon = nguoi_cung_cap if nguoi_cung_cap else entry.get("nguon", "Cộng đồng")


            return jsonify({
                "tu": real_key,
                "nghia": nghia,
                "vi_du": vi_du,
                "dong_nghia": dong_nghia,
                "cap_nhat": cap_nhat,
                "nguon": nguon,
                "nguoi_cung_cap": nguoi_cung_cap
            })

        # Nếu không có trong từ điển, gọi GPT
        client = create_openai_client("gpt-4o")
        prompt = f"""Giải thích nghĩa của cụm từ sau bằng tiếng Việt dễ hiểu, có ví dụ và các từ đồng nghĩa nếu có.

Từ: "{tu}"

Trả lời theo mẫu JSON:
{{
  "nghia": "...",
  "vi_du": "...",
  "dong_nghia": "..."
}}"""
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=300
        )
        reply = res.choices[0].message.content.strip()
        try:
            parsed = json.loads(reply)
        except Exception:
            import re
            nghia = re.search(r'"nghia"\s*:\s*"([^"]+)"', reply)
            vi_du = re.search(r'"vi_du"\s*:\s*"([^"]+)"', reply)
            dong_nghia = re.search(r'"dong_nghia"\s*:\s*"([^"]+)"', reply)

            parsed = {
                "nghia": nghia.group(1) if nghia else "Không rõ nghĩa.",
                "vi_du": vi_du.group(1) if vi_du else "Không rõ ví dụ.",
                "dong_nghia": dong_nghia.group(1) if dong_nghia else "Không có đồng nghĩa."
            }

        parsed["tu"] = tu
        parsed["nguon"] = random.choice([
            "Mai Linh", "Ngọc Hân", "Minh Quân", "Tuấn Kiệt", "Khánh Vy", "AKAZA", "Kukoshibo", "Muichiro",
            "Bảo An", "Trúc Lam", "Phương Thảo", "Gia Hưng", "Thảo Nhi", "Seven", "Lucie", "Muzan", "Gia Bảo"
        ])
        parsed["cap_nhat"] = (datetime.now() - timedelta(days=random.randint(1, 30))).strftime("%Y-%m-%d %H:%M:%S")

        return jsonify(parsed)

    except Exception as e:
        return jsonify({"error": f"Lỗi tra cứu: {e}"}), 500
@app.route("/server_status")
@admin_only
def server_status():
    return render_template("server_status.html")
@app.route('/rules')
def rules():
    return render_template("rules.html")
import os
import time
from flask import jsonify

UPLOADS_FOLDER = os.path.join("static", "images", "uploads")

@app.route("/cleanup_old_ai_images", methods=["POST"])
def cleanup_old_ai_images():
    try:
        now = time.time()
        deleted_files = []
        kept_files = []

        # Duyệt qua toàn bộ file trong uploads/
        for root, dirs, files in os.walk(UPLOADS_FOLDER):
            for file in files:
                file_path = os.path.join(root, file)
                if not os.path.isfile(file_path):
                    continue

                # Thời gian chỉnh sửa cuối
                file_age_days = (now - os.path.getmtime(file_path)) / (60 * 60 * 24)

                if file_age_days >= 30:  # quá 30 ngày hoặc đúng 30 ngày
                    os.remove(file_path)
                    deleted_files.append(file_path.replace("static/", ""))
                else:
                    kept_files.append(file_path.replace("static/", ""))

        return jsonify({
            "success": True,
            "deleted": deleted_files,
            "kept": kept_files,
            "message": f"Đã xóa {len(deleted_files)} ảnh quá hạn (>=30 ngày)."
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

# ====== RUN APP ======
if __name__ == "__main__":
    import os
    if os.environ.get("WERKZEUG_RUN_MAIN") == "true":
        threading.Thread(target=auto_unblock_loop, daemon=True).start()
    socketio.run(app, host="0.0.0.0", port=5000, debug=True, allow_unsafe_werkzeug=True)
